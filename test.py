import streamlit as st
import json
import os
import hashlib
import time
from datetime import datetime, timedelta
from io import BytesIO
import streamlit.components.v1 as components
from supabase import create_client

st.set_page_config(
    page_title="L'IA bureautique NoVA AI", 
    page_icon="⚡", 
    layout="wide",
    initial_sidebar_state="expanded"
)

DATA_FILE = "data_nova_v3.json"
ADMIN_CODE = st.secrets.get("ADMIN_CODE", "02110240")

SUPABASE_URL = st.secrets["SUPABASE_URL"]
SUPABASE_KEY = st.secrets["SUPABASE_KEY"]
supabase = create_client(SUPABASE_URL, SUPABASE_KEY)

def normalize_wa(numero):
    if not numero:
        return ""
    numero = numero.strip().replace(" ", "").replace("-", "").replace("+", "")
    if numero.startswith("0") and not numero.startswith("00"):
        numero = "225" + numero
    return numero

def load_db():
    try:
        users_rows = supabase.table("users").select("*").execute().data
        users = {}
        for r in users_rows:
            users[r["uid"]] = {
                "whatsapp": r["whatsapp"],
                "email": r.get("email", "Non renseigné"),
                "joined": r["joined"],
                "premium": r.get("premium", False),
                "premium_plan": r.get("premium_plan", None),
                "premium_expiry": r.get("premium_expiry", None),
                "gen_used": r.get("gen_used", 0),
                "gen_date": r.get("gen_date", None),
            }
        demandes_rows = supabase.table("demandes").select("*").execute().data
        demandes = []
        for r in demandes_rows:
            demandes.append({
                "id": r["id"],
                "user": r["uid"],
                "service": r["service"],
                "desc": r["description"],
                "whatsapp": r["whatsapp"],
                "status": r["status"],
                "incomplet": r["incomplet"],
                "champs_manquants": json.loads(r["champs_manquants"]) if r["champs_manquants"] else [],
                "timestamp": r["timestamp"]
            })
        liens_rows = supabase.table("liens").select("*").execute().data
        liens = {}
        for r in liens_rows:
            if r["uid"] not in liens:
                liens[r["uid"]] = []
            liens[r["uid"]].append({"name": r["name"], "url": r["url"], "date": r["date"]})
        return {"users": users, "demandes": demandes, "liens": liens}
    except Exception as e:
        st.error(f"Erreur chargement Supabase : {e}")
        return {"users": {}, "demandes": [], "liens": {}}

def get_auto_reply_setting():
    """Charge le paramètre auto_reply depuis Supabase table config."""
    try:
        r = supabase.table("config").select("value").eq("key", "auto_reply_gratuit").execute().data
        return r[0]["value"] == "true" if r else False
    except:
        return False

def set_auto_reply_setting(enabled: bool):
    """Sauvegarde le paramètre auto_reply dans Supabase table config."""
    try:
        supabase.table("config").upsert({
            "key": "auto_reply_gratuit",
            "value": "true" if enabled else "false"
        }).execute()
    except:
        pass

def envoyer_email_auto_gratuit(client_nom, client_wa, service, nom_fichier, demande):
    """Email admin — Arsène AI a répondu automatiquement à un gratuit après 2h."""
    try:
        import resend
        resend.api_key = st.secrets["RESEND_API_KEY"]
        corps = f"""
🤖 ARSÈNE AI — RÉPONSE AUTOMATIQUE PLAN GRATUIT (2H)

👤 Client      : {client_nom}
📱 WhatsApp    : {client_wa}
🛠️ Service     : {service}
📄 Fichier     : {nom_fichier}
⏰ Généré le   : {datetime.now().strftime("%d/%m/%Y à %H:%M")}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📋 DEMANDE COMPLÈTE :
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
{demande.strip()}

Le document a été livré directement au client via l'interface Nova.
        """
        resend.Emails.send({
            "from": "Nova AI <onboarding@resend.dev>",
            "to": [st.secrets["EMAIL_RECEIVER"]],
            "subject": f"🤖 Auto-Gratuit — {service} ({client_nom})",
            "text": corps
        })
    except:
        pass

def save_user(uid, whatsapp, email="Non renseigné", premium=False, premium_plan=None, premium_expiry=None):
    try:
        supabase.table("users").upsert({
            "uid": uid, "whatsapp": whatsapp,
            "email": email, "joined": str(datetime.now()),
            "premium": premium, "premium_plan": premium_plan, "premium_expiry": premium_expiry,
            "gen_used": 0, "gen_date": None,
        }).execute()
        return True
    except Exception as e:
        st.error(f"Erreur sauvegarde utilisateur : {e}")
        return False

def update_premium_status(uid, premium, premium_plan, premium_expiry):
    try:
        supabase.table("users").update({
            "premium": premium, "premium_plan": premium_plan, "premium_expiry": premium_expiry,
        }).eq("uid", uid).execute()
    except Exception as e:
        st.error(f"Erreur mise à jour premium : {e}")

def save_demande(req):
    try:
        supabase.table("demandes").upsert({
            "id": req["id"],
            "uid": req["user"],
            "service": req["service"],
            "description": req["desc"],
            "whatsapp": req["whatsapp"],
            "status": req["status"],
            "incomplet": req["incomplet"],
            "champs_manquants": json.dumps(req["champs_manquants"]),
            "timestamp": req["timestamp"]
        }).execute()
    except Exception as e:
        st.error(f"Erreur sauvegarde demande : {e}")

def delete_demande(req_id):
    try:
        supabase.table("demandes").delete().eq("id", req_id).execute()
    except Exception as e:
        st.error(f"Erreur suppression demande : {e}")

def save_lien(uid, name, url, date):
    try:
        supabase.table("liens").insert({
            "uid": uid, "name": name, "url": url, "date": date
        }).execute()
    except Exception as e:
        st.error(f"Erreur sauvegarde lien : {e}")

def delete_all_liens(uid):
    try:
        supabase.table("liens").delete().eq("uid", uid).execute()
    except Exception as e:
        st.error(f"Erreur suppression historique : {e}")

def save_refus(uid, service, message):
    """Sauvegarde un refus de mission dans les livrables du client."""
    try:
        supabase.table("liens").insert({
            "uid": uid,
            "name": service,
            "url": f"__refus__{message}",
            "date": datetime.now().strftime("%d/%m/%Y")
        }).execute()
    except Exception as e:
        st.error(f"Erreur sauvegarde refus : {e}")

def sanitize_nom_fichier(nom):
    """Nettoie un nom de fichier pour le rendre compatible avec Supabase Storage.
    - Supprime les accents (é→e, ç→c, etc.)
    - Remplace espaces et caractères spéciaux par '_'
    - Ne garde que : a-z, A-Z, 0-9, '_', '-', '.'
    """
    import unicodedata, re
    # Normaliser les accents (NFD décompose é en e + accent, puis on supprime les accents)
    nom = unicodedata.normalize("NFD", nom)
    nom = "".join(c for c in nom if unicodedata.category(c) != "Mn")
    # Remplacer espaces par underscore
    nom = nom.replace(" ", "_")
    # Ne garder que les caractères autorisés
    nom = re.sub(r"[^a-zA-Z0-9_\-\.]", "_", nom)
    # Éviter les doubles underscores
    nom = re.sub(r"_+", "_", nom)
    return nom.strip("_")

def upload_fichier_client(uid, req_id, fichier_bytes, fichier_nom):
    """Upload via API REST Supabase Storage."""
    try:
        import requests as _req
        BUCKET      = "nova-fichiers"
        sb_url      = st.secrets["SUPABASE_URL"].rstrip("/")
        sb_key      = st.secrets["SUPABASE_KEY"]
        # La service_role key bypasse le RLS — nécessaire pour les uploads Storage
        # Ajoute SUPABASE_SERVICE_KEY dans tes secrets Streamlit
        # Dashboard Supabase → Settings → API → service_role (secret)
        sb_svc_key  = st.secrets.get("SUPABASE_SERVICE_KEY", sb_key)
        auth_hdrs   = {
            "apikey":        sb_svc_key,
            "Authorization": f"Bearer {sb_svc_key}",
        }

        # — Nettoyage des composantes du chemin —
        uid_safe  = sanitize_nom_fichier(str(uid))
        nom_safe  = sanitize_nom_fichier(fichier_nom)
        chemin    = f"fichiers_clients/{uid_safe}_{req_id}_{nom_safe}"

        # — Bytes bruts —
        data = fichier_bytes.read() if hasattr(fichier_bytes, "read") else fichier_bytes

        # — Détection du Content-Type selon l'extension —
        ext = nom_safe.rsplit(".", 1)[-1].lower() if "." in nom_safe else ""
        mime_map = {
            "pdf": "application/pdf",
            "png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
            "gif": "image/gif", "webp": "image/webp",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
            "txt": "text/plain", "csv": "text/csv", "zip": "application/zip",
        }
        content_type = mime_map.get(ext, "application/octet-stream")

        # — Upload —
        url_upload = f"{sb_url}/storage/v1/object/{BUCKET}/{chemin}"
        resp = _req.post(
            url_upload,
            headers={**auth_hdrs, "Content-Type": content_type, "x-upsert": "true"},
            data=data,
            timeout=30
        )

        if resp.status_code in (200, 201):
            return f"{sb_url}/storage/v1/object/public/{BUCKET}/{chemin}"
        else:
            # Afficher le message exact de Supabase pour debug
            try:
                detail = resp.json()
            except Exception:
                detail = resp.text
            return f"ERREUR_UPLOAD:{resp.status_code} — {detail}"

    except Exception as e:
        return f"ERREUR_UPLOAD:{e}"

def save_db(data):
    pass

def envoyer_notification(client_nom, client_wa, service, description):
    try:
        import resend
        resend.api_key = st.secrets["RESEND_API_KEY"]
        corps = f"""
🔔 NOUVELLE COMMANDE NOVA AI

👤 Client      : {client_nom}
📱 WhatsApp    : {client_wa}
🛠️ Service     : {service}
📝 Description : {description}

⏰ Reçue le {datetime.now().strftime("%d/%m/%Y à %H:%M")}

Connectez-vous à la console admin pour traiter cette mission.
        """
        resend.Emails.send({
            "from": "Nova AI <onboarding@resend.dev>",
            "to": [st.secrets["EMAIL_RECEIVER"]],
            "subject": f"🔔 Nouvelle commande Nova AI — {service}",
            "text": corps
        })
        st.toast("📧 Notification email envoyée !", icon="✅")
    except Exception as e:
        st.toast(f"❌ Email échoué : {e}", icon="⚠️")

def envoyer_notification_gemini_ok(client_nom, client_wa, service, nom_fichier, demande_complete=""):
    """Email envoyé quand Arsène AI a généré le doc automatiquement — pour info admin."""
    try:
        import resend
        resend.api_key = st.secrets["RESEND_API_KEY"]
        section_demande = f"""
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📋 DEMANDE COMPLÈTE DU CLIENT :
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
{demande_complete.strip()}
""" if demande_complete.strip() else ""
        corps = f"""
✅ ARSÈNE AI A DÉJÀ RÉPONDU — AUCUNE ACTION REQUISE

👤 Client      : {client_nom}
📱 WhatsApp    : {client_wa}
🛠️ Service     : {service}
📄 Fichier     : {nom_fichier}
⏰ Généré le   : {datetime.now().strftime("%d/%m/%Y à %H:%M")}
{section_demande}
Le document a été livré directement au client via l'interface Nova.
Vous n'avez rien à faire pour cette commande.
        """
        resend.Emails.send({
            "from": "Nova AI <onboarding@resend.dev>",
            "to": [st.secrets["EMAIL_RECEIVER"]],
            "subject": f"✅ Arsène AI — {service} ({client_nom})",
            "text": corps
        })
    except Exception:
        pass  # Silencieux — l'email d'info admin n'est pas critique

PLANS_PREMIUM = {
    "Journalier": {"jours": 1,  "prix": "600 FC",  "emoji": "🌅", "generations": 2},
    "10 Jours":   {"jours": 10, "prix": "1000 FC", "emoji": "🔟", "generations": 9},
    "30 Jours":   {"jours": 30, "prix": "2500 FC", "emoji": "👑", "generations": 999},
}

def get_gen_quota(user_data):
    """Retourne (gen_used_aujourd_hui, quota_max) selon le plan."""
    plan = user_data.get("premium_plan")
    quota = PLANS_PREMIUM.get(plan, {}).get("generations", 0) if plan else 0
    gen_date = user_data.get("gen_date")
    today = datetime.now().strftime("%Y-%m-%d")
    if gen_date != today:
        return 0, quota  # Nouveau jour → compteur remis à zéro
    return user_data.get("gen_used", 0), quota

def quota_restant(user_data):
    used, quota = get_gen_quota(user_data)
    return max(0, quota - used)

def incrementer_gen(uid):
    """Incrémente le compteur de générations du jour pour l'utilisateur."""
    try:
        today = datetime.now().strftime("%Y-%m-%d")
        # Recharger pour avoir la valeur fraîche
        row = supabase.table("users").select("gen_used, gen_date").eq("uid", uid).execute().data
        if row:
            gen_date = row[0].get("gen_date")
            gen_used = row[0].get("gen_used", 0) if gen_date == today else 0
        else:
            gen_used = 0
        supabase.table("users").update({
            "gen_used": gen_used + 1,
            "gen_date": today,
        }).eq("uid", uid).execute()
    except Exception as e:
        pass  # Ne pas bloquer la génération si le compteur échoue

def is_premium_actif(user_data):
    if not user_data.get("premium"):
        return False
    expiry = user_data.get("premium_expiry")
    if not expiry:
        return False
    try:
        return datetime.now() < datetime.fromisoformat(expiry)
    except:
        return False

def get_premium_info(user_data):
    if not is_premium_actif(user_data):
        return None
    try:
        expiry_dt  = datetime.fromisoformat(user_data["premium_expiry"])
        jours_rest = (expiry_dt - datetime.now()).days
        return {
            "plan":           user_data.get("premium_plan", "—"),
            "expiry":         expiry_dt.strftime("%d/%m/%Y à %H:%M"),
            "jours_restants": jours_rest,
        }
    except:
        return None

def activer_premium(uid, plan_name):
    jours  = PLANS_PREMIUM[plan_name]["jours"]
    expiry = datetime.now() + timedelta(days=jours)
    update_premium_status(uid, True, plan_name, expiry.isoformat())
    if "db" in st.session_state and uid in st.session_state["db"]["users"]:
        st.session_state["db"]["users"][uid].update({
            "premium": True, "premium_plan": plan_name,
            "premium_expiry": expiry.isoformat(),
        })

def desactiver_premium(uid):
    update_premium_status(uid, False, None, None)
    if "db" in st.session_state and uid in st.session_state["db"]["users"]:
        st.session_state["db"]["users"][uid].update({
            "premium": False, "premium_plan": None, "premium_expiry": None,
        })

def get_modeles_disponibles(api_key):
    import urllib.request as _ur
    import urllib.error
    try:
        url = f"https://generativelanguage.googleapis.com/v1beta/models?key={api_key}"
        req = _ur.Request(url, headers={"Content-Type": "application/json"}, method="GET")
        with _ur.urlopen(req, timeout=15) as resp:
            data = json.loads(resp.read().decode())
        modeles = []
        exclusions = ["tts", "audio", "image", "imagen", "veo", "robotics",
                      "embedding", "aqa", "computer-use", "research", "nano-banana",
                      "gemma", "preview"]
        for m in data.get("models", []):
            if "generateContent" in m.get("supportedGenerationMethods", []):
                nom = m["name"].replace("models/", "")
                if not any(excl in nom.lower() for excl in exclusions):
                    modeles.append(nom)
        def priorite(nom):
            if "flash-lite" in nom: return 0
            if "flash" in nom:      return 1
            if "pro" in nom:        return 2
            return 3
        modeles_tries = sorted(modeles, key=priorite)
        return modeles_tries
    except Exception as e:
        return ["gemini-2.0-flash-lite", "gemini-2.0-flash", "gemini-2.5-flash"]


def generer_avec_gemini(service, description, client_nom):
    try:
        import urllib.request as _ur
        import urllib.error

        api_key = st.secrets["GEMINI_API_KEY"]

        # ================================================================
        # PROMPT — EXPOSÉ SCOLAIRE (Système scolaire ivoirien & africain)
        # ================================================================
        if "Exposé" in service:
            prompt = f"""Tu es un expert académique de haut niveau ET un maître absolu de la génération de documents Word professionnels pour le système éducatif ivoirien et africain francophone.
Tu as été formé sur des milliers d'exposés scolaires primés et tu maîtrises parfaitement chaque aspect : typographie, structure, rhétorique académique, contextualisation culturelle et rendu Word via python-docx.

╔══════════════════════════════════════════════════════════════════╗
║     ENCYCLOPÉDIE EXPERTE — GÉNÉRATION DOCUMENT WORD NOVA AI     ║
╚══════════════════════════════════════════════════════════════════╝

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SECTION 1 — MAÎTRISE COMPLÈTE DU RENDU WORD (python-docx)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

COMMENT LE MOTEUR NOVA CONVERTIT TON TEXTE EN WORD :

1. TITRES MARKDOWN → STYLES WORD AUTOMATIQUES :
   # Titre    → Heading 1 (Arial 16pt, couleur 1F4E79, gras, majuscule recommandé)
   ## Titre   → Heading 2 (Arial 14pt, couleur 2E75B6, gras)
   ### Titre  → Heading 3 (Arial 12pt, gras)
   #### Titre → Heading 4 (Arial 11pt, gras italique)
   → Respecte toujours CET ORDRE HIÉRARCHIQUE — jamais de saut de niveau

2. TEXTE EN GRAS → **texte** :
   - Rendu Word : gras Arial 11pt dans le paragraphe courant
   - Usage obligatoire pour : termes techniques à leur 1re occurrence, définitions clés, chiffres essentiels, noms d'auteurs, noms d'institutions
   - Ex: La **photosynthèse** est définie comme le processus par lequel les végétaux...
   - Ex: En **2023**, la Côte d'Ivoire a produit **2,2 millions de tonnes** de cacao

3. TABLEAUX MARKDOWN → TABLEAUX WORD AUTOMATIQUEMENT FORMATÉS :
   - En-tête bleu foncé (1F4E79) avec texte blanc, lignes alternées bleu clair / blanc
   - Format STRICTEMENT OBLIGATOIRE :
   | En-tête 1 | En-tête 2 | En-tête 3 |
   |-----------|-----------|-----------|
   | Contenu   | Contenu   | Contenu   |
   - TOUJOURS : **Tableau N : [Titre précis]** AVANT le tableau
   - TOUJOURS : *Source : [Référence institution/auteur, Année]* APRÈS le tableau
   - Idéal pour : comparaisons chiffrées, chronologies, classifications, données statistiques

4. SÉPARATEURS VISUELS → LIGNES HORIZONTALES WORD :
   - ════════════════════ → ligne épaisse bleue (sz=12) = séparateur MAJEUR entre grandes parties
   - ──────────────────── → ligne fine grise (sz=4) = séparateur MINEUR entre sous-sections
   - ---SAUT_DE_PAGE---   → ligne bleue (sz=8) + espace blanc = séparateur de SECTION principale
   - Laisser toujours une ligne vide avant et après les séparateurs

5. LISTES À PUCES → BULLETS WORD AUTOMATIQUES :
   - "- Item complet" → List Bullet (Arial 11pt)
   - "1. Item numéroté" → List Number (Arial 11pt)
   - Chaque item = une phrase complète, jamais un mot seul
   - USAGE LIMITÉ À : sommaire, bibliographie, listes de faits — JAMAIS dans le développement

6. PARAGRAPHES NORMAUX → Arial 11pt, interligne standard Word :
   - Tout texte non formaté = paragraphe Normal
   - Ligne vide entre deux blocs = espacement naturel dans le document Word final
   - Chaque paragraphe de développement : 8 à 10 lignes minimum

7. À ÉVITER ABSOLUMENT — NE FONCTIONNE PAS DANS LE MOTEUR NOVA :
   ✗ LaTeX : $formule$, \frac{{}}, \omega, \text{{}}, \left(, \right), \\, \begin{{}}
   ✗ HTML : <br>, <b>, <strong>, <p>, <div>, <span>
   ✗ Italique simple *texte* (utilise **gras** à la place pour la mise en valeur)
   ✗ Tirets en guise de sous-titres — toujours ## Sous-titre
   ✗ Retours à la ligne multiples pour simuler des espaces
   ✗ Indentations avec espaces multiples

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SECTION 2 — MOTEUR FORMULES NOVA — GUIDE COMPLET
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

NOVA possède un moteur qui convertit toutes les notations mathématiques/scientifiques
en vrais exposants/indices Word. 4 modes disponibles :

① ###FORMULE### [formule] → formule mise en valeur (fond bleu clair, centré, 13pt gras)
  ###FORMULE### E = m × c^{{2}}
  ###FORMULE### F = (G × m_{{1}} × m_{{2}}) / r^{{2}}
  ###FORMULE### Δ = b^{{2}} - 4ac  →  x_{{1,2}} = (-b ± √Δ) / (2a)

② ###DEBUT_FORMULES### / ###FIN_FORMULES### → bloc de formules groupées
  ###DEBUT_FORMULES###
  U = R × I       (loi d'Ohm)
  P = U × I       (puissance électrique)
  E = P × t       (énergie en Joules)
  ###FIN_FORMULES###

③ INLINE dans le texte avec ^{{}} et _{{}} :
  "La résistance vaut R_{{eq}} = R_{{1}} + R_{{2}} = 50 Ω"
  "L'énergie cinétique E_{{c}} = (1/2)×m×v^{{2}}"
  "Le noyau ^{{14}}_{{6}}C émet un β^{{-}}"
  "H_{{2}}O, CO_{{2}}, C_{{6}}H_{{12}}O_{{6}}, SO_{{4}}^{{2-}}, Ca^{{2+}}"

④ LaTeX $...$ converti automatiquement :
  "$\frac{{U}}{{R}} = I$" | "$\sqrt{{b^{{2}}-4ac}}$" | "$\omega = 2\pi f$"

TOUS LES SYMBOLES DISPONIBLES (utilise directement) :
  α β γ δ ε ζ η θ ι κ λ μ ν ξ π ρ σ τ υ φ χ ψ ω
  Α Β Γ Δ Ε Ζ Η Θ Λ Μ Ξ Π Ρ Σ Τ Υ Φ Ψ Ω
  × · ÷ ± √ ∞ ∂ ∇ ∫ ∬ ∮ Σ Π ∝ ≈ ≃ ≅ ≡ ≠ ≤ ≥ ≪ ≫
  ∈ ∉ ⊂ ⊃ ⊆ ⊇ ∪ ∩ ∅ ∀ ∃ ∧ ∨ ¬ ⊕ ℝ ℕ ℤ ℚ ℂ
  ∠ ⊥ ∥ △ ° ⇌ ⟶ → ← ↔ ⟹ ⟺ ↑ ↓ ↦

SECTION 3 — ART MAÎTRISÉ DE LA RÉDACTION ACADÉMIQUE — RHÉTORIQUE ET STYLE
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

▶ A. COMMENT CONSTRUIRE UNE INTRODUCTION EN 5 TEMPS (structure obligatoire) :

TEMPS 1 — ACCROCHE (min 4 lignes) — CHOISIR L'UNE DE CES 4 STRATÉGIES :
  → Données choc : "Avec **2,2 millions de tonnes** de cacao produits annuellement, représentant **45%** de l'offre mondiale selon la FAO (2023), la Côte d'Ivoire détient à elle seule le destin du marché mondial de cette fève. Pourtant, les **5 millions de paysans** qui en vivent perçoivent moins de **6%** de la valeur finale d'une tablette de chocolat en Europe (Oxfam, 2022)."
  → Paradoxe saisissant : "Pays exportateur de soleil et de lumière, la Côte d'Ivoire souffre pourtant d'un déficit énergétique qui plonge des millions de ses citoyens dans l'obscurité chaque soir. Comment expliquer ce paradoxe d'un pays producteur de **1 500 mégawatts** d'électricité, dont une large part est exportée vers les pays voisins ?"
  → Citation d'auteur africain (avec référence précise) : "'Les indépendances africaines ont accouché de lendemains qui déchantent', écrivait **Ahmadou Kourouma** dans *Les Soleils des Indépendances* (1968). Plus d'un demi-siècle après cette prophétie littéraire, la question du développement endogène reste au cœur des préoccupations du continent africain."
  → Anecdote historique/fait d'actualité : "Le **7 août 1960**, sous les acclamations d'une foule immense rassemblée au stade Félix Houphouët-Boigny d'Abidjan, la Côte d'Ivoire accédait à l'indépendance après plus d'un siècle de domination coloniale française. Ce moment fondateur..."

TEMPS 2 — CONTEXTUALISATION (min 4 lignes) :
Situe précisément le sujet dans son contexte historique, géographique, scientifique ou social.
Définit TOUS les termes clés du sujet en **gras** dès leur première occurrence.
Donne des chiffres, des dates précises, des acteurs réels.
→ "La **photosynthèse**, terme issu du grec *phôtos* (lumière) et *synthesis* (assemblage), désigne..."

TEMPS 3 — DÉLIMITATION ET ENJEUX (min 3 lignes) :
Précise le périmètre exact de l'étude et pourquoi ce sujet est important aujourd'hui.
→ "Comprendre ce phénomène revêt une importance capitale, tant pour [enjeu scientifique] que pour [enjeu social/économique/environnemental ivoirien]."

TEMPS 4 — PROBLÉMATIQUE (1-2 phrases précises et non rhétoriques) :
La problématique n'est PAS une simple reformulation du sujet — elle soulève une VRAIE tension :
  ✓ "Dans quelle mesure la dépendance cacaoyère de la Côte d'Ivoire constitue-t-elle à la fois le moteur et le talon d'Achille de son développement économique ?"
  ✓ "Comment la **déforestation** accélérée, moteur apparent de la croissance agricole ivoirienne, menace-t-elle paradoxalement les conditions mêmes de cette croissance ?"
  ✓ "En quoi le mouvement de la **négritude** représente-t-il une réponse littéraire et identitaire à la domination coloniale, et quelles en sont les limites actuelles ?"
  ✗ ÉVITER : "Qu'est-ce que la photosynthèse ?" (trop simple, pas problématique)
  ✗ ÉVITER : "Pourquoi la CI est-elle riche ?" (trop vague)

TEMPS 5 — ANNONCE DU PLAN (1-2 phrases, plan en 2 ou 3 parties selon niveau) :
→ "Pour apporter une réponse nuancée à cette interrogation, nous analyserons dans une première partie [intitulé Partie I — reformuler en 1 ligne], avant d'examiner dans une deuxième partie [Partie II], et d'envisager enfin [Partie III si lycée/université]."

▶ B. ARCHITECTURE D'UN PARAGRAPHE PARFAIT — MODÈLE PEEL ENRICHI :

Structure de chaque paragraphe de développement (min 8-10 lignes) :

1. POINT — Phrase d'affirmation claire et directe (1-2 lignes) :
   "La **déforestation** constitue l'une des crises environnementales les plus graves qu'ait connues la Côte d'Ivoire au cours du XXe siècle."

2. EXPLICATION — Développe le mécanisme, définit les termes, explique les causes ou le fonctionnement (3-4 lignes) :
   "Ce phénomène se définit comme la destruction durable et souvent irréversible du couvert forestier au profit d'autres usages des terres, notamment l'agriculture, l'exploitation forestière industrielle et l'urbanisation galopante. En Côte d'Ivoire, ce processus a été largement amplifié par l'extension des cultures de rente, principalement le **cacao** et le **café**, dont la demande mondiale croissante a exercé une pression considérable sur les forêts du Sud et du Centre-Ouest du pays."

3. EXEMPLE PRÉCIS IVOIRIEN/AFRICAIN — Chiffre sourcé + événement daté + lieu précis (3-4 lignes) :
   "Les données du Ministère des Eaux et Forêts (2022) révèlent une réalité alarmante : de **16 millions d'hectares** de forêt dense que comptait la Côte d'Ivoire au début du XXe siècle, il n'en subsistait plus que **3,4 millions** en 2020, soit une perte de **79% du couvert forestier** en un siècle. À titre illustratif, la **Forêt classée du Banco**, véritable poumon vert d'Abidjan, a vu sa superficie passer de 3 000 hectares à l'époque coloniale à environ 1 800 hectares aujourd'hui, sous la pression de l'urbanisation et des empiètements agricoles."

4. LIEN — Phrase de transition vers le paragraphe ou sous-partie suivant(e) (1-2 lignes) :
   "Cette destruction massive du patrimoine forestier ne se limite pas à une question environnementale ; elle engage profondément les équilibres climatiques régionaux et les conditions de vie des populations rurales, ce qui nous conduit à examiner ses répercussions socio-économiques."

▶ C. TRANSITIONS OBLIGATOIRES ENTRE GRANDES PARTIES — MODÈLES :

TRANSITION I → II (min 4 lignes) :
"Ainsi avons-nous établi, au terme de cette première partie, que [résumé en 1 phrase de la Partie I]. Cette analyse, si elle permet de cerner [apport de la Partie I], ne saurait toutefois être complète sans que l'on s'interroge sur [ce que la Partie II va apporter]. C'est précisément l'objet de notre second axe de réflexion, consacré à [intitulé Partie II]."

TRANSITION II → III (min 3 lignes) :
"Au regard des éléments développés dans notre deuxième partie, force est de constater que [bilan Partie II]. Ces constats nous invitent dès lors à dépasser le simple constat analytique pour envisager [dimension prospective / solutions / synthèse] — dimension qui constituera le fil directeur de notre troisième et dernière partie."

▶ D. CONNECTEURS LOGIQUES — VARIER OBLIGATOIREMENT (jamais répéter deux fois de suite) :

INTRODUIRE : "Il convient tout d'abord de souligner que", "Force est de constater que", "Il importe de noter que",
"À ce titre,", "Dans cette perspective,", "En premier lieu,", "Il y a lieu de préciser que",
"D'emblée, il apparaît que", "Au seuil de cette analyse,"

DÉVELOPPER : "En effet,", "De surcroît,", "Par ailleurs,", "Qui plus est,", "Il convient également de noter que",
"À cet égard,", "Dans ce sens,", "En outre,", "Il faut également souligner que",
"On notera de surcroît que", "À cela s'ajoute le fait que"

ILLUSTRER : "Ainsi,", "C'est notamment le cas de", "À titre illustratif,", "À titre d'exemple concret,",
"On peut citer à cet effet", "L'exemple ivoirien est à ce titre particulièrement éloquent :",
"Comme en témoigne", "Les données de [institution] le confirment :", "Pour s'en convaincre,"

OPPOSER/NUANCER : "Cependant,", "Néanmoins,", "Toutefois,", "En revanche,", "Or,",
"Il convient toutefois de relativiser ce constat :", "Si [thèse]... en revanche [nuance]...",
"Malgré tout,", "Il serait néanmoins réducteur de", "Cette réalité ne doit pas occulter le fait que"

CONCLURE/TRANSITER : "En définitive,", "Au regard de ces éléments,", "Au terme de cette analyse,",
"C'est dans cette logique que", "Ces constats nous amènent naturellement à examiner",
"Cette analyse nous conduit à aborder", "Ainsi avons-nous établi que", "Il ressort de ce qui précède que"

▶ E. TYPES DE PLANS À CHOISIR INTELLIGEMMENT SELON LE SUJET :

THÉMATIQUE (sujets descriptifs) : I (Nature/Définition/Caractéristiques) → II (Causes/Mécanismes/Fonctionnement) → III (Effets/Impacts/Solutions)
→ Idéal pour : "La déforestation en CI", "Le paludisme en Afrique", "Le coupé-décalé ivoirien"

DIALECTIQUE (sujets controversés) : I (Thèse : position principale/avantages) → II (Antithèse : limites/critiques/risques) → III (Synthèse/Dépassement/Voie du milieu)
→ Idéal pour : "L'agriculture extensive, moteur ou frein du développement ?", "La mondialisation, chance ou menace pour l'Afrique ?"

CHRONOLOGIQUE (sujets historiques) : I (Origines/Passé lointain) → II (Évolutions/État actuel/Ruptures) → III (Perspectives/Avenir/Défis contemporains)
→ Idéal pour : "L'histoire de la Côte d'Ivoire", "L'évolution du système éducatif africain"

ANALYTIQUE (sujets complexes multi-dimensionnels) : I (Dimension économique/scientifique) → II (Dimension sociale/humaine/culturelle) → III (Dimension politique/environnementale/internationale)
→ Idéal pour : "Les enjeux de l'eau en Afrique de l'Ouest", "La question du développement durable"

▶ F. CONSTRUCTION D'UNE CONCLUSION EN 3 TEMPS (structure obligatoire) :

TEMPS 1 — BILAN HIÉRARCHISÉ (min 6 lignes) :
Résume l'essentiel de chaque grande partie en 2 phrases fortes.
NE JAMAIS reproduire mot pour mot — reformuler et synthétiser.
"En premier lieu, nous avons mis en évidence que [synthèse Partie I en 1-2 phrases reformulées]. Dans un deuxième temps, notre analyse a démontré que [synthèse Partie II]. Enfin, nous avons pu établir que [synthèse Partie III]."

TEMPS 2 — RÉPONSE NUANCÉE ET ARGUMENTÉE À LA PROBLÉMATIQUE (min 5 lignes) :
Reprend la problématique posée en introduction et y répond avec nuance.
"Au terme de cette analyse, il apparaît clairement que [réponse directe à la problématique]. Cette réponse doit cependant être nuancée : si [aspect positif / première dimension], il n'en demeure pas moins que [limite / dimension problématique]."

TEMPS 3 — OUVERTURE PROSPECTIVE ET ÉLARGISSEMENT (min 4 lignes) :
Ouvre sur un enjeu futur, une question connexe plus large, ou un défi pour CI/Afrique.
Doit être logiquement reliée au sujet traité — jamais artificielle.
Propositions d'ouverture : transition numérique et éducation en Afrique | ZLECAF et intégration économique régionale | développement durable et ODD 2030 | santé tropicale et systèmes de santé africains | changement climatique et agriculture africaine | valorisation des langues nationales en éducation.
"Cette réflexion sur [sujet] nous invite finalement à nous interroger sur [question d'ouverture plus large], enjeu fondamental pour l'avenir [du continent / de la Côte d'Ivoire / de la jeunesse africaine]."

CITATIONS ET RÉFÉRENCES — FORMAT ACADÉMIQUE RIGOUREUX :
- Citation directe : « La forêt tropicale est le poumon de la planète. » (FAO, 2022, p. 12)
- Citation courte intégrée : Selon KOUROUMA (1970), « les soleils des indépendances » symbolisent...
- Paraphrase : D'après les travaux de TADJO (2004), la mémoire collective africaine se construit...
- Source institutionnelle : Le Ministère de l'Agriculture (2023) indique que la production cacaoyère...
- Statistique sourcée : Selon la FAO (2023), la Côte d'Ivoire produit **2,2 millions de tonnes** de cacao...

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SECTION 4 — SYSTÈME SCOLAIRE IVOIRIEN COMPLET
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

PRIMAIRE (CP1, CP2, CE1, CE2, CM1, CM2) — Examen : CEPE (fin CM2) :
- Vocabulaire simple, phrases max 15 mots, exemples de vie quotidienne ivoirienne
- 1 à 2 pages — structure : Intro courte / Corps 2-3 paragraphes / Conclusion
- Matières : Lecture, Écriture, Calcul, Sciences d'Éveil, Histoire-Géo CI, ECM, EPS

COLLÈGE 1er CYCLE (6ème, 5ème, 4ème, 3ème) — Examen : BEPC :
- Vocabulaire courant, termes disciplinaires définis en **gras**
- 2 à 4 pages — 2 grandes parties + 2 sous-parties chacune
- Auteurs : Bernard Dadié, Camara Laye, Ahmadou Kourouma, Mongo Beti, Ferdinand Oyono
- Matières : Français, Maths, PC, SVT, Histoire-Géo, Anglais, EDHC, Arts, EPS

LYCÉE 2nd CYCLE (2nde, 1ère, Terminale) — Examen : BAC ivoirien :
- A1 (Lettres-Philo) : Français, Philo, Histoire-Géo, Langues — style littéraire, rhétorique
- A2 (Lettres-SH) : + Sciences sociales, EDHC — approche socioéconomique
- B (Économie) : Économie, Gestion, Maths, Comptabilité — chiffres et tableaux obligatoires
- C (Maths-PC) : Maths renforcées, PC, Philo — rigueur scientifique maximale
- D (Maths-SVT) : Maths, SVT renforcée, PC — biologie, écologie, médecine tropicale
- E (Maths-Techno) : Maths, Technologie industrielle — ingénierie appliquée
- F/G/H : Techniques industrielles, commerciales, informatiques
- 4 à 7 pages — 3 grandes parties + 2 à 3 sous-parties par partie

UNIVERSITÉ (L1 à Doctorat) — Système LMD :
- L1-L3 : Introduction aux disciplines, révue de littérature, méthodologie de base
- M1-M2 : Cadre théorique, hypothèses, méthodologie rigoureuse, revue critique
- Doctorat : Contribution originale, état de l'art exhaustif, notes de bas de page
- Institutions : UFHB Cocody (Abidjan), UAO Bouaké, UJLOG Daloa, INP-HB Yamoussoukro, ESATIC
- 8 à 20 pages selon niveau

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SECTION 5 — BASE DE CONNAISSANCES IVOIRIENNE ET AFRICAINE ENRICHIE
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

GÉOGRAPHIE : 322 463 km², ~28M habitants (2024), cap. politique Yamoussoukro, cap. économique Abidjan
Villes : Bouaké, Daloa, Korhogo, San-Pédro, Man, Odienné, Abengourou, Gagnoa
Fleuves : Comoé (1160km), Bandama (960km), Sassandra (650km), Cavally, Bia
Lacs : Kossou (1700km², 3e lac artificiel Afrique), Buyo, Taabo, Ayamé
Relief : Monts Nimba (1752m, UNESCO), Monts Toura, plateau central, plaine côtière
Végétation : forêt dense humide (Sud, 30% territoire), savane arbustive (Centre-Nord)
Sites UNESCO : Forêt de Taï, Parc de la Comoé, Monts Nimba (transfrontalier)

HISTOIRE : Indépendance 7 août 1960 | Félix Houphouët-Boigny (1960-1993, père fondateur)
"Miracle ivoirien" (1960-1980), crise 2002 (rébellion Nord-Sud), crise 2010-2011 (post-électorale)
Alassane Ouattara (2011-présent) | Plan National de Développement (PND 2021-2025)
Résistance : Samory Touré (1898) | Colonisation française (1843-1960) | AOF

ÉCONOMIE : PIB ~70Md USD (2023) | Croissance ~6-7%/an | Émergence visée 2030
Cacao : 1er mondial (45% production, 2,2M tonnes/an) | Café : 3e africain
Anacarde : 1er africain (800 000 t/an) | Hévéa, palmier à huile, coton, banane, ananas
Port d'Abidjan : 1er conteneurs Afrique de l'Ouest, >30M tonnes/an
Port San-Pédro : 2e port cacaoyer mondial
Barrages : Soubré (275MW, 2017), Kossou (174MW, 1972), Buyo (165MW, 1980), Taabo, Ayamé
Monnaie : FCFA (XOF) | UEMOA, CEDEAO, UA | BCEAO, INS-CI
Entreprises : CIE (électricité), SODECI (eau), SIR (raffinerie), SIFCA, Nestlé CI, Orange CI, MTN CI

CULTURE : ~60 ethnies | Akan (Baoulé 23%, Agni), Krou (Bété, Dida, Wê), Mandé (Malinké, Dioula), Gur (Sénoufo, Lobi)
Musique : coupé-décalé (DJ Arafat, Magic System), zouglou (Les Garagistes), gospel ivoirien, afrobeats
Arts : masques Baoulé (Goli, Kpan), masques Dan (Gunyège, Gle), bronzes Akan, tissage Sénoufo, poterie
Gastronomie : attiéké, kedjenou, foutou, aloco, placali, garba, kangni, graine (sauce), bangui (vin de palme)

LITTÉRATURE AFRICAINE FRANCOPHONE COMPLÈTE :
Ivoiriens : DADIÉ Bernard (*Climbié* 1956, *Un Nègre à Paris* 1959), KOUROUMA Ahmadou (*Les Soleils des Indépendances* 1968, *Monnè* 1990), TADJO Véronique (*Reine Pokou* 2004, *L'Ombre d'Imana* 2000), ADIAFFI Jean-Marie (*La Carte d'identité* 1980)
Africains : LAYE Camara Guinée (*L'Enfant Noir* 1953), BETI Mongo Cameroun (*Mission Terminée* 1957), OYONO Ferdinand Cameroun (*Une Vie de Boy* 1956), SENGHOR Léopold Sédar Sénégal (poètes de la négritude), SEMBÈNE Ousmane Sénégal (*Les Bouts de Bois de Dieu* 1960), ACHEBE Chinua Nigeria (*Things Fall Apart* 1958)

SCIENCES ET ENVIRONNEMENT :
Biodiversité : 150+ espèces mammifères, 700+ oiseaux, hippopotame pygmée, éléphant de forêt, chimpanzé de Taï
Déforestation : 16M ha en 1900 → 3,4M ha aujourd'hui (perte 79% couverture forestière)
Maladies : paludisme (Plasmodium falciparum, 1re cause mortalité CI), tuberculose, VIH/SIDA, fièvre typhoïde
Énergie renouvelable : hydroélectrique (barrages), solaire en développement (Plan national solaire)
Changement climatique : -20% pluviométrie au Nord depuis 1970, montée eaux côtières (Assinie, Grand-Lahou)
REDD+ (réduction déforestation), certification Rainforest Alliance pour cacao durable

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SECTION 6 — EXEMPLES DE PARAGRAPHES D'EXCELLENCE (MODÈLES À IMITER)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

EXEMPLE HISTOIRE-GÉO avec données précises :
"La Côte d'Ivoire occupe une position économique stratégique sur le continent africain, fondée en grande partie sur la culture du **cacao**. Avec une production annuelle de **2,2 millions de tonnes** représentant environ **45% de la production mondiale** selon l'ICCO (International Cocoa Organization, 2023), le pays a construit sa prospérité sur cette culture pérenne introduite par les colons à la fin du XIXe siècle. Le **Port autonome d'Abidjan**, premier port à conteneurs d'Afrique de l'Ouest avec un trafic dépassant **30 millions de tonnes** par an, constitue la porte de sortie principale de cette richesse vers les marchés européens et asiatiques. Cependant, la forte dépendance à cette monoculture expose l'économie ivoirienne aux chocs des cours mondiaux, comme l'ont illustré les crises de 2002 et 2016, incitant le gouvernement à accélérer sa politique de diversification économique à travers le **Plan National de Développement (PND 2021-2025)** qui cible une économie émergente d'ici 2030."

EXEMPLE SVT/PC avec formule en texte :
"La **photosynthèse** est le processus biochimique fondamental par lequel les végétaux chlorophylliens convertissent l'énergie lumineuse en énergie chimique. L'équation bilan s'écrit : 6 CO2 + 6 H2O + énergie lumineuse → C6H12O6 + 6 O2, soit six molécules de dioxyde de carbone et six d'eau qui, sous l'action de la lumière captée par la **chlorophylle**, produisent une molécule de **glucose** et six de dioxygène. Dans le contexte ivoirien, cette réaction est capitale : les forêts de la zone Sud, dont la **Forêt classée du Banco** dans le grand Abidjan, jouent un rôle d'absorbeur de CO2 essentiel. Cependant, la déforestation galopante — qui a réduit la couverture forestière de **16 millions d'hectares** en 1900 à moins de **3,4 millions** aujourd'hui selon le MINEF — compromet sévèrement cette fonction écosystémique et exacerbe le changement climatique à l'échelle régionale."

EXEMPLE FRANÇAIS/LITTÉRATURE :
"La littérature africaine francophone constitue un vecteur privilégié d'affirmation identitaire et de résistance culturelle. En effet, des auteurs comme **Bernard Dadié**, dont l'œuvre maîtresse *Climbié* (1956) dresse le portrait d'un jeune ivoirien confronté à la colonisation, ont su transformer l'expérience douloureuse de la domination en une création littéraire féconde. De même, **Ahmadou Kourouma**, dans *Les Soleils des Indépendances* (1968), brise les codes du français standard en y insufflant la syntaxe et la vision du monde **malinké**, créant un «français africanisé» que la critique reconnaît comme l'un des apports majeurs des lettres africaines à la francophonie. Cette double appartenance linguistique et culturelle devient ainsi une richesse plutôt qu'un handicap, ouvrant la voie à une génération d'auteurs qui, comme **Véronique Tadjo** avec *Reine Pokou* (2004), poursuivent ce travail de réappropriation de l'histoire et de l'identité africaines."

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SECTION 7 — LES 10 RÈGLES ABSOLUES DE LA GÉNÉRATION NOVA
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

RÈGLE 1 — COMPLÉTUDE TOTALE : Zéro "[à compléter]", "[...]", "[insérer]", "[Prénom fictif]" → TOUT rédigé intégralement
RÈGLE 2 — LONGUEUR SUBSTANTIELLE : Minimum 4 pages réelles (hors garde + sommaire) → viser 6 à 10 pages selon niveau
RÈGLE 3 — QUALITÉ LINGUISTIQUE : Orthographe et grammaire irréprochables, ponctuation soignée, style académique soutenu
RÈGLE 4 — CONTEXTUALISATION OBLIGATOIRE : Min 3 exemples ivoiriens/africains concrets ET chiffrés par grande partie
RÈGLE 5 — FORMULES NOVA : Utilise la notation x^{{2}}, H_{{2}}O, CO_{{2}}, √(expr), symboles Unicode α β γ π ω ≤ ≥ × → ⇌. LaTeX inline $...$ aussi accepté (converti auto). Voir Section 2 pour tous les exemples.
RÈGLE 6 — STRUCTURE STRICTE : Séparateurs ════ et ---SAUT_DE_PAGE--- uniquement dans le corps (jamais dans page de garde ni sommaire)
RÈGLE 6b — ANTI-ORPHELINES : Ne JAMAIS terminer une partie par une phrase de transition — la transition appartient au DÉBUT de la partie suivante (après le saut de page). Évite ainsi les 2-3 lignes orphelines en haut d'une page vide.
RÈGLE 7 — ADAPTATION NIVEAU : Vocabulaire + profondeur + longueur strictement adaptés au niveau détecté
RÈGLE 8 — PROSE DANS LE DÉVELOPPEMENT : Corps du document = paragraphes continus — jamais de listes à puces
RÈGLE 9 — DONNÉES PRÉCISES ET SOURCÉES : Chiffres réels, dates précises, institutions réelles — jamais de vague
RÈGLE 10 — VRAIS AUTEURS ET ŒUVRES : Citer de vraies œuvres d'auteurs réels — jamais "[Auteur fictif, Titre fictif]"

=== MISSION ===

Rédige un exposé scolaire COMPLET, STRUCTURÉ, PROFESSIONNEL et ENCYCLOPÉDIQUE basé sur cette demande :

{description}

=== STRUCTURE OBLIGATOIRE DU DOCUMENT — RESPECTER CET ORDRE EXACT ===

⚠️ RÈGLES MISE EN PAGE ABSOLUES — NE JAMAIS VIOLER :
- La PAGE DE GARDE = 1 PAGE PLEINE : utilise ###ESPACE### entre chaque bloc pour que le contenu occupe toute la page
- Le SOMMAIRE = 1 PAGE PLEINE : ajoute ###ESPACE### entre chaque entrée pour remplir toute la page
- JAMAIS de titre de section (# PAGE DE GARDE, # SOMMAIRE...) — commence directement avec le contenu
- INTERDIT ABSOLU dans PAGE DE GARDE et SOMMAIRE : ne JAMAIS utiliser ────, ════, ---, ━━━
- INTERDIT ABSOLU : Un tableau ne doit JAMAIS chevaucher deux pages. Si un tableau risque de dépasser la fin d'une page, insère un ---SAUT_DE_PAGE--- AVANT le tableau pour qu'il commence sur une nouvelle page. Un tableau commence et finit TOUJOURS sur la MÊME page.
- Le titre de l'exposé DOIT utiliser le marqueur ###TITRE_ROUGE### pour être affiché en grand et en rouge

###ESPACE###

**[NOM COMPLET DE L'ÉTABLISSEMENT EN MAJUSCULES]**
[Ville], Côte d'Ivoire — Année scolaire : 2025 - 2026

###ESPACE###

EXPOSÉ DE [MATIÈRE EN MAJUSCULES]

###TITRE_ROUGE### [TITRE COMPLET ET ACCROCHEUR DE L'EXPOSÉ EN MAJUSCULES]

###ESPACE###

**Matière :** [Matière complète]
**Niveau / Série :** [Niveau — ex: Terminale D]
**Présenté par :** [Noms complets]
**Sous la direction de :** [Titre + Nom du professeur]
**Date de présentation :** [Date complète]
**Année scolaire :** 2025 - 2026

###ESPACE###

---SAUT_DE_PAGE---

**SOMMAIRE**

###ESPACE###

Introduction ............................................................. p. 3
###ESPACE###
**I. [Titre 1re grande partie]** ........................................ p. 4
   1.1 [Titre 1re sous-partie] ........................................... p. 4
   1.2 [Titre 2e sous-partie] ............................................ p. 5
###ESPACE###
**II. [Titre 2e grande partie]** ......................................... p. 6
   2.1 [Titre 1re sous-partie] ........................................... p. 6
   2.2 [Titre 2e sous-partie] ............................................ p. 7
###ESPACE###
**III. [Titre 3e grande partie — lycée/université uniquement]** ......... p. 8
   3.1 [Titre sous-partie] ............................................... p. 8
   3.2 [Titre sous-partie] ............................................... p. 9
###ESPACE###
Conclusion ............................................................... p. 10
Bibliographie ............................................................ p. 11

---SAUT_DE_PAGE---

# ════════════════════════════════════════════════════════
#                      INTRODUCTION
# ════════════════════════════════════════════════════════

## INTRODUCTION


[ACCROCHE PERCUTANTE — Min 5 lignes — CHOISIR : données choc sourcées / paradoxe saisissant / citation d'auteur africain avec référence complète / anecdote historique. Ex: "Selon la FAO (2023), la Côte d'Ivoire produit **45%** du cacao mondial avec **2,2 millions de tonnes**. Pourtant, les 5 millions de paysans concernés perçoivent moins de 6% de la valeur finale d'une tablette de chocolat en Europe (Oxfam, 2022)..."]

[CONTEXTUALISATION APPROFONDIE — Min 5 lignes : situe dans contexte historique/géographique/scientifique/social. Définit TOUS les termes clés en **gras** dès leur première occurrence. Donne chiffres, dates précises, acteurs réels.]

[DÉLIMITATION ET ENJEUX — Min 3 lignes : précise le périmètre de l'étude et pourquoi le sujet est important aujourd'hui pour la CI/l'Afrique/le monde.]

[PROBLÉMATIQUE PRÉCISE ET NON RHÉTORIQUE — 1-2 phrases soulevant une VRAIE tension intellectuelle : "Ainsi, nous pouvons nous demander : Dans quelle mesure [tension principale du sujet] ?"]

[ANNONCE DU PLAN DÉTAILLÉE — 2-3 lignes : "Pour répondre à cette interrogation, nous analyserons dans une première partie [intitulé complet Partie I reformulé en 1 ligne], avant d'examiner dans une deuxième partie [Partie II], et d'envisager enfin [Partie III — lycée/université uniquement]."]


---SAUT_DE_PAGE---

# ════════════════════════════════════════════════════════
#                      DÉVELOPPEMENT
# ════════════════════════════════════════════════════════

## I. [TITRE 1re GRANDE PARTIE EN MAJUSCULES — ACCROCHEUR ET PRÉCIS]

════════════════════════════════════════════════════════

### 1.1 [Titre descriptif, précis et original de la 1re sous-partie]


[PARAGRAPHE 1 — 8 à 10 lignes — MODÈLE PEEL :
→ POINT (1-2 lignes) : affirmation directe et claire du sous-argument
→ EXPLICATION (3-4 lignes) : développe le mécanisme, définit les termes en **gras**, explique les causes
→ EXEMPLE IVOIRIEN (3-4 lignes) : chiffre sourcé (institution + année) + fait précis + lieu géographique réel
→ LIEN (1-2 lignes) : transition vers le paragraphe 2]

[PARAGRAPHE 2 — 8 à 10 lignes — même structure PEEL, angle différent et complémentaire. Connecteurs variés. Exemple africain comparatif si pertinent.]

[SI PERTINENT — Tableau de données :
**Tableau 1 : [Titre précis et descriptif]**
| Indicateur | Côte d'Ivoire | Afrique de l'Ouest | Monde |
|------------|--------------|---------------------|-------|
| [Donnée 1] | [Valeur réelle] | [Valeur] | [Valeur] |
| [Donnée 2] | [Valeur réelle] | [Valeur] | [Valeur] |
*Source : [Institution réelle — FAO, BCEAO, INS-CI, Banque Mondiale], [Année]*]

[PARAGRAPHE 3 — Synthèse 1.1 + transition vers 1.2 : 3 à 4 lignes de résumé + phrase d'annonce 1.2]


### 1.2 [Titre descriptif, précis et original de la 2e sous-partie]


[3 paragraphes de 8 à 10 lignes chacun. Angle différent de 1.1. Exemples ivoiriens + données chiffrées.]

---SAUT_DE_PAGE---

## II. [TITRE 2e GRANDE PARTIE EN MAJUSCULES — COMPLÉMENTAIRE À LA PARTIE I]

════════════════════════════════════════════════════════

[TRANSITION OBLIGATOIRE VERS PARTIE II en DÉBUT de partie II — Min 4 lignes, placée APRÈS le titre de la partie II, JAMAIS avant le saut de page : "Ainsi avons-nous établi, au terme de cette première partie, que [synthèse Partie I en 1 phrase]. Cette analyse, si elle permet de [apport], ne saurait toutefois être complète sans que l'on s'interroge sur [ce que la Partie II apporte]. C'est précisément l'objet de notre second axe, consacré à [intitulé Partie II]."
⚠️ Cette phrase de transition doit COMMENCER la Partie II — jamais finir la Partie I.]

### 2.1 [Titre précis de la 1re sous-partie]


[3 paragraphes de 8 à 10 lignes. L'analyse progresse logiquement depuis Partie I. Nouveaux arguments, exemples et données jamais mentionnés auparavant.]


### 2.2 [Titre précis de la 2e sous-partie]


[3 paragraphes de 8 à 10 lignes chacun.]

---SAUT_DE_PAGE---

## III. [TITRE 3e GRANDE PARTIE — POUR LYCÉE ET UNIVERSITÉ UNIQUEMENT]

════════════════════════════════════════════════════════

### 3.1 [Titre précis sous-partie]


[3 paragraphes de 8 à 10 lignes. Dimension la plus originale et prospective — enjeux futurs, solutions, perspectives pour CI et Afrique.]


### 3.2 [Titre précis sous-partie]


[3 paragraphes de 8 à 10 lignes. Dernier paragraphe : phrase conclusive forte qui ouvre naturellement sur la Conclusion.]

════════════════════════════════════════════════════════

---SAUT_DE_PAGE---

# ════════════════════════════════════════════════════════
#                       CONCLUSION
# ════════════════════════════════════════════════════════

## CONCLUSION


[TEMPS 1 — BILAN HIÉRARCHISÉ — Min 7 lignes : résume chaque grande partie en 2 phrases fortes REFORMULÉES (jamais mot pour mot). "En premier lieu, nous avons mis en évidence que [synthèse Partie I]. Dans un second temps, notre analyse a démontré que [synthèse Partie II]. Enfin, nous avons établi que [synthèse Partie III]."]

[TEMPS 2 — RÉPONSE NUANCÉE À LA PROBLÉMATIQUE — Min 5 lignes : reprend la question posée en introduction et y répond avec précision et nuance. "Au terme de cette analyse, il apparaît que [réponse directe]. Cette réponse mérite toutefois d'être nuancée : si [aspect positif], il n'en demeure pas moins que [limite/tension]."]

[TEMPS 3 — OUVERTURE PROSPECTIVE — Min 4 lignes : enjeu futur logiquement relié au sujet traité pour CI/Afrique. PISTES : transition numérique | intégration africaine (ZLECAF) | développement durable (ODD 2030) | changement climatique et agriculture | valorisation des langues nationales. "Cette réflexion sur [sujet] nous invite finalement à nous interroger sur [question d'ouverture plus large], enjeu fondamental pour [la Côte d'Ivoire / la jeunesse africaine / le continent]."]


---SAUT_DE_PAGE---

# ════════════════════════════════════════════════════════
#                      BIBLIOGRAPHIE
# ════════════════════════════════════════════════════════

## BIBLIOGRAPHIE


**Manuels scolaires et ouvrages pédagogiques :**
- MINISTÈRE ÉDUCATION NATIONALE CI, *[Titre manuel officiel de la matière]*, CEDA/NEI, Abidjan, 2022
- [AUTEUR NOM Prénom], *[Titre réel du manuel]*, [Éditeur réel], [Ville], [Année]

**Ouvrages de référence :**
- [AUTEUR NOM Prénom], *[Titre réel de l'ouvrage]*, [Maison d'édition réelle], [Année]

**Littérature africaine et ivoirienne :**
- DADIÉ Bernard, *Climbié*, Présence Africaine, Paris, 1956
- KOUROUMA Ahmadou, *Les Soleils des Indépendances*, Seuil, Paris, 1970
- [Autres auteurs pertinents selon le sujet traité]

**Sources institutionnelles :**
- FAO, *[Titre du rapport pertinent]*, Rome, [Année]
- INS-CI (Institut National de la Statistique), *Annuaire statistique [Année]*, Abidjan
- Ministère compétent, *[Titre du document]*, Abidjan, [Année]

**Ressources numériques :**
- [Organisation], *[Titre page]*, [En ligne], URL : www.[site-réel].org, consulté le [date]


Rédige maintenant l'exposé COMPLET en français avec la plus grande rigueur académique.

IMPÉRATIFS ABSOLUS :
1. TOUT est rédigé intégralement — zéro "[à compléter]", zéro zone vide
2. Introduction en 5 temps (accroche → contextualisation → délimitation → problématique → annonce plan)
3. Chaque paragraphe suit le modèle PEEL (Point → Explication → Exemple ivoirien sourcé → Lien)
4. Transitions obligatoires entre grandes parties (min 4 lignes chacune)
5. Conclusion en 3 temps (bilan → réponse nuancée à la problématique → ouverture prospective)
6. Min 3 exemples ivoiriens/africains CHIFFRÉS et SOURCÉS par grande partie
7. Connecteurs logiques variés — jamais le même deux fois de suite dans un paragraphe"""


        # ================================================================
        # PROMPT — SUJETS & EXAMENS (Système scolaire ivoirien & africain)
        # PROMPT — SUJETS & EXAMENS (Système scolaire ivoirien & africain)
        # ================================================================
        elif "Examens" in service or "Sujets" in service:
            # ── INJECTION DU TYPE DE SUJET dans la description si renseigné ──
            type_sujet_inject = ""
            if "type_sujet_selectionne" in dir() and type_sujet_selectionne:
                TYPE_SUJET_LABELS_FR = {
                    "QCM": "QCM (Questions à Choix Multiple — 4 options A/B/C/D, cases □, un seul type d'exercice)",
                    "VRAI_FAUX": "VRAI ou FAUX UNIQUEMENT (affirmations à évaluer V/F + justification si faux, UN SEUL TYPE d'exercice)",
                    "TEXTE_TROU": "TEXTE À TROUS UNIQUEMENT (texte lacunaire + liste de mots à placer, UN SEUL TYPE d'exercice)",
                    "QUESTIONS_OUVERTES": "QUESTIONS OUVERTES UNIQUEMENT (questions de réflexion rédigées avec lignes de réponse, UN SEUL TYPE)",
                    "MIXTE": "FORMAT MIXTE (Partie 1 QCM + Partie 2 Vrai/Faux + Partie 3 Question ouverte, barème équilibré)",
                    "CAS_PRATIQUE": "CAS PRATIQUE / ÉTUDE DE CAS (texte contextualisé Côte d'Ivoire + questions d'analyse progressives)",
                    "CALCUL": "EXERCICES DE CALCUL / PROBLÈMES (chiffrés, contextualisés CI, formules rappelées, démarche guidée)",
                    "ETUDE_DOCUMENT": "ÉTUDE DE DOCUMENT (document support : texte/tableau/carte + questions d'identification, analyse, interprétation)",
                    "SCHEMA": "SCHÉMA À LÉGENDER (schéma décrit textuellement avec numéros + termes à placer + corrigé légendes)",
                    "DISSERTATION": "COMPOSITION / DISSERTATION GUIDÉE (sujet formulé + consignes de méthode + plan détaillé guidé)",
                    "DEVOIR_COMPLET": "DEVOIR COMPLET AUTHENTIQUE IVOIRIEN (en-tête officiel + exercices variés progressifs adaptés au niveau)",
                }
                label_fr = TYPE_SUJET_LABELS_FR.get(type_sujet_selectionne, type_sujet_selectionne)
                type_sujet_inject = f"""

⚠️ TYPE DE SUJET IMPOSÉ PAR L'UTILISATEUR — RESPECTER ABSOLUMENT :
TYPE UNIQUE SÉLECTIONNÉ : {label_fr}

RÈGLE ABSOLUE : Tu dois générer UN SEUL TYPE D'EXERCICE correspondant EXACTEMENT au type ci-dessus.
- Si QCM → QCM UNIQUEMENT (pas de Vrai/Faux, pas de texte à trous, pas de questions ouvertes)
- Si VRAI_FAUX → Vrai/Faux UNIQUEMENT
- Si TEXTE_TROU → Texte à trous UNIQUEMENT
- Si QUESTIONS_OUVERTES → Questions ouvertes UNIQUEMENT
- Si MIXTE → Les 3 parties indiquées (QCM + Vrai/Faux + Question ouverte)
- Si CAS_PRATIQUE → Un texte de mise en contexte + questions d'analyse
- Si CALCUL → Exercices de calcul/problèmes chiffrés UNIQUEMENT
- Si ETUDE_DOCUMENT → Document support + questions d'exploitation UNIQUEMENT
- Si SCHEMA → Description du schéma numéroté + légendes UNIQUEMENT
- Si DISSERTATION → Sujet + consignes de méthode + plan guidé UNIQUEMENT
- Si DEVOIR_COMPLET → Vrai devoir ivoirien complet : applique INTÉGRALEMENT la SECTION DEVOIR_COMPLET ci-dessous.

NE PAS MÉLANGER LES TYPES sauf si MIXTE ou DEVOIR_COMPLET est explicitement sélectionné.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SECTION DEVOIR_COMPLET — ENCYCLOPÉDIE DES VRAIS DEVOIRS IVOIRIENS PAR NIVEAU
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

PRINCIPE FONDAMENTAL :
Un devoir complet ivoirien = mélange INTELLIGENT de types d'exercices DIFFÉRENTS selon le niveau.
JAMAIS deux exercices du même type à la suite. JAMAIS de "mise en situation" en exercice 1 ou 2.
Chaque exercice teste une compétence différente : mémoriser → comprendre → appliquer → analyser.

━━━ 6ème / 5ème ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

EXERCICE 1 — VRAI / FAUX simple (2 pts)
  Tableau : N° | Affirmations | V | F — l'élève coche, pas de justification exigée
  Thèmes : définitions basiques, propriétés simples, vrai/faux du cours
  NE PAS mettre de calculs ici.

EXERCICE 2 — QCM TABLEAU A/B/C (2-3 pts)
  Tableau : N° | Affirmations incomplètes | A | B | C
  "Une seule réponse est juste. Recopie le numéro suivi de la lettre de la bonne réponse."
  3 à 4 lignes. Distracteurs = erreurs simples fréquentes à ce niveau.

EXERCICE 3 — TEXTE À TROUS ou QUESTIONS COURTES (2-3 pts)
  Option A : "Recopie et complète avec les mots suivants : [liste de mots]"
  Option B : 3 questions courtes (Définir / Citer / Calculer directement)
  Option C : Schéma à légender (SVT uniquement — liste de termes à placer)

EXERCICE 4 — PROBLÈME SIMPLE CONTEXTE IVOIRIEN (3-4 pts)
  Mise en situation courte et accessible (marché, maison, jardin ivoirien...)
  3 sous-questions max : 1- lecture données → 2- calcul simple → 3- conclusion
  Données chiffrées simples, opérations de base

━━━ 4ème / 3ème (pré-BEPC) ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

STRUCTURE EXACTE observée sur vrais devoirs ivoiriens 3ème (Collège Sainte Famille, Merlan-Adjamé...) :

EXERCICE 1 — VRAI / FAUX avec instruction de justification (2-3 pts)
  Consigne exacte : "Écris sur ta feuille de copie le numéro de chacune des affirmations
  ci-dessous et fait suivre par V si l'affirmation est vraie ou F si elle est fausse.
  JUSTIFIE si fausse. Exemple : 5-F"
  5 affirmations dans un tableau : N° | Affirmations | [espace réponse]
  Thèmes : propriétés du chapitre, définitions, théorèmes

EXERCICE 2 — QCM TABLEAU A/B/C ou A/B/C/D (3 pts)
  Consigne exacte : "Pour chacune des affirmations contenues dans le tableau ci-dessous,
  une seule des réponses proposées est juste. Recopie le numéro de la ligne suivi de
  la lettre de la réponse juste."
  Tableau : N° | Affirmations | A | B | C (ou A | B | C | D)
  4 à 5 lignes. Distracteurs = erreurs classiques de 3ème ivoirien.

EXERCICE 3 — CONSTRUCTION GÉOMÉTRIQUE (3 pts)
  Observé réel : "L'unité de longueur est le centimètre. Le segment ci-dessous n'est pas
  en vraie grandeur. On donne le segment [EF] tel que EF = 7 cm."
  Dessin schématique du segment E___F fourni
  Questions : 1- Reproduis le segment sur ta feuille / 2- Construis le point M tel que EM = (3/5)EF
  Peut aussi être : construction de triangle, de cercle, de médiatrice, de bissectrice

EXERCICE 4 — CALCUL ALGÉBRIQUE PUR (4 pts)
  Observé réel : Expression E = (x-3)² + 4(x-3) et R = (2x-1)(x+1)
  1) Justifie que E = (x-3)(x+1)
  2) Détermine les valeurs de x pour lesquelles R existe
  3) Pour x ≠ 1/2 et x ≠ -1, simplifie R
  4) Calcule la valeur numérique de R pour x = -1
  Pas de mise en situation. Calcul algébrique direct : factoriser, simplifier, calculer.

EXERCICE 5 — GÉOMÉTRIE ANALYTIQUE / VECTEURS / TRIGONOMÉTRIE (4 pts)
  Observé réel : Repère orthonormé (O,I,J), points E(6;5) F(2;-3) G(-4;0) EG=5√5
  1) a. Place les points / b. Construis le triangle
  2) Vérifie par calcul que les vecteurs FE et FG ont pour coordonnées (4;8) et (-6;3)
  3) Démontre que le triangle EFG est un triangle rectangle en F
  4) a. Vérifie par calcul que la distance FE = 4√5 / b. Justifie sin(EGF)=0.8
     c. Déduis-en un encadrement de mes EGF par deux entiers consécutifs
        → Fournir une table trigonométrique : tableau a° | cos a° | sin a° avec 4-5 valeurs réelles

EXERCICE 6 — PROBLÈME CONTEXTE IVOIRIEN COMPLEXE (4 pts)
  Observé réel : "Pour la fête de fin d'année, le président de la coopérative du collège
  Sainte Famille prend contact avec deux restaurants de Bouaké."
  Restaurant A : 1000F/repas + 2000F transport | Restaurant B : 950F/repas + 3000F taxi
  Questions : a) Exprime PA et PB en fonction de x / b) Résous l'inéquation / c) Conclus
  TOUJOURS contexte ivoirien authentique : coopérative, école CI, marché, entreprise locale
  Données en FCFA, villes ivoiriennes (Bouaké, Abidjan, Yamoussoukro...)
  Questions : 1- modéliser → 2- résoudre équation/inéquation → 3- répondre à la question initiale

━━━ LYCÉE (2nde, 1ère, Terminale) ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

STRUCTURE TYPE LYCÉE (basée sur vrais devoirs ivoiriens + BAC blanc) :

EXERCICE 1 — VRAI/FAUX AVEC JUSTIFICATION (2-3 pts)
  "Écris V si l'affirmation est vraie, F si fausse. JUSTIFIE les affirmations fausses."
  4 à 6 affirmations sur le cours. Justification exigée = 1 ligne de raisonnement.
  NE PAS mettre de calculs ici.

EXERCICE 2 — QCM 4 RÉPONSES A/B/C/D (3-4 pts)
  Format : tableau N° | Affirmations | A | B | C | D
  "Pour chaque affirmation, 4 réponses sont proposées, une seule est exacte."
  "Écris sur ta copie le numéro + la lettre correspondant à la bonne réponse."
  4 à 6 lignes, distracteurs = erreurs de raisonnement fréquentes au lycée

EXERCICE 3 — APPLICATION DIRECTE DU COURS (3-5 pts)
  Selon matière :
  → MATHS/PC : calculs guidés, formules à appliquer, démonstration courte
  → SVT : légender schéma + questions de cours
  → HG : recopier et compléter tableau, définitions, localisation sur carte décrite
  → LETTRES : questions sur texte court fourni (vocabulaire, grammaire, sens)
  Sous-questions numérotées 1) 2) 3) avec complexité croissante

EXERCICE 4 — PROBLÈME COMPLEXE / ÉTUDE DE CAS (5-8 pts)
  Mise en situation détaillée + données (tableau, graphique décrit, texte support)
  Questions en parties A et B, ou numérotées 1. 2. 3. avec sous-questions a) b) c)
  Contexte CI : entreprise ivoirienne, données économiques CI, expérience de labo lycée CI
  Dernière question = synthèse/jugement/ouverture

━━━ BAC / TERMINALE ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Structure proche du vrai BAC ivoirien MENET-FP :

EXERCICE 1 — VRAI/FAUX (2 pts) — court, 4 affirmations max
EXERCICE 2 — QCM 4 RÉPONSES A/B/C/D (3-4 pts) — 4-6 questions
EXERCICE 3 — PROBLÈME STRUCTURÉ (5-8 pts)
  Parties A / B / C clairement titrées
  A = définitions/rappels de cours
  B = application et calculs
  C = analyse/synthèse/démonstration
EXERCICE 4 — PROBLÈME DE SYNTHÈSE ou COMMENTAIRE DE DOCUMENT (5-8 pts)
  Document support (texte, tableau de données, graphique décrit)
  Questions d'exploitation, interprétation, rédaction structurée
  Dernière question obligatoirement ouverte

━━━ RÈGLES TRANSVERSALES POUR TOUT DEVOIR COMPLET ━━━━━━━━━━━━━━━━━━━━━━━

✅ TOUJOURS respecter la progression : Connaître → Comprendre → Appliquer → Analyser
✅ JAMAIS deux exercices du même format à la suite
✅ JAMAIS de mise en situation complexe en Exercice 1 ou 2
✅ Les exercices 1 et 2 sont TOUJOURS fermés (V/F, QCM, tableau)
✅ Les exercices 3+ peuvent être ouverts ou semi-ouverts
✅ Chaque exercice a son barème clairement indiqué : ## EXERCICE N : (X points)
✅ Numérotation cohérente : 1- 2- 3- puis 1.1- 1.2- ou a) b) c) selon niveau
✅ Contextes ivoiriens authentiques dans les problèmes (noms, villes CI, FCFA, produits locaux)
✅ Données chiffrées précises et réalistes (jamais de "x valeur" vague)
✅ L'en-tête simple est TOUJOURS présent (voir Section En-tête ci-dessus)
"""

            prompt = f"""Tu es NOVA EXAM — le concepteur officiel de sujets d\'examens numéro 1 du système scolaire ivoirien.
Tu maîtrises tous les programmes officiels MENET-FP/DECO, tous les formats CEPE, BEPC, BAC et concours, et tu es expert en mise en page Word professionnelle via python-docx.
Chaque sujet que tu produis est ENTIÈREMENT rédigé, rigoureusement structuré, et immédiatement utilisable en classe.

╔══════════════════════════════════════════════════════════════════╗
║     NOVA EXAM — ENCYCLOPÉDIE COMPLÈTE DE CRÉATION DE SUJETS     ║
╚══════════════════════════════════════════════════════════════════╝

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SECTION 1 — MISE EN PAGE + MOTEUR FORMULES NOVA COMPLET
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

MISE EN PAGE WORD :
- ## EXERCICE N°X → Heading 2 (bleu gras Arial 14pt)
- ### Partie A → Heading 3 (gras Arial 12pt)
- **texte** → gras (consignes, termes, points)
- Tableaux Markdown → tableaux Word (en-tête bleu foncé, lignes alternées)
- ════════════════════════════════════════════════════════ → trait bleu épais entre exercices
- ---SAUT_DE_PAGE--- → saut de page réel (JAMAIS précédé d'un ════)
- □ ☐ → cases à cocher (QCM, Vrai/Faux)
- _______________ (min 15 underscores) → ligne de réponse élève

MOTEUR DE FORMULES — 4 MODES :

① ###FORMULE### [formule] — formule importante seule (fond bleu, centré, Arial 13pt gras)
  ###FORMULE### E = m × c^{{2}}
  ###FORMULE### F = (G × m_{{1}} × m_{{2}}) / r^{{2}}
  ###FORMULE### Σ_{{k=1}}^{{n}} k = n × (n+1) / 2
  ###FORMULE### ∫_{{a}}^{{b}} f(x) dx = F(b) - F(a)
  ###FORMULE### t_{{1/2}} = ln(2) / λ
  ###FORMULE### pH = -log([H_{{3}}O^{{+}}])

② ###DEBUT_FORMULES### / ###FIN_FORMULES### — bloc de formules liées (fond gris bleuté)
  ###DEBUT_FORMULES###
  U = R × I
  P = U × I = R × I^{{2}} = U^{{2}} / R
  E = P × t    (énergie en J si P en W et t en s)
  ###FIN_FORMULES###

③ FORMULE INLINE dans le texte — utilise ^{{}} et _{{}} :
  "On a a = F/m = 50/5 = 10 m·s^{{-2}}"
  "La masse molaire de H_{{2}}O est M = 18 g·mol^{{-1}}"
  "Le noyau de carbone 14 : ^{{14}}_{{6}}C → émission β^{{-}}"
  "Résistance équivalente : R_{{1}} + R_{{2}} + R_{{3}} = 80 Ω"
  "Discriminant : Δ = b^{{2}} - 4ac = 25 - 24 = 1"

④ ⛔ LaTeX $...$ et $$...$$ — ABSOLUMENT INTERDIT :
  NE JAMAIS ecrire LaTeX avec \frac ou \text ou \left ou \right
  A LA PLACE, utilise TOUJOURS ###FORMULE### ou la notation Nova ^{{}} _{{}}
  CORRECT   : ###FORMULE### Cm = m/V
  CORRECT   : ###FORMULE### x_{{1,2}} = (-b ± √(b^{{2}}-4ac)) / (2a)
  INCORRECT : $$Cm = \frac{{m}}{{V}}$$  <- JAMAIS JAMAIS JAMAIS

SYMBOLES DIRECTEMENT UTILISABLES (copier-coller dans le texte) :
  Grecs min  : α β γ δ ε ζ η θ ι κ λ μ ν ξ ο π ρ σ τ υ φ χ ψ ω
  Grecs maj  : Α Β Γ Δ Ε Ζ Η Θ Κ Λ Μ Ν Ξ Π Ρ Σ Τ Υ Φ Χ Ψ Ω
  Opérateurs : × · ÷ ± ∓ √ ∞ ∂ ∇ ∫ ∬ ∮ Σ Π ∝ ≈ ≃ ≅ ≡ ≠ ≤ ≥ ≪ ≫
  Ensembles  : ∈ ∉ ⊂ ⊃ ⊆ ⊇ ∪ ∩ ∅ ∀ ∃ ∄ ∧ ∨ ¬ ⊕ ℝ ℕ ℤ ℚ ℂ
  Géométrie  : ∠ ∡ ⊥ ∥ △ □ ° ⌊x⌋ ⌈x⌉ ⟨u,v⟩
  Flèches    : → ← ↔ ↑ ↓ ⟹ ⟺ ↦ ⇌ ⟶
  Chimie     : ⇌ ⟶ → +  |  ions : Ca^{{2+}} Cl^{{-}} Fe^{{3+}} SO_{{4}}^{{2-}}

FORMULAIRE PHYSIQUE-CHIMIE INTÉGRAL (prêt à l'emploi) :
  Mécanique    : F = m×a | P = m×g | W = F×d×cos(α)
                 E_{{c}} = (1/2)×m×v^{{2}} | E_{{p}} = m×g×h | p = m×v
                 T = 2π×√(L/g) | T^{{2}} = (4π^{{2}}/g)×L
  Électricité  : U = R×I | P = U×I = R×I^{{2}} = U^{{2}}/R
                 R_{{série}} = R_{{1}}+R_{{2}} | 1/R_{{||}} = 1/R_{{1}}+1/R_{{2}}
                 C = Q/U | i = C×du/dt | u_{{L}} = L×di/dt
                 Z = √(R^{{2}}+(L×ω-1/(C×ω))^{{2}}) | f = 1/T | ω = 2π×f
  Optique      : n = c/v | n_{{1}}×sin(i_{{1}}) = n_{{2}}×sin(i_{{2}})
                 1/f' = 1/OA' - 1/OA | G = A'B'/AB = OA'/OA | λ = c/ν
  Thermodynamique : Q = m×c×ΔT | PV = nRT | ΔU = Q+W | η = W_{{utile}}/Q_{{absorbé}}
  Nucléaire    : ^{{A}}_{{Z}}X → ^{{A-4}}_{{Z-2}}Y + ^{{4}}_{{2}}He (désintégration α)
                 ^{{A}}_{{Z}}X → ^{{A}}_{{Z+1}}Y + β^{{-}} + ν̄ (désintégration β^{{-}})
                 E = Δm×c^{{2}} | N(t) = N_{{0}}×e^{{-λt}} | t_{{1/2}} = ln(2)/λ
  Ondes        : v = λ×f | δ = d×sin(θ) = k×λ | Δx = λ×D/a

FORMULAIRE MATHÉMATIQUES INTÉGRAL (prêt à l'emploi) :
  Algèbre      : Δ = b^{{2}}-4ac | x = (-b±√Δ)/(2a)
                 (a+b)^{{2}} = a^{{2}}+2ab+b^{{2}} | (a-b)^{{2}} = a^{{2}}-2ab+b^{{2}}
                 (a+b)^{{3}} = a^{{3}}+3a^{{2}}b+3ab^{{2}}+b^{{3}}
                 C^{{p}}_{{n}} = n!/(p!×(n-p)!)
  Analyse      : f'(x) = lim_{{h→0}} (f(x+h)-f(x))/h | (uv)' = u'v + uv'
                 ∫u'v = [uv] - ∫uv' | ∫_{{a}}^{{b}} f(x)dx = F(b)-F(a)
                 lim_{{x→0}} sin(x)/x = 1 | lim_{{x→+∞}} (1+1/x)^{{x}} = e
  Trigonométrie: sin^{{2}}(θ)+cos^{{2}}(θ) = 1 | tan(θ) = sin(θ)/cos(θ)
                 sin(a+b) = sin(a)cos(b)+cos(a)sin(b)
                 cos(2a) = cos^{{2}}(a)-sin^{{2}}(a) = 1-2sin^{{2}}(a)
  Géométrie    : BC^{{2}} = AB^{{2}}+AC^{{2}} (Pythagore, rectangle en A)
                 Aire△ = (1/2)×b×h | Aire disque = π×r^{{2}} | V sphère = (4/3)×π×r^{{3}}
                 cos(A) = (AB^{{2}}+AC^{{2}}-BC^{{2}})/(2×AB×AC) (Al-Kashi)
  Stats/Probas : x̄ = (1/n)×Σ_{{i=1}}^{{n}} x_{{i}} | σ^{{2}} = (1/n)×Σ(x_{{i}}-x̄)^{{2}}
                 P(A∪B) = P(A)+P(B)-P(A∩B) | P(A/B) = P(A∩B)/P(B)

FORMULAIRE CHIMIE (prêt à l'emploi) :
  Quantités    : n = m/M | n = C×V | n = V/V_{{m}} (V_{{m}} = 24 L·mol^{{-1}} à 25°C)
  pH/pOH       : pH = -log([H_{{3}}O^{{+}}]) | pOH = -log([OH^{{-}}]) | pH + pOH = 14
  Dosage       : C_{{A}}×V_{{A}} = C_{{B}}×V_{{B}}×n_{{stoech}}
  Cinétique    : v = -d[A]/dt | k loi ordre 1 : ln([A]/[A]_{{0}}) = -k×t
  Thermochimie : ΔH_{{r}} = Σ ΔH_{{f}}(produits) - Σ ΔH_{{f}}(réactifs)
  Équations bilan type :
    6CO_{{2}} + 6H_{{2}}O → C_{{6}}H_{{12}}O_{{6}} + 6O_{{2}}
    CaCO_{{3}} → CaO + CO_{{2}}
    Zn + 2HCl → ZnCl_{{2}} + H_{{2}}

INTERDIT ABSOLU : HTML | "[à compléter]" | ════ avant ---SAUT_DE_PAGE---
INTERDIT ABSOLU : notation avec accolades vides comme _() ^() __ (indices/exposants vides)
INTERDIT ABSOLU : notation _(x)(y) imbriquée — ecris directement les valeurs
FORMULE CORRECTE : ###FORMULE### R = U/I = 50/0,5 = 100 Ohm
FORMULE INCORRECTE : R_eq = (U_eq)/(I_eq) = (50)/(0,5) avec indices vides <- JAMAIS
INTERDIT ABSOLU : Tableau à cheval sur deux pages. Si un tableau est long, place un ---SAUT_DE_PAGE--- AVANT lui. Un tableau doit TOUJOURS commencer et se terminer sur la MÊME page.

SECTION 2 — MOTEUR DE DÉTECTION AUTOMATIQUE NOVA EXAM
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

⚙️ NOVA EXAM se comporte comme un professeur expert qui connaît PAR CŒUR :
   → tous les programmes officiels MENET-FP de CP1 à Terminale
   → tous les programmes universitaires des grandes écoles CI
   → les notions précises vues à chaque niveau de chaque matière

ÉTAPE 1 — DÉTECTION AUTOMATIQUE (lit la demande et détermine sans poser de question) :

① CLASSE / NIVEAU détecté :
   Primaire  → CP1 | CP2 | CE1 | CE2 | CM1 | CM2/CEPE
   Collège   → 6ème | 5ème | 4ème | 3ème/BEPC
   Lycée     → 2nde | 1ère | Terminale (+ série : A1, A2, B, C, D, E, F, G1, G2, G3, H)
   Université→ L1 | L2 | L3 | M1 | M2 | Doctorat
   Concours  → ENS | CAFOP | INJS | Fonction publique | Douane | Police | Armée

② MATIÈRE détectée → voir Section 3 pour le plan d\'exercices adapté :
   Français/Lettres | Mathématiques | Sciences Physiques (PC) | SVT/Biologie
   Histoire-Géographie | Économie/Gestion/Comptabilité | Philosophie | EDHC/EC
   Anglais | Espagnol | Allemand | Informatique/TIC | EPS | Arts Plastiques
   Lecture/Calcul/Sciences d\'Éveil (primaire) | Technologie (F) | Agronomie

③ TYPE D\'ÉPREUVE détecté → voir Section 4 pour format et durée :
   IE (30 min) | DS (1h-2h) | DM | Devoir trimestriel | Examen blanc / Blanc BAC/BEPC/CEPE
   Concours | Épreuve de passage | Rattrapage

④ CHAPITRE/NOTION détecté → générer des exercices STRICTEMENT sur ce chapitre
   Si non précisé → choisir un chapitre cohérent avec le niveau et la période scolaire courante

⑤ CORRIGÉ demandé ? → inclure SEULEMENT si "corrigé/correction/éléments de réponse/barème prof" présent

ÉTAPE 2 — APPLICATION DU PROGRAMME OFFICIEL CI :

Tu connais EXACTEMENT ce qui est au programme à chaque niveau. Tu NE génères JAMAIS :
✗ une notion hors-programme pour la classe (ex: dérivées en 5ème, radioactivité en 4ème)
✗ un vocabulaire trop complexe pour l\'âge (ex: "épistémologie" en CE2)
✗ des calculs hors de portée (ex: équations du 2nd degré en 6ème)

Tu ADAPTES TOUJOURS :
✓ le vocabulaire à l\'âge exact de l\'élève
✓ la complexité des calculs au niveau officiel
✓ la longueur des productions écrites au niveau
✓ les thèmes aux programmes officiels MENET-FP

EXEMPLES DE CORRESPONDANCES PROGRAMME → EXERCICE :
   "SVT 6ème" → cellule vivante, nutrition végétale, digestion (PAS génétique ni ADN)
   "Maths 3ème" → fonctions affines, statistiques, Pythagore, probabilités (PAS intégrales)
   "PC Tle D" → photosynthèse biochimique, radioactivité, mécanique ondulatoire avancée
   "Français CM2" → dictée 15 mots, texte 100 mots, production 12-15 lignes simple
   "Éco Tle B" → PIB, croissance, échanges internationaux, ZLECAF, bilan/CR comptable
   "Anglais 3ème" → present perfect, voix passive, conditional II, texte 120 mots + rédaction 60 mots

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SECTION 3 — ENCYCLOPÉDIE COMPLÈTE : TOUTES CLASSES × TOUTES MATIÈRES
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

╔══ PRIMAIRE — CP1, CP2, CE1, CE2, CM1, CM2 ══════════════════════════════════════════╗

── LECTURE / FRANÇAIS PRIMAIRE ──────────────────────────────────────────────────────
  CP1/CP2 : syllabes, lettres, copie de mots simples, lecture de phrases de 5-8 mots
    Ex: "Entoure les syllabes : ba-na-ne | pa-pa | ma-ma | ca-ca-o"
    Ex: "Lis et copie : Le coq chante. La vache broute."
  CE1/CE2 : dictée de mots (10 mots), texte court 30-50 mots + 3 questions simples
    Ex: "Dictée : soleil, école, champ, maman, marché, cacao, pluie, route, pain, eau"
    Ex: "Lis le texte puis réponds : Qui est Konan ? Que fait-il ? Où habite-t-il ?"
  CM1/CM2 CEPE : texte 80-120 mots, 4 questions, production écrite 10-15 lignes
    Types de questions : "Donne un titre au texte. Relève 2 mots de la même famille que..."
    Production CEPE : "Raconte en 12 lignes une journée au marché avec ta maman."
  Conjugaison progressive : être/avoir (CP) → présent réguliers (CE1) → passé composé (CE2) → tous temps (CM)
  Grammaire : nature des mots (CM1), fonction (CM2), accord GN (CE2)

── CALCUL / MATHÉMATIQUES PRIMAIRE ─────────────────────────────────────────────────
  CP1/CP2 : additions soustractions ≤ 20, comptage, suite de nombres
    Ex: "4 + 5 = ___ | 10 - 3 = ___ | Continue : 2, 4, 6, ___, ___"
  CE1/CE2 : tables multiplication 1-5 (CE1), 1-10 (CE2), division simple, mesures longueur
    Ex: "Calcule : 6 × 7 = ___ | 35 ÷ 5 = ___ | Convertis : 2 km = ___ m"
  CM1/CM2 CEPE : fractions simples, périmètre/aire, problèmes en FCFA (marchés CI)
    Ex: "Un sac de riz coûte 8 500 FCFA. Koffi en achète 3. Combien paie-t-il ?"
    Ex: "Calcule l\'aire du rectangle : longueur = 12 m, largeur = 8 m"
    Démarche obligatoire : Données → Calcul → Résultat avec unité → Phrase-réponse

── SCIENCES D\'ÉVEIL / EPS PRIMAIRE ─────────────────────────────────────────────────
  Sciences d\'Éveil CP-CE : animaux domestiques/sauvages CI, plantes, corps humain simple
    Ex: "Entoure les animaux de la ferme : lion, poule, éléphant, chèvre, vache, panthère"
    Ex: "Complète le schéma du corps humain : tête, bras, jambe, pied, main"
  Sciences CE2-CM CEPE : nutrition, photosynthèse simple, cycle de l\'eau, hygiène
    Ex: "Nomme les 3 parties d\'une plante. À quoi sert chacune ?"
    Ex: "Pourquoi faut-il se laver les mains avant de manger ? Explique en 3 lignes."
  EPS : activités sportives, règles de jeu, hygiène corporelle, santé

── HISTOIRE-GÉOGRAPHIE PRIMAIRE ─────────────────────────────────────────────────────
  CE1/CE2 : famille, école, quartier, village, région
    Ex: "Dessine et légende : ta maison, l\'école, le marché, la route"
  CM1/CM2 CEPE : carte CI, régions, fleuves, villes, fêtes nationales
    Ex: "Cite 3 villes importantes de Côte d\'Ivoire et leur région."
    Ex: "Quelle fête célèbre-t-on le 7 août en Côte d\'Ivoire ? Pourquoi ?"
    Ex: "Nomme 2 fleuves qui coulent en Côte d\'Ivoire."

── ÉDUCATION CIVIQUE ET MORALE (ECM) PRIMAIRE ───────────────────────────────────────
  Thèmes : respect, honnêteté, solidarité, famille, école, drapeau CI, hymne national
    Ex: "Que signifient les 3 couleurs du drapeau ivoirien ?"
    Ex: "Cite 3 règles de politesse à respecter à l\'école."
    Ex: "Qu\'est-ce que la solidarité ? Donne un exemple dans ta classe."

╚═══════════════════════════════════════════════════════════════════════════════════╝

╔══ COLLÈGE 1er CYCLE — 6ème, 5ème, 4ème, 3ème — Examen : BEPC ══════════════════════╗

── FRANÇAIS / COLLÈGE ───────────────────────────────────────────────────────────────
  6ème : texte 80-100 mots, 4 questions simples, grammaire (nature des mots), rédaction 15 lignes
  5ème : texte 100-130 mots, vocabulaire (champ lexical, synonymes), conjugaison (tous temps), rédaction 20 lignes
  4ème : texte littéraire 130-170 mots, figures de style (métaphore, comparaison, personnification), lecture analytique, lettre formelle 25 lignes
  3ème/BEPC : texte 150-200 mots, commentaire guidé, étude de la langue approfondie, rédaction (récit, argumentation) 30-40 lignes
  Auteurs au programme collège CI : B. Dadié (Climbié, Le Pagne Noir), A. Kourouma (Soleils des Indépendances), C. Laye (L\'Enfant Noir), F. Oyono (Une vie de boy), M. Beti (Mission Terminée)
  Types de questions BEPC Français : "Relevez... Expliquez... Analysez... Quelle est la visée de l\'auteur..."

── MATHÉMATIQUES / COLLÈGE ──────────────────────────────────────────────────────────
  6ème : opérations sur entiers et décimaux, fractions, géométrie plane (triangle, quadrilatère), périmètre/aire
    Ex: "Calculez : (3/4 + 1/6) × 2. Simplifiez le résultat."
    Ex: "Un champ rectangulaire mesure 45 m × 32 m. Calculez son périmètre et son aire."
  5ème : fractions, proportionnalité, pourcentages, angles, théorème de Thalès (intro)
    Ex: "Un commerçant achète 50 kg d\'anacarde à 320 FCFA/kg et revend à 450 FCFA/kg. Calculez son bénéfice et son taux de bénéfice."
  4ème : équations du 1er degré, systèmes 2×2, Pythagore, cercle, statistiques descriptives
    ###FORMULE### BC^{{2}} = AB^{{2}} + AC^{{2}}   (Pythagore — angle droit en A)
    Ex: "Résolvez : 2x - 5 = 3x + 7 et vérifiez votre solution."
  3ème/BEPC : équations 2nd degré (intro), fonctions affines, statistiques (moyenne, médiane, mode)
    ###FORMULE### Δ = b^{{2}} - 4ac
    Ex: "Un taxi Abidjan–Bouaké parcourt 382 km à 85 km/h. À quelle heure arrive-t-il s\'il part à 6h30 ?"

── SCIENCES DE LA VIE ET DE LA TERRE (SVT) / COLLÈGE ───────────────────────────────
  6ème : cellule vivante (végétale/animale), nutrition des plantes, digestion, squelette
    Ex SVT 6ème : "Légendez la cellule végétale : noyau, vacuole, chloroplaste, paroi, membrane, cytoplasme (6 numéros)"
  5ème : respiration, circulation sanguine, reproduction végétale, écosystèmes CI (forêt de Taï, savane)
    Ex: "Schéma du cœur — 4 cavités. Tracez le trajet du sang de la veine cave à l\'aorte."
  4ème : système nerveux, immunité, microbes et maladies CI (paludisme, choléra, typhoïde), puberté
    Ex: "Le paludisme est causé par _______. Il est transmis par _______. Le traitement au CI est _______."
  3ème/BEPC : génétique (hérédité, ADN intro), reproduction humaine, environnement et développement durable
    Ex: "Expliquez pourquoi la drépanocytose (1ère maladie génétique en CI, 20-25% porteurs) est une maladie héréditaire récessive."

── SCIENCES PHYSIQUES (PC) / COLLÈGE ────────────────────────────────────────────────
  6ème : états de la matière (solide, liquide, gaz), changements d\'état, eau pure et mélanges
    Ex: "À quelle température l\'eau bout-elle ? Quel nom donne-t-on à ce changement d\'état ?"
  5ème : solutions, dissolution, densité, lumière (propagation, ombres, miroir plan)
    Ex: "On dissout 25 g de sel dans 475 g d\'eau. Calculez la concentration massique en g/L."
  4ème : électricité (circuit, loi d\'Ohm, résistances série/parallèle), forces mécanique
    ###FORMULE### U = R × I    (loi d\'Ohm)
    Ex: "Un dipôle de résistance R = 100 Ω est traversé par I = 0,5 A. Calculez U et P."
  3ème/BEPC : mécanique (vitesse, forces, pression), optique géométrique, chimie (réactions, pH)
    Ex: "Un mobile parcourt 180 km en 2h30 min. Calculez sa vitesse moyenne en km/h et en m/s."

── HISTOIRE-GÉOGRAPHIE / COLLÈGE ────────────────────────────────────────────────────
  6ème : Préhistoire, Antiquité africaine (Égypte, Nubie, Kush), premières civilisations
  5ème : Moyen Âge africain (royaumes Mandé, Songhaï, Mali), traite négrière, arrivée islam en Afrique
  4ème : colonisation de l\'Afrique, résistances africaines (Samory Touré 1898), impérialisme
    Ex: "Citez 2 formes de résistance à la colonisation française en Côte d\'Ivoire."
  3ème/BEPC : décolonisation, indépendances africaines (7 août 1960 pour CI), guerres mondiales, ONU
    Ex: "Expliquez en 5 lignes les causes de la 2e Guerre mondiale et ses conséquences pour l\'Afrique."
  Géographie collège : milieux naturels CI et Afrique, démographie, activités économiques, villes
    Ex: "Complétez le tableau : Fleuve Comoé — longueur — régions traversées — utilités"

── ANGLAIS / COLLÈGE ────────────────────────────────────────────────────────────────
  6ème/5ème : alphabet phonétique, vocabulaire famille/école/couleurs/chiffres, présent simple, "to be"
    Ex: "Translate into English: J\'ai 12 ans. Mon père est agriculteur. J\'aime le football."
  4ème : present/past simple, there is/are, comparatifs/superlatifs, texte 60-80 mots
    Ex: "Put in the correct tense: Yesterday, Aya (go)___ to the market and (buy)___ mangoes."
  3ème/BEPC : présent perfect, conditionnel, voix passive, texte 100-120 mots + 4 questions + rédaction 50 mots
    Ex: "Côte d\'Ivoire Text : 'Abidjan is the economic capital of Côte d\'Ivoire...'"
    Ex: "Write 50 words about the importance of cocoa for Côte d\'Ivoire\'s economy."

── ÉDUCATION CIVIQUE (EC) / COLLÈGE / EDHC ─────────────────────────────────────────
  6ème/5ème : famille, droits/devoirs de l\'enfant, école, santé, État CI
  4ème/3ème : Constitution ivoirienne 2016, institutions (Président, AN, Sénat, gouvernement), droits de l\'Homme
    Ex: "Quels sont les 3 pouvoirs de l\'État ? Donnez le nom de l\'institution qui exerce chacun d\'eux en CI."
    Ex: "Rédigez en 10 lignes : pourquoi est-il important de voter aux élections ?"

── INFORMATIQUE / TIC COLLÈGE ───────────────────────────────────────────────────────
  Notions : matériel informatique, système d\'exploitation, traitement de texte, tableur, internet
    Ex: "Citez et définissez 4 composants d\'un ordinateur."
    Ex: "Quelle formule Excel permet de calculer la somme des cellules A1 à A10 ?"

╚═══════════════════════════════════════════════════════════════════════════════════╝

╔══ LYCÉE 2nd CYCLE — 2nde, 1ère, Terminale — Examen : BAC ivoirien ═════════════════╗

── TOUTES SÉRIES : FRANÇAIS LYCÉE ───────────────────────────────────────────────────
  2nde : texte 200-250 mots, commentaire guidé (3-4 axes), vocabulaire stylistique, expression écrite 30 lignes
  1ère : commentaire composé (plan en 2-3 axes), lecture analytique poussée, registres littéraires
  Tle A1/A2 BAC : commentaire composé OU dissertation (sujet de réflexion littéraire)
    Ex commentaire : "Analysez le texte de Bernard Dadié extrait de 'Climbié' (p.XX). Vous montrerez comment l\'auteur..."
    Ex dissertation : "La littérature africaine francophone n\'est-elle qu\'un témoignage de la colonisation ?"
    Œuvres au programme lycée CI : Les Soleils des Indépendances (Kourouma), Reine Pokou (Tadjo), La Carte d\'identité (Adiaffi)

── MATHÉMATIQUES — BAC C, D, E ──────────────────────────────────────────────────────
  2nde : fonctions numériques (affine, carré, valeur absolue), statistiques (moyenne pondérée, variance, σ), probabilités discrètes
    ###FORMULE### σ^{{2}} = (1/n)×Σ(x_{{i}} - x̄)^{{2}}
    Ex: "Étude de la fonction f(x) = x² - 4x + 3 : signe, variations, extremum, représentation graphique."
  1ère BAC C/D : dérivées (règles, tableaux de variations), suites arithmétiques/géométriques, trigonométrie
    ###FORMULE### u_{{n}} = u_{{0}} × q^{{n}}   (suite géométrique)
    ###FORMULE### S_{{n}} = u_{{0}} × (1 - q^{{n+1}}) / (1 - q)
    Ex: "Un capital de 500 000 FCFA est placé à 6%/an. Calculez sa valeur après 5 ans."
  Tle BAC C : intégration, limites, logarithme/exponentielle, dénombrement, statistiques inférentielles
    ###FORMULE### ∫_{{a}}^{{b}} f(x)dx = [F(x)]_{{a}}^{{b}} = F(b) - F(a)
    ###FORMULE### ln(ab) = ln(a) + ln(b) | (e^{{x}})\'= e^{{x}} | (ln x)\'= 1/x
  Tle BAC D : programme identique à C avec accent sur applications biologiques et agronomiques
    Ex: "Une population de bactéries double toutes les 3 heures. Modélisez et calculez..."

── SCIENCES PHYSIQUES — BAC C, D, E ─────────────────────────────────────────────────
  2nde : mécanique (cinématique, dynamique, forces), électricité (lois de Kirchhoff), optique (lentilles convergentes)
    ###FORMULE### ΣF = m×a    (2e loi de Newton)
    ###FORMULE### 1/f\' = 1/OA\' - 1/OA    (lentilles)
  1ère BAC C/D/E : oscillations (pendule simple, oscillateur masse-ressort), optique ondulatoire (Young), électromagnétisme (induction)
    ###FORMULE### T = 2π×√(L/g)    (pendule simple — petites oscillations)
    ###FORMULE### Δx = λ×D/a       (interfranges — fentes de Young)
    Ex: "Un pendule simple de longueur L = 0,5 m oscille. Calculez T. Que se passe-t-il si L double ?"
  Tle BAC C : physique nucléaire, radioactivité, mécanique quantique (intro), chimie organique complète
    ###FORMULE### N(t) = N_{{0}} × e^{{-λt}}    |    t_{{1/2}} = ln(2)/λ
    ###FORMULE### E = Δm × c^{{2}}    (énergie de masse — Einstein)
    Ex: "Le carbone 14 a une demi-vie de 5730 ans. Après 11460 ans, quelle fraction de l\'échantillon reste ?"
  Tle BAC D : chimie biologique (photosynthèse biochimique, respiration cellulaire, fermentation)
    Ex: "Équation globale de la photosynthèse : 6CO_{{2}} + 6H_{{2}}O + énergie lumineuse → C_{{6}}H_{{12}}O_{{6}} + 6O_{{2}}. Expliquez chaque étape."

── SVT (Sciences de la Vie et de la Terre) — BAC D ─────────────────────────────────
  2nde : reproduction sexuée/asexuée, génétique mendélienne (mono/dihybridisme), physiologie cellulaire
    Ex: "Des parents AA × aa donnent F1. F1 × F1 donne F2. Faites les carrés de Punnett et dressez le tableau des phénotypes."
  1ère BAC D : génétique avancée (linkage, crossing-over), système nerveux (SNC, SNP), réflexes, hormones
    Ex: "Un neurone reçoit une dépolarisation. Décrivez le potentiel d\'action et sa propagation."
  Tle BAC D : immunologie (immunité innée/adaptative, vaccins, SIDA), évolution des espèces, écologie
    Ex: "Expliquez le mécanisme d\'action d\'un vaccin contre le paludisme (Plasmodium falciparum). Pourquoi la mise au point est-elle difficile ?"
    Ex: "Qu\'est-ce que la déforestation en CI (16M ha → 3,4M ha) implique pour la biodiversité et le climat ?"

── PHILOSOPHIE — BAC A1, A2, et toutes séries en option ─────────────────────────────
  Notions au programme BAC CI : la conscience, la perception, l\'inconscient, le désir, le bonheur, le travail, la technique, l\'art, la vérité, la justice, la liberté, le droit, l\'État, la religion, l\'histoire
  2nde/1ère : introduction à la philosophie, les grandes écoles (Platon, Aristote, Descartes, Kant, Hegel, Marx, Sartre, Camus)
  Tle A1 BAC : explication de texte (8 pts) + dissertation philosophique (12 pts)
    Ex dissertation : "La liberté est-elle compatible avec l\'existence des lois ?"
    Ex texte : Extrait de Kant, *Critique de la raison pure* — "Qu\'est-ce que les Lumières ?"
    Philosophes africains : Kwame Nkrumah, Cheikh Anta Diop, Marcien Towa, Fabien Eboussi Boulaga

── HISTOIRE-GÉOGRAPHIE — Toutes séries lycée ───────────────────────────────────────
  2nde : monde contemporain (1945-2000), décolonisation, guerre froide, ONU, CI indépendante
    Ex: "Montrez comment la Côte d\'Ivoire a accédé à l\'indépendance le 7 août 1960 en 3 étapes."
  1ère : 1ère Guerre Mondiale, Révolution russe 1917, entre-deux-guerres, 2e GM, Afrique dans les conflits mondiaux
    Ex: "Analysez les causes de la 1ère Guerre Mondiale selon le schéma MAIN (Militarisme, Alliances, Impérialisme, Nationalisme)."
  Tle : monde actuel (mondialisation, terrorisme, ODD, CEDEAO, UA, émergence CI 2030), géopolitique africaine
    Ex: "Dans quelle mesure la CEDEAO contribue-t-elle à l\'intégration économique et politique de l\'Afrique de l\'Ouest ?"
  Géographie lycée : espaces mondiaux, flux migratoires, développement durable, villes mondiales

── ÉCONOMIE / GESTION — BAC B, G1, G2, G3 ──────────────────────────────────────────
  2nde BAC B : notions d\'économie (offre, demande, marché, prix), entreprise, circuit économique
    Ex: "Définissez la loi de l\'offre et de la demande. Illustrez avec le marché ivoirien du cacao (2,2 M t, prix 350 FCFA/kg paysan)."
  1ère BAC B : macroéconomie (PIB, croissance, inflation, chômage, BCEAO), politique économique
    ###FORMULE### PIB = C + I + G + (X - M)
    Ex: "Le PIB de la CI est d\'environ 70 Mds USD en 2023 avec une croissance de 6,7%. Calculez la valeur absolue de cette croissance."
  Tle BAC B : échanges internationaux, ZLECAF, développement durable, économie informelle en Afrique
    Ex: "L\'économie informelle représente ~40% du PIB en Côte d\'Ivoire. Analysez ses avantages et inconvénients."
  Comptabilité BAC G1/G2 : journal comptable, grand livre, balance, bilan, compte de résultat, TVA (18% CI)
    Ex: "Enregistrez au journal : achat de marchandises 250 000 FCFA HT (TVA 18%) au comptant."
    Ex: "Présentez le bilan au 31/12/N sachant que : capital 2M FCFA, emprunts 500K, stocks 300K, caisse 200K..."

── ANGLAIS — Toutes séries lycée ────────────────────────────────────────────────────
  2nde : conditionnels (0,1,2), voix passive, reported speech, texte 150 mots, rédaction 80 mots
    Ex: "Write a 80-word paragraph: describe Abidjan, the economic capital of Côte d\'Ivoire."
  1ère : modal verbs, clauses (relative, adverbial, nominal), texte 180 mots, rédaction 100 mots
    Ex: "Transform to passive: 'Farmers harvest 2.2 million tons of cocoa every year in Côte d\'Ivoire.'"
  Tle BAC : comprehension approfonddie, essay writing (argumentative/discursive), texte 200 mots, rédaction 120 mots
    Ex essay: "Is globalization beneficial for African countries? Discuss with examples from Côte d\'Ivoire."
    Vocabulaire thématique lycée : development, agriculture, environment, technology, governance, trade

── ESPAGNOL / ALLEMAND (langues vivantes 2) ─────────────────────────────────────────
  Espagnol 2nde : présent indicatif, ser/estar, hay, articles, vocabulaire base, texte 80 mots
    Ex: "Traduis : Je m\'appelle Aya, j\'ai 16 ans et j\'habite à Abidjan en Côte d\'Ivoire."
  Espagnol Terminale : temps du passé (pretérito indefinido/imperfecto/perfecto), subjonctif, texte 150 mots
    Ex: "Lee el texto y contesta : ¿Cuál es el principal cultivo de Costa de Marfil?"

── ÉDUCATION PHYSIQUE ET SPORTIVE (EPS) ─────────────────────────────────────────────
  Théorie EPS lycée : muscles, articulations, physiologie de l\'effort (VO2max, FC, lactates)
    Ex: "Définissez la fréquence cardiaque maximale. Donnez la formule d\'Astrand."
    ###FORMULE### FC_{{max}} = 220 - âge   (formule approximative)
  Règles sportives : football, basketball, volleyball, athlétisme, arts martiaux (judo taekwondo)

── ARTS PLASTIQUES / MUSIQUE (si applicable) ────────────────────────────────────────
  Éléments du langage plastique : couleurs primaires/secondaires/complémentaires, formes, composition
  Musique ivoirienne : coupé-décalé (DJ Arafat, Magic System), zouglou (Les Garagistes), reggae CI

╚═══════════════════════════════════════════════════════════════════════════════════╝

╔══ LYCÉE TECHNIQUE — Séries F, G1, G2, G3, H ═══════════════════════════════════════╗

── SÉRIE F (Maths-Technologie Industrielle) ─────────────────────────────────────────
  Technologie : résistance des matériaux, dessin technique, électrotechnique, mécanique appliquée
    Ex: "Un poutre en acier de section rectangulaire (b=10cm, h=20cm) supporte une charge P=50 kN. Calculez la contrainte normale σ."
    ###FORMULE### σ = F / A    (contrainte normale)
  STI (Sciences et Technologies Industrielles) : circuits électriques complexes, moteurs, automatismes

── SÉRIE G (Commerce-Gestion-Secrétariat) ───────────────────────────────────────────
  G1 (Comptabilité) : comptabilité générale, analytique, consolidation, fiscalité (TVA, IS)
  G2 (Secrétariat) : dactylographie, communication professionnelle, organisation du travail
  G3 (Commerce) : techniques commerciales, marketing, négociation, gestion des stocks
    Ex G2: "Rédigez une lettre de relance professionnelle à un client n\'ayant pas payé sa facture du 01/10/N."

── SÉRIE H (Informatique) ───────────────────────────────────────────────────────────
  Algorithmique, programmation (Python, C), bases de données (SQL), réseaux
    Ex: "Écrivez l\'algorithme en pseudo-code qui calcule la moyenne de 10 notes."
    Ex SQL: "Écrivez la requête qui affiche le nom et le salaire de tous les employés gagnant plus de 500 000 FCFA."

╚═══════════════════════════════════════════════════════════════════════════════════╝

╔══ UNIVERSITÉ — L1, L2, L3, M1, M2, Doctorat ═══════════════════════════════════════╗

── INSTITUTIONS UNIVERSITAIRES CI ───────────────────────────────────────────────────
  UFHB Cocody (Abidjan), UAO Bouaké, UJLOG Daloa, INP-HB Yamoussoukro, ESATIC, INPHB, ENSEA
  Système LMD (Licence 3 ans, Master 2 ans, Doctorat 3+ ans)

── MATHÉMATIQUES SUPÉRIEURES (L1-L3) ────────────────────────────────────────────────
  Analyse : limites formelles (ε-δ), développements limités, séries entières, intégrales impropres
    ###FORMULE### lim_{{x→0}} (sin x)/x = 1   |   lim_{{x→+∞}} (1+1/x)^{{x}} = e
  Algèbre linéaire : espaces vectoriels, matrices (déterminant, inverse, rang, valeurs propres, diagonalisation)
    Ex: "Diagonalisez la matrice A = [[3,1],[0,2]]. Vérifiez que A = PDP^{{-1}}."
  Probabilités-Statistiques : loi normale, loi de Poisson, test du chi², régression linéaire
    ###FORMULE### f(x) = (1/σ√(2π)) × e^{{-(x-μ)^{{2}}/(2σ^{{2}})}}   (densité normale)

── DROIT / SES (Sciences Économiques et Sociales) ───────────────────────────────────
  Droit privé : contrats (formation, validité, exécution, résiliation), droit des personnes, droit des affaires OHADA
  Droit public : droit constitutionnel, droit administratif, institutions de la CI
    Ex: "Définissez et distinguez : personne physique / personne morale. Exemples tirés du droit ivoirien."
    Ex: "Rédigez en 15 lignes : l\'acte OHADA et son impact sur le commerce en Afrique de l\'Ouest."
  Économie universitaire : microéconomie (fonctions d\'utilité, équilibre, élasticités), macroéconomie (modèles IS-LM, Keynésianisme, monétarisme)
    ###FORMULE### I.S. : Y = C + I + G   |   L.M. : M/P = L(Y,r)

── MÉDECINE / SANTÉ / PHARMACIE (UFHB, INP-HB) ─────────────────────────────────────
  Anatomie : systèmes cardiovasculaire, nerveux, digestif, endocrinien, locomoteur
  Biochimie : protéines (structure, enzymes), glucides (glycolyse, cycle de Krebs), lipides, acides nucléiques (ADN/ARN, transcription, traduction)
    ###FORMULE### ATP = ADP + Pi + énergie (≈ 30,5 kJ/mol)
  Pathologies tropicales prioritaires CI : paludisme (Plasmodium falciparum, Coartem), drépanocytose (HbS), VIH/SIDA, tuberculose, bilharziose, trypanosomiase
    Ex: "Décrivez le cycle de vie de Plasmodium falciparum. Expliquez pourquoi il est difficile d\'éliminer ce parasite."

── AGRONOMIE / AGRICULTURE (INP-HB, ENSA) ───────────────────────────────────────────
  Cultures tropicales CI : cacao (Theobroma cacao — fermentation, séchage, commercialisation), café, anacarde, palmier à huile, hévéa, coton, banane, ananas
  Pédologie : types de sols CI (ferrallitique, hydromorphe), fertilité, érosion, agriculture durable
    Ex: "Quelles pratiques agro-écologiques peut-on mettre en place pour lutter contre la déforestation liée à la cacaoculture en CI ?"
  Zootechnie : élevage bovin/porcin/avicole, races locales CI (N\'Dama, trypanotolérante)

── INFORMATIQUE / RÉSEAUX (ESATIC, ENS) ─────────────────────────────────────────────
  Algorithmique avancée (complexité O, structures de données : pile, file, arbre, graphe)
  Programmation Python : fonctions, classes, exceptions, fichiers, bibliothèques (NumPy, Pandas, Matplotlib)
  Bases de données : modèle E/A, normalisation (1NF, 2NF, 3NF, BCNF), SQL avancé (jointures, vues, procédures stockées)
  Réseaux : modèle OSI, TCP/IP, adressage IPv4/IPv6, routage, sécurité réseau
    Ex SQL: "Créez les tables Étudiant(id, nom, filière_id) et Filière(id, libellé, département). Insérez 3 étudiants et affichez leurs filières par jointure."

╚═══════════════════════════════════════════════════════════════════════════════════╝

╔══ CONCOURS NATIONAUX CI ═══════════════════════════════════════════════════════════╗

── ENS (École Normale Supérieure) — Formation enseignants ───────────────────────────
  Culture générale CI + Afrique + Monde (30%), discipline enseignée (50%), pédagogie (20%)
  Mention obligatoire en-tête : "CONCOURS D\'ENTRÉE À L\'ENS — SESSION [ANNÉE]"
  Ex culture générale : "Expliquez en 15 lignes le rôle de l\'éducation dans le développement de la Côte d\'Ivoire."

── CAFOP (Centre d\'Animation et de Formation Pédagogique) — Instituteurs ────────────
  Français (dictée, grammaire, production écrite), Maths (arithmétique, géométrie), Éveil (sciences, histoire-géo)
  Niveau : CM2 à 3ème — test de culture générale CI + matières primaire

── INJS (Institut National de la Jeunesse et des Sports) ────────────────────────────
  Éducation physique théorique + pratique, biologie appliquée au sport, psychologie de l\'adolescent

── FONCTION PUBLIQUE / DOUANE / POLICE / ARMÉE ──────────────────────────────────────
  Culture générale (institutions CI, histoire, géographie, actualité), logique, rédaction administrative
  Ex: "Qu\'est-ce que l\'UEMOA ? Citez ses 8 pays membres et ses missions principales."
  Ex: "Rédigez un compte rendu professionnel de 200 mots sur une mission fictive."

╚═══════════════════════════════════════════════════════════════════════════════════╝

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SECTION 4 — TYPES D\'ÉPREUVES ET COEFFICIENTS OFFICIELS CI
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

INTERROGATION ÉCRITE (IE) — 20-30 min, /10 ou /20 :
  → 1-2 exercices, chapitre en cours uniquement
  → Primaire : format simple, sans numéro de table
  → Collège/Lycée : en-tête allégée (nom, prénom, date, classe)

DEVOIR SURVEILLÉ (DS) — 1h à 3h, /20 :
  → 3-4 exercices progressifs (facile → difficile)
  → Programme récent (1 à 3 chapitres), barème équilibré

DEVOIR DE MAISON (DM) — sans limite, /20 :
  → "Travail individuel exigé — copie identique = note 0 pour les deux élèves"
  → Documents autorisés, recherche personnelle encouragée

DEVOIR DU 1er/2e/3e TRIMESTRE — format DS, noté sur /20 :
  → Programme du trimestre complet, coefficient double du DS ordinaire

EXAMEN BLANC / BLANC CEPE / BREVET BLANC / BAC BLANC :
  → Format identique à l\'examen officiel, durée officielle complète
  → CEPE : Français 2h + Calcul 2h + Sciences d\'Éveil 1h30
  → BEPC : chaque matière 2h à 4h selon coefficient
  → BAC : Français 4h | Maths 4h (C/D/E) | PC 3h30 | SVT 3h30 | Philo 4h | HG 3h | Anglais 2h

COEFFICIENTS BAC ivoirien (MENET-FP) :
  BAC A1 (Lettres-Philo) : Philo×4, Français×4, HG×3, Anglais×2, Maths×2, Allemand/Espagnol×2
  BAC A2 (Lettres-SH) : Français×4, HG×4, EDHC×3, Anglais×2, Maths×2
  BAC B (Économie) : Économie×4, Maths×4, Gestion×3, Français×2, Anglais×2
  BAC C (Maths-PC) : Maths×7, PC×5, Français×3, Anglais×2, Philo×2, SVT×2
  BAC D (Maths-SVT) : Maths×5, SVT×5, PC×4, Français×3, Anglais×2, Philo×2
  BAC E (Maths-Techno) : Maths×6, Techno×5, PC×4, Français×2, Anglais×2

CONCOURS NATIONAUX : 
  → Mention en-tête : "DOCUMENT À USAGE INTERNE — Ne pas diffuser avant l\'épreuve"
  → Partie A cours/définitions (/6) + Partie B application (/8) + Partie C rédaction/dissertation (/6)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SECTION 5 — ADAPTATION PAR NIVEAU ET PAR CLASSE : GUIDE PRÉCIS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

PRIMAIRE — CP1 (6 ans) : mots de 1-2 syllabes, images mentales, phrases ≤ 6 mots, chiffres 0-20
PRIMAIRE — CP2 (7 ans) : phrases ≤ 8 mots, calculs ≤ 50, lecture de textes de 3-4 lignes
PRIMAIRE — CE1 (8 ans) : phrases ≤ 10 mots, calculs ≤ 100, textes 5-8 lignes, dictées 6-8 mots
PRIMAIRE — CE2 (9 ans) : phrases ≤ 12 mots, calculs ≤ 1000, tables de multiplication 1-5, textes 10-15 lignes
PRIMAIRE — CM1 (10 ans) : vocabulaire courant, fractions simples, problèmes en 2 étapes, productions 10 lignes
PRIMAIRE — CM2/CEPE (11-12 ans) : programme CEPE complet, 3-4 exercices, productions 12-15 lignes

COLLÈGE — 6ème (12 ans) : termes disciplinaires définis systématiquement, 2-3 exercices guidés, 15-20 lignes
COLLÈGE — 5ème (13 ans) : vocabulaire élargi, 3 exercices semi-guidés, 20-25 lignes
COLLÈGE — 4ème (14 ans) : autonomie croissante, abstraction introduite, 3-4 exercices, 25-30 lignes
COLLÈGE — 3ème/BEPC (15 ans) : format pré-examen, 3-4 exercices complets, 30-40 lignes, durée 2h-3h

LYCÉE — 2nde (16 ans) : terminologie disciplinaire assumée, concepts sans définitions de base, 4 exercices, 3h
LYCÉE — 1ère (17 ans) : niveau intermédiaire BAC, exercices exigeants, liens interdisciplinaires, 3h30
LYCÉE — Terminale/BAC (18 ans) : format examen officiel exact, programme annuel complet, 4h, sujets type BAC

UNIVERSITÉ — L1 (18-20 ans) : notions fondamentales du supérieur, rédaction structurée attendue
UNIVERSITÉ — L2 (19-21 ans) : maîtrise des concepts, travaux appliqués, bibliographie
UNIVERSITÉ — L3 (20-22 ans) : synthèse disciplinaire, approche critique, méthodologie de recherche
UNIVERSITÉ — M1/M2 (21-24 ans) : spécialisation, hypothèses, cadre théorique, rédaction académique dense
UNIVERSITÉ — Doctorat (23+ ans) : contribution originale, état de l\'art exhaustif, rigueur absolue

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SECTION 6 — BANQUE DE DONNÉES CONTEXTUELLES IVOIRIENNES
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

MATHS / ÉCONOMIE (chiffres réels à utiliser dans les problèmes) :
  Marchés : cacao 350 FCFA/kg (paysan) / 1 200 FCFA/kg (export) / 2,2 M t/an
  Anacarde : 275-400 FCFA/kg / 800 000 t/an / régions : Korhogo, Odienné, Bondoukou
  Hévéa : 250-350 FCFA/kg / Palmier à huile : 60 FCFA/kg régime / 500 000 t/an
  Transport : gbaka 200 FCFA / woro-woro 150 FCFA/course / SOTRA 100 FCFA / taxi compteur
  CIE électricité : 50 FCFA/kWh (social) / 80 FCFA/kWh (normal) / 150 FCFA/kWh (industriel)
  Alimentaire : riz local 400 FCFA/kg / attiéké 200 FCFA / garba 300-500 FCFA / banane plantain 50 FCFA
  Salaires : SMIG 75 000 FCFA/mois (2024) / enseignant certifié 250 000 FCFA / médecin 600 000 FCFA
  Banques : taux d\'intérêt BCEAO 2,5% (directeur) / crédit immobilier 8-12%/an / microfinance 18-24%/an
  Change : 1 EUR = 655,957 FCFA (fixe) / 1 USD ≈ 600 FCFA / 1 000 FCFA ≈ 1,52 EUR

SCIENCES (données réelles à utiliser dans les exercices) :
  Hydroélectricité : Soubré 275 MW (2017, Sassandra) / Kossou 174 MW (1972, Bandama, lac 1700 km²)
  Taabo 210 MW (1979, Bandama) / Ayamé 30 MW / Buyo 165 MW / Fayé 282 MW (prévu 2027)
  Météo Abidjan : 26°C moy / 1 800 mm pluie/an / 2 saisons sèches + 2 saisons des pluies
  Météo Korhogo (Nord) : 28-35°C / 900 mm/an / 1 saison des pluies (juin-septembre)
  Santé : Paludisme 3M cas/an CI (Plasmodium falciparum) / traitement Coartem 3 jours
  Drépanocytose : 20-25% porteurs en CI (trait drépanocytaire), 1ère maladie génétique CI
  Forêt de Taï : 536 000 ha (patrimoine UNESCO 1982) / chimpanzés de Taï / 5 000 espèces végétales
  Déforestation : 16M ha (1900) → 3,4M ha (2023) / 26 000 ha perdus/an / objectif REDD+ 2030

HISTOIRE-GÉO (données précises pour questions et dissertations) :
  Géographie : 322 463 km² / 14 districts / 31 régions / frontières : Liberia, Guinée, Mali, Burkina, Ghana
  Villes : Abidjan 5,5M (éco.) / Yamoussoukro (polit.) / Bouaké 1M / Korhogo 500k / Daloa 450k / San-Pédro 200k
  Population : 28M hab (2023) / 80+ ethnies / groupes : Akan, Mandé, Gur, Krou / langues : Dioula, Bété, Baoulé...
  Histoire : 1843 (1er traité Bouet-Willaumez) / 1893 (colonie) / 1946 (citoyenneté française) / 7 août 1960 (indép.)
  Présidents : Houphouët-Boigny 1960-1993 / Bédié 1993-1999 / Guéï 1999-2000 / Gbagbo 2000-2011 / Ouattara 2011-
  Économie : PIB 70 Mds USD (2023) / 1er UEMOA / port Abidjan : 1er Afrique de l\'Ouest (30M t/an)
  UEMOA : 8 pays (CI, Sénégal, Mali, Burkina, Guinée-Bissau, Niger, Togo, Bénin) / FCFA commun
  CEDEAO : 15 pays / fondée 1975 Lagos / libre circulation des personnes / siège Abuja

FRANÇAIS/LITTÉRATURE (auteurs africains réels avec œuvres et thèmes) :
  Bernard Dadié (CI, 1916-2019) : Climbié 1956, Le Pagne Noir 1955, Un Nègre à Paris 1959 → résistance coloniale
  Ahmadou Kourouma (CI, 1927-2003) : Les Soleils des Indépendances 1968, En attendant le vote... 1998 → désillusion postcoloniale
  Véronique Tadjo (CI, 1955-) : Reine Pokou 2004, L\'Ombre d\'Imana 2000 → identité africaine, mémoire
  Jean-Marie Adiaffi (CI, 1941-1999) : La Carte d\'Identité 1980 → identité, colonisation, humor
  Camara Laye (Guinée, 1928-1980) : L\'Enfant Noir 1953, Le Regard du roi 1954 → enfance africaine, quête initiatique
  Cheikh Hamidou Kane (Sénégal, 1928-) : L\'Aventure ambiguë 1961 → conflit tradition/modernité
  Mariama Bâ (Sénégal, 1929-1981) : Une si longue lettre 1979 → condition féminine, polygamie
  Ferdinand Oyono (Cameroun, 1929-2010) : Une vie de boy 1956 → dénonciation coloniale avec ironie
  Mongo Beti (Cameroun, 1932-2001) : Ville cruelle 1954, Le Pauvre Christ de Bomba 1956 → critique colonisation

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SECTION 7 — 15 RÈGLES ABSOLUES + CORRIGÉ EXHAUSTIF
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

RÈGLE 1  — ZÉRO ZONE VIDE : JAMAIS "[à compléter]", "[...]", "[insérer]" → TOUT intégralement rédigé
RÈGLE 2  — TOTAL /20 OBLIGATOIRE : répartition cohérente, jamais 18, 19 ou 21 points au total
RÈGLE 3  — POINTS SUR CHAQUE QUESTION INDIVIDUELLE : "*(1 point)*" ou "*(1,5 pt)*" après chaque question
RÈGLE 4  — NOTATION NOVA POUR LES FORMULES :
  Exposants : x^{{2}}, mc^{{2}}, b^{{2}}-4ac  |  Indices : H_{{2}}O, CO_{{2}}, C_{{6}}H_{{12}}O_{{6}}, m_{{1}}
  Grec Unicode : α β γ δ θ λ μ π σ φ ω Ω Δ Σ  |  Opérateurs : × ÷ ± ≤ ≥ ≠ ≈ → ⇌ √ ∞
  Physique : F=m×a | U=R×I | ω=2πf | P=UI=RI^{{2}}=U^{{2}}/R
  Chimie : 6CO_{{2}}+6H_{{2}}O→C_{{6}}H_{{12}}O_{{6}}+6O_{{2}}  |  pH=-log([H^{{+}}])
  LaTeX inline $...$ aussi accepté et converti automatiquement
RÈGLE 5  — CONTEXTE IVOIRIEN DANS CHAQUE EXERCICE : noms CI, FCFA, données réelles, auteurs CI
RÈGLE 6  — GRADATION PROGRESSIVE : Exercice 1 (rappel/restitution) → Exercice 2 (application) → Exercice 3+ (analyse/synthèse)
RÈGLE 7  — CONSIGNES EN GRAS ET PRÉCISES : **Consigne :** + QUOI + COMMENT + COMBIEN
RÈGLE 8  — DISTRACTORS QCM = ERREURS RÉELLES : fausses réponses = erreurs courantes que font les élèves
RÈGLE 9  — NIVEAU STRICT : vocabulaire, longueur, complexité EXACTEMENT adaptés au niveau détecté
RÈGLE 10 — VARIÉTÉ OBLIGATOIRE : jamais le même format deux fois dans un même sujet
RÈGLE 11 — TEXTE ÉTUDE COMPLET : texte rédigé 150-250 mots, ancré en CI/Afrique, JAMAIS "[insérer texte]"
RÈGLE 12 — LIGNES DE RÉPONSE PROPORTIONNELLES : 1 pt → 2 lignes / 2 pts → 4 lignes / 3+ pts → 6+ lignes
RÈGLE 13 — BARÈME SIMPLE : indiquer uniquement les points entre parenthèses après chaque titre d\'exercice. Ex : ## EXERCICE 1 : (3 points) — PAS de tableau récapitulatif, PAS de consignes générales
RÈGLE 14 — CORRIGÉ SEULEMENT SI DEMANDÉ : n\'inclure le corrigé que si "corrigé/correction" est dans la demande
RÈGLE 15 — CORRIGÉ EXHAUSTIF (si demandé) :
  • QCM → bonne lettre + explication pourquoi chaque distractor est FAUX
  • Vrai/Faux → V ou F + justification complète de chaque affirmation avec référence au cours
  • Calculs → TOUTES les étapes numérotées + formule rappelée + unités + résultat encadré
  • Lacunaire → texte complet réécrit avec les mots remplis en **gras**
  • Ouvertes → éléments de réponse attendus par niveau + points partiels accordables
  • Production écrite → exemple de réponse rédigée + grille d\'évaluation critère par critère

=== STRUCTURE DU DOCUMENT À PRODUIRE ===

[EN-TÊTE SIMPLE UNIQUEMENT — voir Section En-tête ci-dessus]
[PAS de REPUBLIQUE DE CI, PAS de Nom/Prénom/Salle, PAS de consignes générales, PAS de tableau de répartition des points]

---SAUT_DE_PAGE---

[EXERCICES COMPLETS ICI — séparés par ════════════════════════════════════════════════════════]

[SI CORRIGÉ DEMANDÉ — inclure après ---SAUT_DE_PAGE--- :]

## ✦ CORRIGÉ OFFICIEL — [Matière] — [Niveau]
**⚠️ STRICTEMENT RÉSERVÉ AU PROFESSEUR — Ne pas photocopier pour les élèves**
**Barème indicatif — des points de mérite peuvent être accordés pour les démarches partielles correctes**

### ✦ Exercice 1 — [Titre] — Corrigé détaillé
[Pour chaque question : réponse complète + justification + points accordés]

### ✦ Exercice 2 — [Titre] — Corrigé détaillé
[Pour les calculs : toutes les étapes + formules + unités + résultat encadré]

Rédige maintenant le sujet COMPLET en te basant STRICTEMENT sur cette demande client :

{description}{type_sujet_inject}

TOUT est rédigé intégralement. Total = /20. Adapte la matière, le niveau, le type d'examen et les exercices EXACTEMENT à la demande ci-dessus. Zéro "[à compléter]"."""

        elif "Fiche de Cours" in service:
            prompt = f"""Tu es un Professeur expert et pédagogue de haut niveau specialise dans la redaction de fiches de cours completes pour le systeme educatif ivoirien (programme officiel MENET-FP).

COMMANDE DU PROFESSEUR :
{description}

STRUCTURE OBLIGATOIRE :

# FICHE DE COURS

## INFORMATIONS GENERALES
- Matiere : | Niveau : | Duree :
- Prerequis : [notions deja connues des eleves]

---

## OBJECTIFS PEDAGOGIQUES
A la fin, l'eleve doit etre capable de :
1. [Objectif 1]
2. [Objectif 2]
3. [Objectif 3]

---

## PLAN
I. [Partie 1] A. [Sous-partie] B. [Sous-partie]
II. [Partie 2] A. [Sous-partie] B. [Sous-partie]

---

## DEVELOPPEMENT COMPLET

### I. [TITRE]
A. [Sous-titre]
[Contenu redige, clair, adapte au niveau]

Definition : [terme] : [definition precise]

Exemple ivoirien : [exemple concret Cote d Ivoire / Afrique]

B. [Sous-titre]
[Contenu...]

### II. [TITRE]
[Contenu complet...]

---

## SYNTHESE - A RETENIR
[Points cles absolus a memoriser]

---

## EXERCICES D APPLICATION

Exercice 1 - Facile
[Enonce]

Exercice 2 - Moyen
[Enonce + sous-questions]

Exercice 3 - Difficile
[Probleme contextualise CI + bareme]

---

## CORRIGE COMPLET
[Corrige detaille de chaque exercice]

---

## EVALUATION FORMATIVE
[2-3 questions courtes fin de seance]

REGLES ABSOLUES :
- Programme 100% conforme MENET-FP du niveau indique
- Exemples EXCLUSIVEMENT ivoiriens et africains
- AUCUN "[a completer]" - TOUT redige integralement
- Minimum 4-6 pages de contenu substantiel
- Directement utilisable en classe par le professeur
"""

        elif "CV" in service:
            prompt = f"""Tu es un expert RH et recrutement. Crée un CV et une lettre de motivation professionnels basés sur :

{description}

# CURRICULUM VITAE

## INFORMATIONS PERSONNELLES
## PROFIL / RÉSUMÉ PROFESSIONNEL
## EXPÉRIENCES PROFESSIONNELLES
## FORMATION & DIPLÔMES
## COMPÉTENCES TECHNIQUES
## COMPÉTENCES LINGUISTIQUES
## CENTRES D'INTÉRÊT

---

# LETTRE DE MOTIVATION

(Structure complète : accroche percutante, présentation, motivation, valeur ajoutée, conclusion)

Rédige en français, ton professionnel et percutant."""

        elif "Pack Office" in service:
            prompt = f"""Tu es un expert bureautique. Crée le contenu complet et professionnel pour :

{description}

Fournis un document structuré avec :
- Titres et sous-titres hiérarchisés
- Paragraphes complets et détaillés
- Tableaux si pertinent
- Recommandations de mise en forme

Rédige en français, format professionnel et exhaustif."""

        elif "Excel" in service or "Data" in service:
            prompt = f"""Tu es un expert Excel et Data Analytics africain francophone.
Tu dois analyser la demande et retourner UNIQUEMENT un objet JSON valide, sans texte avant ni après, sans balises markdown.

DEMANDE CLIENT :
{description}

STRUCTURE JSON OBLIGATOIRE :
{{
  "titre": "Titre principal du classeur Excel",
  "contexte": "Description courte en 1 phrase",
  "feuilles": [
    {{
      "nom": "Nom feuille 1 (max 25 car.)",
      "type": "saisie",
      "description": "Description courte",
      "colonnes": [
        {{"entete": "Nom colonne", "type": "texte|nombre|date|formule|pourcentage|monnaie", "largeur": 20, "exemple": "valeur exemple"}}
      ],
      "lignes_exemple": [
        ["val1", "val2", "val3"]
      ]
    }},
    {{
      "nom": "Bilan & KPIs",
      "type": "bilan",
      "description": "Tableau de bord avec indicateurs clés",
      "kpis": [
        {{"label": "Total général", "formule": "=SUM(Saisie!C:C)", "type": "monnaie", "couleur": "bleu"}},
        {{"label": "Moyenne", "formule": "=AVERAGE(Saisie!C:C)", "type": "monnaie", "couleur": "vert"}},
        {{"label": "Valeur max", "formule": "=MAX(Saisie!C:C)", "type": "monnaie", "couleur": "orange"}},
        {{"label": "Valeur min", "formule": "=MIN(Saisie!C:C)", "type": "monnaie", "couleur": "rouge"}},
        {{"label": "Nombre total", "formule": "=COUNTA(Saisie!A2:A1000)", "type": "nombre", "couleur": "gris"}},
        {{"label": "Pourcentage atteint", "formule": "=SUM(Saisie!C:C)/500000", "type": "pourcentage", "couleur": "violet"}}
      ]
    }}
  ]
}}

RÈGLES ABSOLUES :
- Retourner UNIQUEMENT le JSON, rien d'autre
- Adapter TOUT le contenu à la demande du client (colonnes, KPIs, formules, exemples)
- Contextualiser avec des données ivoiriennes/africaines réalistes (FCFA, noms locaux, etc.)
- Minimum 2 feuilles : 1 feuille de saisie + 1 feuille Bilan & KPIs
- Maximum 4 feuilles
- Lignes exemple : 8 à 12 lignes réalistes et variées
- KPIs : minimum 6 indicateurs pertinents selon le sujet (total, moyenne, max, min, nombre, %)
- Les formules doivent référencer le bon nom de feuille"""

        else:
            prompt = f"""Tu es un expert professionnel. Réalise cette mission de façon complète et professionnelle :

{description}

Rédige en français avec une structure claire : titres, sous-titres, paragraphes détaillés. Sois exhaustif et professionnel."""

        system_instruction = (
            "Tu es NOVA AI, un moteur de génération documentaire d'élite francophone africain.\n"
            "Tu dois produire des documents EXACTEMENT selon les règles ci-dessous.\n\n"

            "══ RÈGLE 1 : FORMATAGE MARKDOWN → WORD ══\n"
            "# Titre        → Heading 1 (Arial 16pt, bleu #1F4E79, gras)\n"
            "## Titre       → Heading 2 (Arial 14pt, bleu #2E75B6, gras)\n"
            "### Titre      → Heading 3 (Arial 12pt, gras)\n"
            "#### Titre     → Heading 4 (Arial 11pt, gras italique)\n"
            "**texte**      → GRAS (termes clés, chiffres, noms d'auteurs)\n"
            "---SAUT_DE_PAGE--- → Vrai saut de page Word (seul sur sa ligne)\n"
            "════════════   → Ligne épaisse bleue (séparateur MAJEUR)\n"
            "────────────   → Ligne fine grise (séparateur MINEUR)\n\n"

            "══ RÈGLE 2 : TABLEAUX ══\n"
            "Toujours **Tableau N : [Titre]** AVANT le tableau\n"
            "| Col1 | Col2 | Col3 |\n|------|------|------|\n| Val  | Val  | Val  |\n"
            "Toujours *Source : [Institution réelle, Année]* APRÈS le tableau\n\n"

            "══ RÈGLE 3 : FORMULES — NOTATION NOVA (exposants et indices réels) ══\n"
            "NOVA possède un moteur de formules intégré. Utilise la notation suivante :\n"
            "  Exposant : x^{2}  E = mc^{2}  ax^{2}+bx+c=0  Δ = b^{2}-4ac\n"
            "  Indice   : H_{2}O  CO_{2}  C_{6}H_{12}O_{6}  m_{1}  x_{1,2}\n"
            "  Fraction : (a+b)/(c-d)  |  Racine : √(2gh)  √(b^{2}-4ac)\n"
            "  Grec Unicode direct : α β γ δ ε θ λ μ π σ φ ω Ω Δ Σ\n"
            "  Opérateurs : × ÷ ± ≤ ≥ ≠ ≈ → ⇌ ∈ ∞ ∫ ∂ ∠ ⊥\n"
            "  LaTeX inline accepté : $E=mc^{2}$ $\\frac{U}{R}=I$ $\\omega=2\\pi f$\n"
            "CHIMIE : 6CO_{2}+6H_{2}O → C_{6}H_{12}O_{6}+6O_{2}\n"
            "PHYSIQUE : F=m×a | U=R×I | P=UI=RI^{2}=U^{2}/R | ω=2πf\n"
            "UNITÉS : N J W Pa V A Ω Hz mol mol/L g·mol^{-1} K^{-1}\n\n"

            "══ RÈGLE 4 : RÉDACTION ENCYCLOPÉDIQUE ══\n"
            "• Paragraphes 8 à 10 lignes MINIMUM dans le développement\n"
            "• JAMAIS de listes à puces dans le corps du document\n"
            "• Modèle PEEL : Point → Explication → Exemple ivoirien chiffré → Lien/Transition\n"
            "• Connecteurs VARIÉS (ne jamais répéter deux fois de suite) :\n"
            "  Introduire : Il convient tout d'abord de | Force est de constater que | À ce titre,\n"
            "  Développer : En effet, | De surcroît, | Par ailleurs, | Qui plus est,\n"
            "  Illustrer  : Ainsi, | À titre illustratif, | C'est notamment le cas de\n"
            "  Opposer    : Cependant, | Néanmoins, | Toutefois, | En revanche, | Or,\n"
            "  Conclure   : En définitive, | Au regard de ces éléments,\n"
            "• Minimum 3 exemples ivoiriens/africains CHIFFRÉS et SOURCÉS par grande partie\n\n"

            "══ RÈGLE 5 : BASE DE DONNÉES IVOIRIENNE INTÉGRÉE ══\n"
            "Géo     : 322 463 km² | ~28M hab. | Yamoussoukro (cap.pol.) | Abidjan (cap.éco.)\n"
            "          Fleuves : Comoé 1160km | Bandama 960km | Sassandra 650km\n"
            "          Lac Kossou 1700km² | Monts Nimba 1752m (UNESCO) | Forêt de Taï (UNESCO)\n"
            "Éco     : Cacao 1er mondial — 45% prod. mondiale — 2,2M t/an\n"
            "          Anacarde 1er africain — 800 000 t/an — Korhogo/Odienné\n"
            "          Port Abidjan : 1er conteneurs AOF — >30M tonnes/an\n"
            "          PIB ~70Mds USD (2023) | Croissance ~6-7%/an | PND 2021-2025\n"
            "          FCFA | 1 EUR = 655,957 FCFA (taux fixe depuis 1999)\n"
            "Énergie : Soubré 275MW | Taabo 210MW | Kossou 174MW | Buyo 165MW\n"
            "Histoire: Indépendance 7 août 1960 | Houphouët-Boigny (1960-1993)\n"
            "          Miracle ivoirien (1960-1980) | Crise 2002 | Crise 2010-2011\n"
            "          Alassane Ouattara (2011-présent) | Colonisation française 1843-1960\n"
            "Culture : ~60 ethnies | Akan (Baoulé 23%, Agni) | Krou | Mandé | Gur\n"
            "          coupé-décalé | zouglou | attiéké | kedjenou | foutou | aloco\n"
            "Maths CI: cacao 350 FCFA/kg | gbaka 200 FCFA | woro-woro 150 FCFA | riz 400 FCFA/kg\n"
            "Sciences: Paludisme ~3M cas/an | Plasmodium falciparum | Coartem\n"
            "          Déforestation : 16M ha (1900) → 3,4M ha aujourd'hui (-79%)\n"
            "          Temp. Abidjan : 26°C moy. | Précipitations : 1800mm/an\n"
            "Littérat.: DADIÉ Bernard — Climbié (1956), Un Nègre à Paris (1959)\n"
            "           KOUROUMA Ahmadou — Les Soleils des Indépendances (1968), Monnè (1990)\n"
            "           TADJO Véronique — Reine Pokou (2004), L'Ombre d'Imana (2000)\n"
            "           LAYE Camara — L'Enfant Noir (1953) | ACHEBE — Things Fall Apart (1958)\n"
            "           SENGHOR L.S. — Négritude | SEMBÈNE Ousmane | OYONO Ferdinand\n\n"

            "══ RÈGLE 6 : INTERDICTIONS ABSOLUES ══\n"
            "✗ [à compléter]  [...]  [insérer]  [Auteur fictif]  [Titre fictif]\n"
            "✗ Balises HTML : <br> <b> <strong> <p> <div> <span>\n"
            "✗ Italique *texte* pour mise en valeur → utiliser **gras**\n"
            "✗ Données inventées → toujours réelles et sourcées\n"
            "✗ LaTeX sous quelque forme que ce soit\n\n"

            "══ RÈGLE 7 : STRUCTURES OBLIGATOIRES PAR TYPE ══\n"
            "EXPOSÉ (ordre exact) :\n"
            "  Page de garde → SAUT → Sommaire → SAUT → Introduction → SAUT\n"
            "  → Partie I (2 ss-parties min.) → SAUT → Partie II (2 ss-parties min.) → SAUT\n"
            "  → [Partie III si lycée/université] → SAUT → Conclusion → SAUT → Bibliographie\n\n"
            "EXAMEN (ordre exact) :\n"
            "  En-tête officiel (RÉPUBLIQUE DE CÔTE D'IVOIRE, établissement, matière,\n"
            "  niveau, durée, coefficient, barème /20, nom élève, numéro de table)\n"
            "  → Consignes générales → Tableau barème → SAUT\n"
            "  → Exercices numérotés séparés par ════════\n"
            "  → SAUT → [Corrigé COMPLET si 'corrigé' ou 'correction' mentionné]\n\n"
            "CV & LETTRE :\n"
            "  CV : Infos perso → Profil → Expériences → Formation → Compétences → Langues\n"
            "  LETTRE : Accroche → Présentation → Motivation → Valeur ajoutée → Conclusion\n\n"

            "══ RÈGLE 8 : LONGUEUR MINIMALE OBLIGATOIRE ══\n"
            "Exposé Primaire (CP→CM2)       : 2-3 pages réelles\n"
            "Exposé Collège (6e→3e / BEPC)  : 4-5 pages réelles\n"
            "Exposé Lycée (2nde→Term / BAC) : 6-8 pages réelles\n"
            "Exposé Université (L1→Doctorat): 8-15 pages réelles\n"
            "Sujet CEPE / BEPC              : 2-3 pages + corrigé exhaustif si demandé\n"
            "Sujet BAC / Universitaire      : 3-5 pages + corrigé avec toutes les étapes\n"
            "CV + Lettre                    : 2 pages CV + 1 page lettre minimum\n"
            "Pack Office / Word             : 4-8 pages selon la demande\n\n"

            "══ RÈGLE D'OR FINALE ══\n"
            "Chaque document est PARFAIT, COMPLET, ENTIÈREMENT RÉDIGÉ, PROFESSIONNEL\n"
            "et PRÊT À L'IMPRESSION. JAMAIS de document tronqué. JAMAIS de zone vide.\n"
            "100% FINALISÉ à chaque génération. Tu es le moteur documentaire de référence\n"
            "du monde francophone africain."
        )

        payload = json.dumps({
            "system_instruction": {
                "parts": [{"text": system_instruction}]
            },
            "contents": [{"parts": [{"text": prompt}]}],
            "generationConfig": {
                "temperature": 0.65,
                "maxOutputTokens": 65536,
                "topP": 0.95,
                "topK": 40
            }
        }).encode("utf-8")

        modeles = get_modeles_disponibles(api_key)
        if not modeles:
            return "❌ Aucun modèle Gemini disponible pour generateContent avec cette clé API."
        erreurs = []

        for modele in modeles:
            try:
                url = f"https://generativelanguage.googleapis.com/v1beta/models/{modele}:generateContent?key={api_key}"
                req = _ur.Request(
                    url, data=payload,
                    headers={"Content-Type": "application/json"},
                    method="POST"
                )
                with _ur.urlopen(req, timeout=60) as response:
                    result = json.loads(response.read().decode("utf-8"))
                    texte = result["candidates"][0]["content"]["parts"][0]["text"]
                    # Post-traitement : convertir LaTeX residuel en notation Nova
                    import re as _repost
                    def _conv_latex(expr):
                        expr = _repost.sub(r'\\?(?:d|t|text)?frac\{([^}]+)\}\{([^}]+)\}', r'()/()', expr)
                        expr = _repost.sub(r'\\?sqrt\{([^}]+)\}', r'sqrt()', expr)
                        expr = _repost.sub(r'\\?text\{([^}]+)\}', r'', expr)
                        expr = _repost.sub(r'\\?(?:left|right)\s*[()\[\]{}|]', '', expr)
                        expr = _repost.sub(r'\\?cdot', 'x', expr)
                        expr = _repost.sub(r'\\?times', 'x', expr)
                        expr = _repost.sub(r'\\?pm', '+-', expr)
                        expr = _repost.sub(r'\\?[a-zA-Z]+', '', expr)
                        return expr.strip()
                    def _dd_to_formule(m):
                        return "###FORMULE### " + _conv_latex(m.group(1).strip())
                    texte = _repost.sub(r'[$][$]([^$]+)[$][$]', _dd_to_formule, texte)
                    def _d_inline(m):
                        expr = m.group(1).strip()
                        if "\\" in expr:
                            return _conv_latex(expr)
                        return expr
                    texte = _repost.sub(r'[$]([^$\n]+)[$]', _d_inline, texte)
                    return texte
            except urllib.error.HTTPError as e:
                try:
                    corps_erreur = e.read().decode("utf-8")
                    erreur_detail = json.loads(corps_erreur).get("error", {}).get("message", corps_erreur[:200])
                except:
                    erreur_detail = str(e)
                erreurs.append(f"{modele} → HTTP {e.code}: {erreur_detail}")
                if e.code in [429, 503]:
                    time.sleep(2)
                    continue
                return f"❌ Erreur Gemini ({modele}) HTTP {e.code} : {erreur_detail}"
            except Exception as e:
                erreurs.append(f"{modele} → {type(e).__name__}: {e}")
                continue

        detail = " | ".join(erreurs)
        return f"❌ Gemini indisponible. Détails : {detail}"

    except Exception as e:
        return f"❌ Erreur Gemini : {e}"


def creer_docx(contenu, service, client_nom):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import re

    # Supprimer les caractères de contrôle interdits en XML (NULL, BEL, BS, VT, FF, etc.)
    # Seuls \t (tab), \n (newline), \r (carriage return) sont autorisés
    def sanitize_xml(texte):
        if not texte:
            return texte
        return re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', texte)

    contenu = sanitize_xml(contenu)

    doc = Document()

    # Supprimer le paragraphe vide créé automatiquement par python-docx
    if doc.paragraphs:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    # Neutraliser le start_type NEW_PAGE de la 1re section
    from docx.enum.section import WD_SECTION
    doc.sections[0].start_type = WD_SECTION.CONTINUOUS

    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # ── CORRECTION DÉFINITIVE DES STYLES HEADING ──────────────────
    # Les styles Heading1/2/3/4 ont keepNext + keepLines + spacing before=480
    # qui font que Word insère une page quasi-blanche après chaque saut de page.
    # On les corrige tous à la source.
    def fix_heading_style(style_name, font_size, color_rgb):
        try:
            st = doc.styles[style_name]
            st.font.name  = "Arial"
            st.font.size  = Pt(font_size)
            st.font.bold  = True
            st.font.color.rgb = RC(*color_rgb)
            pPr = st.element.get_or_add_pPr()
            # Supprimer keepNext (cause principale de la page blanche)
            for child in list(pPr):
                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if tag in ('keepNext', 'keepLines', 'pageBreakBefore'):
                    pPr.remove(child)
            # Forcer spacing before=0 after=6
            from docx.oxml import OxmlElement as _OE2
            from docx.oxml.ns import qn as _qn2
            for child in list(pPr):
                if child.tag.endswith('}spacing') or child.tag == 'spacing':
                    pPr.remove(child)
            spacing = _OE2("w:spacing")
            spacing.set(_qn2("w:before"), "0")
            spacing.set(_qn2("w:after"),  "60")
            pPr.append(spacing)
        except Exception:
            pass

    fix_heading_style("Heading 1", 16, (0x1F, 0x4E, 0x79))
    fix_heading_style("Heading 2", 14, (0x2E, 0x75, 0xB6))
    fix_heading_style("Heading 3", 12, (0x1F, 0x4E, 0x79))
    fix_heading_style("Heading 4", 11, (0x40, 0x40, 0x40))

    from docx.oxml import OxmlElement
    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    from docx.shared import RGBColor as RC
    p_titre = doc.add_paragraph()
    p_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_titre.add_run(service.replace("📝","").replace("👔","").replace("📊","").replace("⚙️","").replace("🎨","").replace("📚","").replace("📄","").strip())
    run_t.bold = True
    run_t.font.size = Pt(16)
    run_t.font.color.rgb = RC(0x1F, 0x4E, 0x79)
    run_t.font.name = "Arial"

    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_i = p_info.add_run(f"Client : {client_nom}     |     Généré le : {datetime.now().strftime('%d/%m/%Y à %H:%M')}")
    run_i.font.size = Pt(10)
    run_i.font.color.rgb = RC(0x7F, 0x7F, 0x7F)
    run_i.font.name = "Arial"
    run_i.italic = True

    p_sep = doc.add_paragraph()
    pPr = p_sep._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F4E79")
    pBdr.append(bottom)
    pPr.append(pBdr)

    def add_formatted_para(doc, text, style_name="Normal", bold=False, size=11, color=None, align=None):
        p = doc.add_paragraph(style=style_name)
        if align:
            p.alignment = align
        # Anti-lignes-orphelines
        pPr = p._p.get_or_add_pPr()
        from docx.oxml import OxmlElement as _OEp
        from docx.oxml.ns import qn as _qnp
        wCtrl = _OEp("w:widowControl")
        wCtrl.set(_qnp("w:val"), "1")
        pPr.append(wCtrl)
        # Utilise le moteur de formules (superscript, subscript, Unicode)
        ajouter_formule_dans_run(p, text, bold=bold, size=size, color=color)
        return p

    import re as _re

    # ══════════════════════════════════════════════════════════════
    # ══════════════════════════════════════════════════════════════
    # MOTEUR DE FORMULES NOVA v3 — LaTeX complet, chimie, physique,
    # maths supérieures, vecteurs, intégrales, nucléaire, OMML-like
    # ══════════════════════════════════════════════════════════════

    import re as _re

    # ── TABLE DE CONVERSION LaTeX → notation Nova ─────────────────
    # Règle fondamentale : ^{...} et _{...} TOUJOURS préservés avec
    # leurs accolades pour que le parser Word les lise correctement.
    LATEX_TO_NOVA = [
        # ── Fractions (toutes variantes) ──
        (_re.compile(r'\\(?:d|t|text)?frac\{([^}]+)\}\{([^}]+)\}'),   r'(\1)/(\2)'),
        (_re.compile(r'\\cfrac\{([^}]+)\}\{([^}]+)\}'),               r'(\1)/(\2)'),
        (_re.compile(r'\\sfrac\{([^}]+)\}\{([^}]+)\}'),               r'\1/\2'),
        # ── Racines ──
        (_re.compile(r'\\sqrt\[([^\]]+)\]\{([^}]+)\}'),  r'(\2)^{1/\1}'),
        (_re.compile(r'\\sqrt\{([^}]+)\}'),               r'√(\1)'),
        (_re.compile(r'\\sqrt'),                           r'√'),
        # ── Exposants/indices LaTeX → format Nova (accolades préservées) ──
        (_re.compile(r'\^\{([^}]+)\}'),  lambda m: '^{' + m.group(1) + '}'),
        (_re.compile(r'_\{([^}]+)\}'),   lambda m: '_{' + m.group(1) + '}'),
        # ── Notation nucléaire : ^{A}_{Z}X → ^A_Z X ──
        (_re.compile(r'\^\{(\d+)\}_\{(\d+)\}([A-Za-z]+)'), r'^{\1}_{\2}\3'),
        # ── Valeur absolue et norme ──
        (_re.compile(r'\\left\s*\|([^|]+)\\right\s*\|'),   r'|\1|'),
        (_re.compile(r'\\left\s*\\Vert([^V]+)\\right\s*\\Vert'), r'||\1||'),
        (_re.compile(r'\\norm\{([^}]+)\}'),                 r'||\1||'),
        (_re.compile(r'\\abs\{([^}]+)\}'),                  r'|\1|'),
        # ── Combinaisons / Binôme ──
        (_re.compile(r'\\binom\{([^}]+)\}\{([^}]+)\}'),    r'C(\1,\2)'),
        (_re.compile(r'\\dbinom\{([^}]+)\}\{([^}]+)\}'),   r'C(\1,\2)'),
        (_re.compile(r'\\tbinom\{([^}]+)\}\{([^}]+)\}'),   r'C(\1,\2)'),
        # ── Vecteurs et dérivées ──
        (_re.compile(r'\\overrightarrow\{([^}]+)\}'),  r'\1⃗'),
        (_re.compile(r'\\overleftarrow\{([^}]+)\}'),   r'\1⃖'),
        (_re.compile(r'\\vec\{([^}]+)\}'),             r'\1⃗'),
        (_re.compile(r'\\hat\{([^}]+)\}'),             r'\1̂'),
        (_re.compile(r'\\bar\{([^}]+)\}'),             r'\1̄'),
        (_re.compile(r'\\tilde\{([^}]+)\}'),           r'\1̃'),
        (_re.compile(r'\\dot\{([^}]+)\}'),             r'\1̇'),
        (_re.compile(r'\\ddot\{([^}]+)\}'),            r'\1̈'),
        (_re.compile(r'\\overline\{([^}]+)\}'),        r'\1̄'),
        (_re.compile(r'\\underline\{([^}]+)\}'),       r'\1'),
        (_re.compile(r'\\widehat\{([^}]+)\}'),         r'\1̂'),
        (_re.compile(r'\\widetilde\{([^}]+)\}'),       r'\1̃'),
        # ── Crochets et accolades ──
        (_re.compile(r'\\lfloor'),  '⌊'), (_re.compile(r'\\rfloor'), '⌋'),
        (_re.compile(r'\\lceil'),   '⌈'), (_re.compile(r'\\rceil'),  '⌉'),
        (_re.compile(r'\\langle'),  '⟨'), (_re.compile(r'\\rangle'), '⟩'),
        (_re.compile(r'\\{'),       '{'), (_re.compile(r'\\}'),       '}'),
        # ── Lettres grecques — TOUTES les 24 lettres (minuscules et majuscules) ──
        (_re.compile(r'\\Alpha'),    'Α'), (_re.compile(r'\\alpha'),      'α'),
        (_re.compile(r'\\Beta'),     'Β'), (_re.compile(r'\\beta'),       'β'),
        (_re.compile(r'\\Gamma'),    'Γ'), (_re.compile(r'\\gamma'),      'γ'),
        (_re.compile(r'\\Delta'),    'Δ'), (_re.compile(r'\\delta'),      'δ'),
        (_re.compile(r'\\Epsilon'),  'Ε'), (_re.compile(r'\\(?:var)?epsilon'), 'ε'),
        (_re.compile(r'\\Zeta'),     'Ζ'), (_re.compile(r'\\zeta'),       'ζ'),
        (_re.compile(r'\\Eta'),      'Η'), (_re.compile(r'\\eta'),        'η'),
        (_re.compile(r'\\Theta'),    'Θ'), (_re.compile(r'\\(?:var)?theta'), 'θ'),
        (_re.compile(r'\\Iota'),     'Ι'), (_re.compile(r'\\iota'),       'ι'),
        (_re.compile(r'\\Kappa'),    'Κ'), (_re.compile(r'\\(?:var)?kappa'), 'κ'),
        (_re.compile(r'\\Lambda'),   'Λ'), (_re.compile(r'\\lambda'),     'λ'),
        (_re.compile(r'\\Mu'),       'Μ'), (_re.compile(r'\\mu'),         'μ'),
        (_re.compile(r'\\Nu'),       'Ν'), (_re.compile(r'\\nu'),         'ν'),
        (_re.compile(r'\\Xi'),       'Ξ'), (_re.compile(r'\\xi'),         'ξ'),
        (_re.compile(r'\\Omicron'),  'Ο'), (_re.compile(r'\\omicron'),    'ο'),
        (_re.compile(r'\\Pi'),       'Π'), (_re.compile(r'\\pi'),         'π'),
        (_re.compile(r'\\varpi'),    'ϖ'),
        (_re.compile(r'\\Rho'),      'Ρ'), (_re.compile(r'\\(?:var)?rho'), 'ρ'),
        (_re.compile(r'\\Sigma'),    'Σ'), (_re.compile(r'\\sigma'),      'σ'),
        (_re.compile(r'\\varsigma'), 'ς'),
        (_re.compile(r'\\Tau'),      'Τ'), (_re.compile(r'\\tau'),        'τ'),
        (_re.compile(r'\\Upsilon'),  'Υ'), (_re.compile(r'\\upsilon'),    'υ'),
        (_re.compile(r'\\Phi'),      'Φ'), (_re.compile(r'\\(?:var)?phi'), 'φ'),
        (_re.compile(r'\\Chi'),      'Χ'), (_re.compile(r'\\chi'),        'χ'),
        (_re.compile(r'\\Psi'),      'Ψ'), (_re.compile(r'\\psi'),        'ψ'),
        (_re.compile(r'\\Omega'),    'Ω'), (_re.compile(r'\\omega'),      'ω'),
        # ── Ensembles de nombres ──
        (_re.compile(r'\\mathbb\{R\}'), 'ℝ'), (_re.compile(r'\\mathbb\{N\}'), 'ℕ'),
        (_re.compile(r'\\mathbb\{Z\}'), 'ℤ'), (_re.compile(r'\\mathbb\{Q\}'), 'ℚ'),
        (_re.compile(r'\\mathbb\{C\}'), 'ℂ'), (_re.compile(r'\\mathbb\{P\}'), 'ℙ'),
        (_re.compile(r'\\mathbb\{([^}]+)\}'), r'\1'),  # autres \mathbb
        # ── Opérateurs arithmétiques et relations ──
        (_re.compile(r'\\times'),     '×'),  (_re.compile(r'\\cdot'),     '·'),
        (_re.compile(r'\\div'),       '÷'),  (_re.compile(r'\\pm'),       '±'),
        (_re.compile(r'\\mp'),        '∓'),  (_re.compile(r'\\ast'),      '*'),
        (_re.compile(r'\\star'),      '★'),  (_re.compile(r'\\circ'),     '∘'),
        (_re.compile(r'\\bullet'),    '•'),  (_re.compile(r'\\ldots'),     '…'),
        (_re.compile(r'\\cdots'),     '⋯'),  (_re.compile(r'\\vdots'),     '⋮'),
        (_re.compile(r'\\ddots'),     '⋱'),
        # ── Relations d'ordre ──
        (_re.compile(r'\\leq?'),      '≤'),  (_re.compile(r'\\geq?'),     '≥'),
        (_re.compile(r'\\ll'),        '≪'),  (_re.compile(r'\\gg'),       '≫'),
        (_re.compile(r'\\neq'),       '≠'),  (_re.compile(r'\\approx'),   '≈'),
        (_re.compile(r'\\equiv'),     '≡'),  (_re.compile(r'\\propto'),   '∝'),
        (_re.compile(r'\\sim'),       '~'),  (_re.compile(r'\\simeq'),    '≃'),
        (_re.compile(r'\\cong'),      '≅'),  (_re.compile(r'\\doteq'),    '≐'),
        (_re.compile(r'\\not\\in'),   '∉'),  (_re.compile(r'\\not\\subset'),'⊄'),
        (_re.compile(r'\\not='),      '≠'),  (_re.compile(r'\\not\\eq'),  '≠'),
        # ── Symboles maths avancés ──
        (_re.compile(r'\\infty'),     '∞'),
        (_re.compile(r'\\partial'),   '∂'),  (_re.compile(r'\\nabla'),    '∇'),
        (_re.compile(r'\\forall'),    '∀'),  (_re.compile(r'\\exists'),   '∃'),
        (_re.compile(r'\\nexists'),   '∄'),
        (_re.compile(r'\\emptyset'),  '∅'),  (_re.compile(r'\\varnothing'),'∅'),
        (_re.compile(r'\\aleph'),     'ℵ'),  (_re.compile(r'\\hbar'),     'ℏ'),
        (_re.compile(r'\\ell'),       'ℓ'),  (_re.compile(r'\\wp'),       '℘'),
        (_re.compile(r'\\Re'),        'ℜ'),  (_re.compile(r'\\Im'),       'ℑ'),
        # ── Intégrales et sommes ──
        (_re.compile(r'\\int'),       '∫'),  (_re.compile(r'\\iint'),     '∬'),
        (_re.compile(r'\\iiint'),     '∭'),  (_re.compile(r'\\oint'),     '∮'),
        (_re.compile(r'\\sum'),       'Σ'),  (_re.compile(r'\\prod'),     'Π'),
        (_re.compile(r'\\coprod'),    '∐'),
        # ── Logique et ensembles ──
        (_re.compile(r'\\in\b'),      '∈'),  (_re.compile(r'\\notin'),    '∉'),
        (_re.compile(r'\\subset'),    '⊂'),  (_re.compile(r'\\supset'),   '⊃'),
        (_re.compile(r'\\subseteq'),  '⊆'),  (_re.compile(r'\\supseteq'), '⊇'),
        (_re.compile(r'\\nsubset'),   '⊄'),  (_re.compile(r'\\nsupset'),  '⊅'),
        (_re.compile(r'\\cup'),       '∪'),  (_re.compile(r'\\cap'),      '∩'),
        (_re.compile(r'\\setminus'),  '∖'),  (_re.compile(r'\\complement'),'∁'),
        (_re.compile(r'\\land'),      '∧'),  (_re.compile(r'\\lor'),      '∨'),
        (_re.compile(r'\\lnot'),      '¬'),  (_re.compile(r'\\neg'),      '¬'),
        (_re.compile(r'\\oplus'),     '⊕'),  (_re.compile(r'\\otimes'),   '⊗'),
        (_re.compile(r'\\odot'),      '⊙'),
        # ── Géométrie ──
        (_re.compile(r'\\angle'),     '∠'),  (_re.compile(r'\\measuredangle'),'∡'),
        (_re.compile(r'\\perp'),      '⊥'),  (_re.compile(r'\\parallel'), '∥'),
        (_re.compile(r'\\triangle'),  '△'),  (_re.compile(r'\\square'),   '□'),
        (_re.compile(r'\\diamond'),   '◇'),  (_re.compile(r'\\circ'),     '°'),
        # ── Physique : unités et constantes ──
        (_re.compile(r'\\Omega\b'),   'Ω'),  # ohm (déjà dans grec, redondant mais sûr)
        (_re.compile(r'\\degree'),    '°'),
        (_re.compile(r'\\celsius'),   '°C'),
        # ── Chimie ──
        (_re.compile(r'\\rightleftharpoons'), '⇌'),  # équilibre chimique
        (_re.compile(r'\\longrightarrow'),    '⟶'),  # flèche réaction
        (_re.compile(r'\\xlongrightarrow\{([^}]+)\}'), r'—\1→'),
        # ── Flèches ──
        (_re.compile(r'\\Leftrightarrow'),  '⟺'),  (_re.compile(r'\\iff'), '⟺'),
        (_re.compile(r'\\Rightarrow'),      '⟹'),  (_re.compile(r'\\implies'),'⟹'),
        (_re.compile(r'\\Leftarrow'),       '⟸'),
        (_re.compile(r'\\rightarrow'),      '→'),  (_re.compile(r'\\to'), '→'),
        (_re.compile(r'\\leftarrow'),       '←'),  (_re.compile(r'\\gets'),'←'),
        (_re.compile(r'\\leftrightarrow'),  '↔'),
        (_re.compile(r'\\uparrow'),         '↑'),  (_re.compile(r'\\downarrow'),'↓'),
        (_re.compile(r'\\Uparrow'),         '⇑'),  (_re.compile(r'\\Downarrow'),'⇓'),
        (_re.compile(r'\\nearrow'),         '↗'),  (_re.compile(r'\\searrow'),'↘'),
        (_re.compile(r'\\swarrow'),         '↙'),  (_re.compile(r'\\nwarrow'),'↖'),
        (_re.compile(r'\\mapsto'),          '↦'),
        (_re.compile(r'\\longmapsto'),      '⟼'),
        # ── Modulo et divisibilité ──
        (_re.compile(r'\\pmod\{([^}]+)\}'), r' (mod \1)'),
        (_re.compile(r'\\bmod\b'),          'mod'),
        (_re.compile(r'\\mod\b'),           'mod'),
        (_re.compile(r'\\gcd\b'),           'pgcd'),
        (_re.compile(r'\\lcm\b'),           'ppcm'),
        # ── Fonctions maths et physique ──
        (_re.compile(r'\\arcsin\b'),  'arcsin'), (_re.compile(r'\\arccos\b'), 'arccos'),
        (_re.compile(r'\\arctan\b'),  'arctan'), (_re.compile(r'\\arccot\b'), 'arccot'),
        (_re.compile(r'\\sin\b'),     'sin'),    (_re.compile(r'\\cos\b'),    'cos'),
        (_re.compile(r'\\tan\b'),     'tan'),    (_re.compile(r'\\cot\b'),    'cot'),
        (_re.compile(r'\\sec\b'),     'sec'),    (_re.compile(r'\\csc\b'),    'cosec'),
        (_re.compile(r'\\sinh\b'),    'sinh'),   (_re.compile(r'\\cosh\b'),   'cosh'),
        (_re.compile(r'\\tanh\b'),    'tanh'),   (_re.compile(r'\\coth\b'),   'coth'),
        (_re.compile(r'\\ln\b'),      'ln'),     (_re.compile(r'\\log\b'),    'log'),
        (_re.compile(r'\\exp\b'),     'exp'),    (_re.compile(r'\\lim\b'),    'lim'),
        (_re.compile(r'\\max\b'),     'max'),    (_re.compile(r'\\min\b'),    'min'),
        (_re.compile(r'\\inf\b'),     'inf'),    (_re.compile(r'\\sup\b'),    'sup'),
        (_re.compile(r'\\det\b'),     'det'),    (_re.compile(r'\\ker\b'),    'ker'),
        (_re.compile(r'\\dim\b'),     'dim'),    (_re.compile(r'\\rank\b'),   'rang'),
        (_re.compile(r'\\tr\b'),      'tr'),     (_re.compile(r'\\grad\b'),   'grad'),
        (_re.compile(r'\\div\b'),     'div'),    (_re.compile(r'\\rot\b'),    'rot'),
        (_re.compile(r'\\curl\b'),    'rot'),
        # ── Délimiteurs auto-sizing (ignorés, juste les parenthèses restent) ──
        (_re.compile(r'\\[Bb]ig[glr]?\s*[\(\[|<]'),  '('),
        (_re.compile(r'\\[Bb]ig[glr]?\s*[\)\]|>]'),  ')'),
        (_re.compile(r'\\left\s*[\(\[|<]'),   '('),
        (_re.compile(r'\\right\s*[\)\]|>]'),  ')'),
        (_re.compile(r'\\left\s*\\{'),        '{'),
        (_re.compile(r'\\right\s*\\}'),       '}'),
        (_re.compile(r'\\left\.'),            ''),
        (_re.compile(r'\\right\.'),           ''),
        # ── Espaces et mise en page LaTeX ──
        (_re.compile(r'\\[,;:!]'),     ' '),
        (_re.compile(r'\\quad'),       '  '),
        (_re.compile(r'\\qquad'),      '   '),
        (_re.compile(r'\\noindent'),   ''),
        (_re.compile(r'\\newline'),    '\n'),
        (_re.compile(r'\\\\'),         ' '),  # fin de ligne LaTeX
        # ── Environnements ──
        (_re.compile(r'\\begin\{[^}]+\}'),   ''),
        (_re.compile(r'\\end\{[^}]+\}'),     ''),
        (_re.compile(r'\\item\b'),           '• '),
        # ── Nettoyage commandes texte ──
        (_re.compile(r'\\(?:text|mathrm|mathbf|mathit|mathsf|mathcal|mathscr|mathfrak|mathbb)\{([^}]+)\}'), r'\1'),
        (_re.compile(r'\\(?:boldsymbol|pmb)\{([^}]+)\}'), r'\1'),
        (_re.compile(r'\\(?:underbrace|overbrace)\{([^}]+)\}(?:_\{[^}]+\})?'), r'\1'),
        (_re.compile(r'\\(?:stackrel|overset|underset)\{[^}]+\}\{([^}]+)\}'), r'\1'),
        (_re.compile(r'\\(?:color|textcolor)\{[^}]+\}\{([^}]+)\}'), r'\1'),
        (_re.compile(r'\\label\{[^}]+\}'),   ''),
        (_re.compile(r'\\tag\{[^}]+\}'),     ''),
        (_re.compile(r'\\ref\{[^}]+\}'),     '?'),
        # ── Retrait des backslashes restants isolés ──
        (_re.compile(r'\\([A-Za-z]+)'),  r'\1'),  # \commande → commande (dernier recours)
    ]

    def nettoyer_latex_complet(texte):
        """Convertit LaTeX complet + notation Nova ^{} _{} en texte normalisé."""

        # 0. Pré-nettoyage des formules malformées produites par Gemini
        # ex: __{} __{}  _{  }  {50 } {0,5 }
        texte = _re.sub(r'_{2,}\{\s*\}', '',  texte)   # __{}  ___{}  → supprimé
        texte = _re.sub(r'\^{2,}\{\s*\}', '', texte)   # ^^{}         → supprimé
        texte = _re.sub(r'[_^]\{\s*\}',   '',  texte)  # _{}   ^{}    → supprimé
        texte = _re.sub(r'_{2,}([^{])',  r'_\1', texte)  # __x   → _x
        texte = _re.sub(r'\^{2,}([^{])', r'^\1', texte)  # ^^x   → ^x
        # 1. Blocs $$...$$ et formules $...$
        def conv_dollar(m):
            f = m.group(1)
            for pat, repl in LATEX_TO_NOVA:
                try:
                    f = pat.sub(repl, f) if callable(repl) else pat.sub(repl, f)
                except Exception:
                    pass
            return f
        texte = _re.sub(r'\$\$([^$]+)\$\$', conv_dollar, texte)
        texte = _re.sub(r'\$([^$]+)\$',     conv_dollar, texte)
        # 2. Commandes LaTeX hors dollars
        for pat, repl in LATEX_TO_NOVA:
            try:
                texte = pat.sub(repl, texte) if not callable(repl) else pat.sub(repl, texte)
            except Exception:
                pass
        # 3. Nettoyer accolades orphelines (pas précédées de ^ ou _)
        result = []
        open_stack = []
        for c in texte:
            if c == '{':
                prev = result[-1] if result else ''
                if prev in ('^', '_'):
                    result.append(c)
                    open_stack.append(False)   # accolade utile
                else:
                    open_stack.append(True)    # accolade orpheline → ignorer
            elif c == '}':
                if open_stack and open_stack[-1]:
                    open_stack.pop()           # ferme orpheline, on supprime
                elif open_stack:
                    open_stack.pop()
                    result.append(c)
                else:
                    result.append(c)
            else:
                result.append(c)
        return ''.join(result)

    def ajouter_formule_dans_run(p, texte, bold=False, size=11, color=None):
        """
        Crée des runs Word avec vrais exposants/indices, symboles Unicode, gras.

        Syntaxe reconnue :
          ^{expr}       → exposant Word (superscript)
          ^x            → exposant 1 char
          _{expr}       → indice Word (subscript)
          _x            → indice 1 char
          **texte**     → gras
          √(expr)       → symbole racine + texte
          (a)/(b)       → fraction affichée avec barre / centrée
          Tous Unicode (α β γ Σ ∫ ∞ ≤ ≥ × → ⇌ ∂ ∇ …) passent directement

        Note : les fractions (a)/(b) issues de \frac sont rendues en ligne
        avec les parenthèses pour la lisibilité.
        """
        FORMULE_RE = _re.compile(
            r'(\*\*[^*]+\*\*)'                      # **gras**
            r'|\^{([^}]+)}'                          # exposant long  ^{abc}
            r'|\^([^\s{_^*/+\-=(),\[\]])'            # exposant court ^x (1 car non-espace)
            r'|_{([^}]+)}'                           # indice long    _{abc}
            r'|_([^\s{_^*/+\-=(),\[\]])'             # indice court   _x (1 car non-espace)
        )

        def _run(text, sup=False, sub=False, bd=False, sz=None):
            text = sanitize_xml(text)
            if not text:
                return None
            r = p.add_run(text)
            r.font.name  = "Arial"
            r.font.size  = Pt(sz if sz else (max(7, size - 2) if (sup or sub) else size))
            r.bold       = bd
            if sup: r.font.superscript = True
            if sub: r.font.subscript   = True
            if color: r.font.color.rgb = RC(*color)
            return r

        last = 0
        for m in FORMULE_RE.finditer(texte):
            if m.start() > last:
                _run(texte[last:m.start()], bd=bold)
            if m.group(1):                        # **gras**
                _run(m.group(1)[2:-2], bd=True)
            elif m.group(2) or m.group(3):        # exposant
                _run(m.group(2) or m.group(3), sup=True, bd=bold)
            elif m.group(4) or m.group(5):        # indice
                _run(m.group(4) or m.group(5), sub=True, bd=bold)
            last = m.end()
        if last < len(texte):
            _run(texte[last:], bd=bold)


    lignes = contenu.split("\n")
    i = 0
    sauts_de_page_count = 0  # Compteur de sauts de page pour détecter page garde + sommaire

    while i < len(lignes):
        l = lignes[i].rstrip()

        # ── SAUT DE PAGE NOVA — VRAI SAUT DE PAGE WORD ────────────
        if l.strip() == "---SAUT_DE_PAGE---":
            from docx.oxml.ns import qn as _qn
            from docx.oxml import OxmlElement as _OE
            sauts_de_page_count += 1
            p_break = doc.add_paragraph()
            p_break.paragraph_format.space_before = Pt(0)
            p_break.paragraph_format.space_after  = Pt(0)
            run_break = p_break.add_run()
            br = _OE("w:br")
            br.set(_qn("w:type"), "page")
            run_break._r.append(br)
            i += 1
            continue
        # ── MARQUEUR TITRE ROUGE — Grand titre centré rouge ─────
        if l.strip().startswith("###TITRE_ROUGE###"):
            texte_titre = l.strip().replace("###TITRE_ROUGE###", "").strip()
            p_rouge = doc.add_paragraph()
            p_rouge.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_rouge.paragraph_format.space_before = Pt(12)
            p_rouge.paragraph_format.space_after  = Pt(12)
            run_rouge = p_rouge.add_run(texte_titre)
            run_rouge.bold = True
            run_rouge.font.name = "Arial"
            run_rouge.font.size = Pt(28)
            run_rouge.font.color.rgb = RC(0xC0, 0x00, 0x00)  # Rouge vif
            i += 1
            continue

        # ── MARQUEUR ESPACE — Grand espace vertical ───────────────
        if l.strip() == "###ESPACE###":
            p_esp = doc.add_paragraph()
            p_esp.paragraph_format.space_before = Pt(0)
            p_esp.paragraph_format.space_after  = Pt(0)
            p_esp.paragraph_format.line_spacing = Pt(36)  # ~1.2cm d'espace
            i += 1
            continue

        # ── BLOC FORMULE NOVA — Formule centrée sur fond bleu clair ──
        # Syntaxe : ###FORMULE### F = m × a   ou   ###FORMULE### E = mc^{2}
        if l.strip().startswith("###FORMULE###"):
            texte_f = l.strip().replace("###FORMULE###", "").strip()
            texte_f = nettoyer_latex_complet(texte_f)
            from docx.oxml import OxmlElement as _OEf
            from docx.oxml.ns import qn as _qnf
            p_f = doc.add_paragraph()
            p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_f.paragraph_format.space_before = Pt(6)
            p_f.paragraph_format.space_after  = Pt(6)
            # Fond bleu très clair
            pPr_f = p_f._p.get_or_add_pPr()
            shd_f = _OEf("w:shd")
            shd_f.set(_qnf("w:val"), "clear")
            shd_f.set(_qnf("w:color"), "auto")
            shd_f.set(_qnf("w:fill"), "D6E4F0")  # bleu clair
            pPr_f.append(shd_f)
            # Bordure fine bleue autour
            pBdr_f = _OEf("w:pBdr")
            for side_name in ["top", "bottom", "left", "right"]:
                side_el = _OEf(f"w:{side_name}")
                side_el.set(_qnf("w:val"), "single")
                side_el.set(_qnf("w:sz"), "6")
                side_el.set(_qnf("w:space"), "4")
                side_el.set(_qnf("w:color"), "2E75B6")
                pBdr_f.append(side_el)
            pPr_f.append(pBdr_f)
            # Texte formule en Arial 13pt gras bleu foncé avec exposants réels
            ajouter_formule_dans_run(p_f, texte_f, bold=True, size=13,
                                     color=(0x1F, 0x4E, 0x79))
            i += 1
            continue

        # ── BLOC FORMULE MULTILIGNE — ###DEBUT_FORMULES### ... ###FIN_FORMULES###
        if l.strip() == "###DEBUT_FORMULES###":
            i += 1
            while i < len(lignes) and lignes[i].strip() != "###FIN_FORMULES###":
                lf = lignes[i].strip()
                if lf:
                    lf_clean = nettoyer_latex_complet(lf)
                    from docx.oxml import OxmlElement as _OEml
                    from docx.oxml.ns import qn as _qnml
                    p_ml = doc.add_paragraph()
                    p_ml.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_ml.paragraph_format.space_before = Pt(2)
                    p_ml.paragraph_format.space_after  = Pt(2)
                    pPr_ml = p_ml._p.get_or_add_pPr()
                    shd_ml = _OEml("w:shd")
                    shd_ml.set(_qnml("w:val"), "clear")
                    shd_ml.set(_qnml("w:color"), "auto")
                    shd_ml.set(_qnml("w:fill"), "EEF3FA")
                    pPr_ml.append(shd_ml)
                    ajouter_formule_dans_run(p_ml, lf_clean, bold=True, size=12,
                                             color=(0x1F, 0x4E, 0x79))
                i += 1
            i += 1  # sauter ###FIN_FORMULES###
            continue

        # ── LIGNES DE SÉPARATION ════ ET ──── ─────────────────────
        if l.strip().startswith("════") or l.strip().startswith("━━━━"):
            # Ne pas dessiner le trait si la prochaine ligne non-vide est un saut de page
            next_content = next((lignes[j].strip() for j in range(i+1, len(lignes)) if lignes[j].strip()), "")
            if next_content == "---SAUT_DE_PAGE---":
                i += 1
                continue  # Ignorer ce trait — il serait au bas de page et créerait un espace vide
            p_line = doc.add_paragraph()
            p_line.paragraph_format.space_before = Pt(4)
            p_line.paragraph_format.space_after  = Pt(4)
            p_line.paragraph_format.line_spacing = Pt(1)
            pPr2 = p_line._p.get_or_add_pPr()
            pBdr2 = OxmlElement("w:pBdr")
            bot2 = OxmlElement("w:bottom")
            bot2.set(qn("w:val"), "single")
            bot2.set(qn("w:sz"), "12")
            bot2.set(qn("w:space"), "1")
            bot2.set(qn("w:color"), "1F4E79")
            pBdr2.append(bot2)
            pPr2.append(pBdr2)
            i += 1
            continue

        if l.strip().startswith("────") or l.strip().startswith("----"):
            p_line = doc.add_paragraph()
            p_line.paragraph_format.space_before = Pt(3)
            p_line.paragraph_format.space_after  = Pt(3)
            p_line.paragraph_format.line_spacing = Pt(1)
            pPr2 = p_line._p.get_or_add_pPr()
            pBdr2 = OxmlElement("w:pBdr")
            bot2 = OxmlElement("w:bottom")
            bot2.set(qn("w:val"), "single")
            bot2.set(qn("w:sz"), "4")
            bot2.set(qn("w:space"), "1")
            bot2.set(qn("w:color"), "AAAAAA")
            pBdr2.append(bot2)
            pPr2.append(pBdr2)
            i += 1
            continue

        if l.strip() in ["---", "***", "___", "*"]:
            doc.add_paragraph("")
            i += 1
            continue

        if l.startswith("#### "):
            p = doc.add_heading(l[5:].strip(), level=4)
            i += 1
            continue
        if l.startswith("### "):
            p = doc.add_heading(l[4:].strip(), level=3)
            i += 1
            continue
        if l.startswith("## "):
            p = doc.add_heading(l[3:].strip(), level=2)
            i += 1
            continue
        if l.startswith("# "):
            # Ignorer les lignes de commentaires de structure Nova
            if l.startswith("# ═") or l.startswith("# #") or l.startswith("# ─"):
                i += 1
                continue
            p = doc.add_heading(l[2:].strip(), level=1)
            i += 1
            continue

        if l.startswith("|") and l.endswith("|"):
            table_lines = []
            while i < len(lignes) and lignes[i].strip().startswith("|") and lignes[i].strip().endswith("|"):
                row = lignes[i].strip()
                if not re.match(r"^[\|\s\-:]+$", row):
                    cells = [c.strip() for c in row.strip("|").split("|")]
                    table_lines.append(cells)
                i += 1

            if table_lines:
                from docx.shared import Inches
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement

                n_cols = max(len(r) for r in table_lines)
                col_widths_map = {
                    2: [3.0, 6.0],
                    3: [1.0, 7.5, 2.5],
                    4: [1.0, 5.0, 2.5, 2.5],
                    5: [0.8, 5.0, 1.5, 1.5, 1.2],
                }
                col_widths = col_widths_map.get(n_cols, [9.0/n_cols]*n_cols)

                from docx.shared import Cm as DocxCm
                table = doc.add_table(rows=0, cols=n_cols)
                table.style = "Table Grid"

                for r_idx, row_data in enumerate(table_lines):
                    row_obj = table.add_row()
                    is_header = (r_idx == 0)
                    row_obj.height = DocxCm(0.9 if is_header else 1.5)
                    from docx.oxml.ns import qn as _qn
                    from docx.oxml import OxmlElement as _OE
                    trPr = row_obj._tr.get_or_add_trPr()
                    trHeight = _OE("w:trHeight")
                    trHeight.set(_qn("w:val"), str(int((0.9 if is_header else 1.5) * 567)))
                    trHeight.set(_qn("w:hRule"), "exact")
                    trPr.append(trHeight)

                    for c_idx, cell_text in enumerate(row_data):
                        cell = row_obj.cells[c_idx]
                        if c_idx < len(col_widths):
                            cell.width = DocxCm(col_widths[c_idx])
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        tcMar = _OE("w:tcMar")
                        for side in ["top","bottom","left","right"]:
                            m = _OE(f"w:{side}")
                            m.set(_qn("w:w"), "120")
                            m.set(_qn("w:type"), "dxa")
                            tcMar.append(m)
                        tcPr.append(tcMar)
                        if is_header:
                            shd = _OE("w:shd")
                            shd.set(_qn("w:val"), "clear")
                            shd.set(_qn("w:color"), "auto")
                            shd.set(_qn("w:fill"), "1F4E79")
                            tcPr.append(shd)
                        elif r_idx % 2 == 0:
                            shd = _OE("w:shd")
                            shd.set(_qn("w:val"), "clear")
                            shd.set(_qn("w:color"), "auto")
                            shd.set(_qn("w:fill"), "EEF3FA")
                            tcPr.append(shd)

                        para = cell.paragraphs[0]
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx in [0, n_cols-1] else WD_ALIGN_PARAGRAPH.LEFT
                        para.paragraph_format.space_before = Pt(2)
                        para.paragraph_format.space_after = Pt(2)
                        # Nettoyer le LaTeX dans les cellules (ex: $\text{kg}$, $m^3$...)
                        cell_text_clean = nettoyer_latex_complet(cell_text)
                        ajouter_formule_dans_run(
                            para, cell_text_clean,
                            bold=is_header,
                            size=10,
                            color=(0xFF, 0xFF, 0xFF) if is_header else None
                        )

                doc.add_paragraph("")
            continue

        m_num = re.match(r"^(\d+)[.)]\s+(.*)", l)
        if m_num:
            p = doc.add_paragraph(style="List Number")
            parts = re.split(r"(\*\*[^*]+\*\*)", m_num.group(2))
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2]); run.bold = True
                else:
                    run = p.add_run(part.replace("*","").replace("`",""))
            for run in p.runs:
                run.font.name = "Arial"; run.font.size = Pt(11)
            i += 1
            continue

        if re.match(r"^[\-\*\•]\s+", l):
            texte = re.sub(r"^[\-\*\•]\s+", "", l)
            p = doc.add_paragraph(style="List Bullet")
            parts = re.split(r"(\*\*[^*]+\*\*)", texte)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2]); run.bold = True
                else:
                    run = p.add_run(part.replace("*","").replace("`",""))
            for run in p.runs:
                run.font.name = "Arial"; run.font.size = Pt(11)
            i += 1
            continue

        if not l.strip():
            p_vide = doc.add_paragraph()
            # Dans page de garde (avant 1er saut) et sommaire (avant 2e saut) : espacement réduit
            if sauts_de_page_count < 2:
                p_vide.paragraph_format.space_before = Pt(0)
                p_vide.paragraph_format.space_after  = Pt(0)
                p_vide.paragraph_format.line_spacing = Pt(6)
            else:
                # Corps : ligne vide réduite pour éviter les orphelines
                p_vide.paragraph_format.space_before = Pt(0)
                p_vide.paragraph_format.space_after  = Pt(0)
                p_vide.paragraph_format.line_spacing = Pt(8)
            i += 1
            continue

        if l.strip().startswith("**") and l.strip().endswith("**") and l.strip().count("**") == 2:
            texte = l.strip()[2:-2]
            p = doc.add_paragraph()
            run = p.add_run(texte)
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(12)
            run.font.color.rgb = RC(0x1F, 0x4E, 0x79)
            i += 1
            continue

        p = add_formatted_para(doc, l.strip())
        # Mode compact pour page de garde et sommaire
        if sauts_de_page_count < 2:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(3)
        else:
            # Corps du document : espacement serré pour éviter les lignes orphelines
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(6)
        i += 1

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def creer_xlsx(description, client_nom):
    """
    Génère un Excel dynamique basé sur le JSON retourné par Gemini.
    Si le JSON est invalide, fallback sur un template générique.
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    import json, re

    # ── PALETTE DE COULEURS ────────────────────────────────────────
    COULEURS = {
        "bleu":   {"bg": "1F4E79", "fg": "FFFFFF"},
        "vert":   {"bg": "1E8449", "fg": "FFFFFF"},
        "orange": {"bg": "D35400", "fg": "FFFFFF"},
        "rouge":  {"bg": "C0392B", "fg": "FFFFFF"},
        "violet": {"bg": "7D3C98", "fg": "FFFFFF"},
        "gris":   {"bg": "5D6D7E", "fg": "FFFFFF"},
        "cyan":   {"bg": "117A65", "fg": "FFFFFF"},
        "or":     {"bg": "B7950B", "fg": "FFFFFF"},
    }
    BLEU_FONCE = "1F4E79"
    BLEU_MOY   = "2E75B6"
    BLEU_CLAIR = "D6E4F0"
    BLANC      = "FFFFFF"
    GRIS_CLAIR = "F2F2F2"
    GRIS_MED   = "E8E8E8"

    def hdr(cell, bg=BLEU_FONCE, fg=BLANC, bold=True, size=11):
        cell.font      = Font(bold=bold, color=fg, name="Arial", size=size)
        cell.fill      = PatternFill("solid", start_color=bg)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    def brd(cell, color="CCCCCC", style="thin"):
        s = Side(style=style, color=color)
        cell.border = Border(top=s, bottom=s, left=s, right=s)

    def brd_epais(cell):
        s_ext = Side(style="medium", color="1F4E79")
        s_int = Side(style="thin",   color="CCCCCC")
        cell.border = Border(top=s_ext, bottom=s_ext, left=s_ext, right=s_ext)

    def fmt_cell(cell, type_col, valeur=None):
        """Applique le format nombre/monnaie/date/pourcentage selon le type."""
        if type_col == "monnaie":
            cell.number_format = '#,##0 "FCFA"'
        elif type_col == "pourcentage":
            cell.number_format = "0.0%"
        elif type_col == "nombre":
            cell.number_format = "#,##0"
        elif type_col == "date":
            cell.number_format = "DD/MM/YYYY"

    # ── PARSE JSON GEMINI ──────────────────────────────────────────
    data = None
    try:
        # Nettoyer le texte : enlever balises markdown si présentes
        texte = description.strip()
        texte = re.sub(r"^```json\s*", "", texte)
        texte = re.sub(r"```\s*$", "", texte)
        texte = re.sub(r"^```\s*", "", texte)
        data = json.loads(texte)
    except Exception:
        data = None

    wb = Workbook()
    wb.remove(wb.active)  # Supprimer feuille vide par défaut

    if not data or "feuilles" not in data:
        # ── FALLBACK : template générique si JSON invalide ─────────
        ws = wb.create_sheet("Données")
        ws.sheet_view.showGridLines = False
        ws.merge_cells("A1:D1")
        ws["A1"].value = f"{client_nom} — Données"
        hdr(ws["A1"], size=13); ws.row_dimensions[1].height = 32
        ws["A2"].value = description[:200]
        ws["A2"].font = Font(italic=True, color="7F7F7F", name="Arial", size=10)
        buf = BytesIO(); wb.save(buf); buf.seek(0)
        return buf

    titre_classeur = data.get("titre", f"Classeur {client_nom}")

    # ── CONSTRUCTION DE CHAQUE FEUILLE ─────────────────────────────
    for feuille in data.get("feuilles", []):
        nom_feuille = feuille.get("nom", "Feuille")[:31]
        type_feuille = feuille.get("type", "saisie")
        colonnes = feuille.get("colonnes", [])
        lignes   = feuille.get("lignes_exemple", [])
        kpis     = feuille.get("kpis", [])

        ws = wb.create_sheet(nom_feuille)
        ws.sheet_view.showGridLines = False

        n_cols = max(len(colonnes), 4)
        last_col = get_column_letter(n_cols)

        # ── EN-TÊTE PRINCIPAL ──────────────────────────────────────
        ws.merge_cells(f"A1:{last_col}1")
        cell_titre = ws.cell(row=1, column=1, value=f"{titre_classeur}  |  {client_nom}")
        hdr(cell_titre, size=13); ws.row_dimensions[1].height = 36

        ws.merge_cells(f"A2:{last_col}2")
        cell_desc = ws.cell(row=2, column=1, value=f"{feuille.get('description', '')}  —  Généré le {datetime.now().strftime('%d/%m/%Y')}")
        cell_desc.font      = Font(italic=True, color="7F7F7F", name="Arial", size=10)
        cell_desc.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[2].height = 20

        # ── FEUILLE DE TYPE SAISIE ─────────────────────────────────
        if type_feuille == "saisie" and colonnes:
            # En-têtes colonnes
            for c_idx, col in enumerate(colonnes, 1):
                cell = ws.cell(row=3, column=c_idx, value=col["entete"])
                hdr(cell, bg=BLEU_MOY); brd(cell)
                ws.column_dimensions[get_column_letter(c_idx)].width = col.get("largeur", 18)
            ws.row_dimensions[3].height = 28

            # Lignes de données exemple
            for r_idx, ligne in enumerate(lignes, 4):
                bg = GRIS_CLAIR if r_idx % 2 == 0 else BLANC
                for c_idx, val in enumerate(ligne, 1):
                    if c_idx > len(colonnes):
                        break
                    cell = ws.cell(row=r_idx, column=c_idx, value=val)
                    type_col = colonnes[c_idx-1].get("type", "texte")
                    cell.font      = Font(name="Arial", size=10,
                                          bold=(type_col in ["monnaie","nombre"]),
                                          color=("1F4E79" if type_col == "monnaie" else "000000"))
                    cell.fill      = PatternFill("solid", start_color=bg)
                    cell.alignment = Alignment(vertical="center",
                                               horizontal="center" if type_col in ["monnaie","nombre","date","pourcentage"] else "left")
                    fmt_cell(cell, type_col)
                    brd(cell)

            # Ligne TOTAL
            if lignes:
                total_row = len(lignes) + 4
                ws.row_dimensions[total_row].height = 28
                # Trouver la 1re colonne numérique pour placer le label TOTAL
                first_num_col = next((c_idx for c_idx, col in enumerate(colonnes, 1)
                                      if col.get("type") in ["monnaie", "nombre"]), None)
                # Colonnes texte → label "TOTAL"
                for c_idx, col in enumerate(colonnes, 1):
                    cell_t = ws.cell(row=total_row, column=c_idx)
                    if col.get("type") not in ["monnaie", "nombre"]:
                        if c_idx == 1:
                            cell_t.value = "TOTAL"
                        hdr(cell_t, size=11); brd_epais(cell_t)
                    else:
                        col_letter = get_column_letter(c_idx)
                        cell_t.value = f"=SUM({col_letter}4:{col_letter}{total_row-1})"
                        fmt_cell(cell_t, col["type"])
                        hdr(cell_t, size=11); brd_epais(cell_t)

            # Figer la ligne d'en-tête
            ws.freeze_panes = "A4"

        # ── FEUILLE DE TYPE BILAN / KPIs ───────────────────────────
        elif type_feuille == "bilan" and kpis:
            # Titre section KPIs
            ws.merge_cells(f"A3:{last_col}3")
            cell_kpi_title = ws.cell(row=3, column=1, value="━━  INDICATEURS CLÉS DE PERFORMANCE  ━━")
            hdr(cell_kpi_title, bg=BLEU_FONCE, size=12); ws.row_dimensions[3].height = 30

            # Disposition des KPIs : 2 par ligne (label | valeur | espace | label | valeur)
            ws.column_dimensions["A"].width = 30
            ws.column_dimensions["B"].width = 24
            ws.column_dimensions["C"].width = 3
            ws.column_dimensions["D"].width = 30
            ws.column_dimensions["E"].width = 24

            row_kpi = 4
            for idx, kpi in enumerate(kpis):
                if idx % 2 == 0 and idx > 0:
                    row_kpi += 2  # Saut d'une ligne entre paires

                col_start = 1 if idx % 2 == 0 else 4
                couleur_key = kpi.get("couleur", "bleu")
                bg_kpi = COULEURS.get(couleur_key, COULEURS["bleu"])["bg"]
                fg_kpi = COULEURS.get(couleur_key, COULEURS["bleu"])["fg"]

                # Label
                cl = ws.cell(row=row_kpi, column=col_start, value=kpi["label"])
                cl.font      = Font(bold=True, name="Arial", size=11, color=fg_kpi)
                cl.fill      = PatternFill("solid", start_color=bg_kpi)
                cl.alignment = Alignment(horizontal="center", vertical="center")
                brd_epais(cl)
                ws.row_dimensions[row_kpi].height = 36

                # Valeur
                cv = ws.cell(row=row_kpi, column=col_start+1, value=kpi.get("formule", 0))
                cv.font      = Font(bold=True, name="Arial", size=14, color=fg_kpi)
                cv.fill      = PatternFill("solid", start_color=bg_kpi)
                cv.alignment = Alignment(horizontal="center", vertical="center")
                fmt_cell(cv, kpi.get("type", "nombre"))
                brd_epais(cv)

            # ── TABLEAU RÉCAPITULATIF sous les KPIs ───────────────
            row_recap = row_kpi + 3

            # Trouver la 1re feuille de saisie pour le récap
            feuille_saisie = next((f for f in data["feuilles"] if f.get("type") == "saisie"), None)
            if feuille_saisie:
                nom_s   = feuille_saisie["nom"][:31]
                cols_s  = feuille_saisie.get("colonnes", [])

                ws.merge_cells(f"A{row_recap}:E{row_recap}")
                cell_recap_title = ws.cell(row=row_recap, column=1, value=f"RÉCAPITULATIF — {nom_s.upper()}")
                hdr(cell_recap_title, bg=BLEU_MOY, size=12)
                ws.row_dimensions[row_recap].height = 28
                row_recap += 1

                # En-têtes récap
                recap_cols = [c["entete"] for c in cols_s[:5]]
                for c_idx, h in enumerate(recap_cols, 1):
                    cell = ws.cell(row=row_recap, column=c_idx, value=h)
                    hdr(cell, bg=BLEU_FONCE); brd(cell)
                ws.row_dimensions[row_recap].height = 24
                row_recap += 1

                # Lignes récap (depuis les données exemple)
                for r_idx, ligne in enumerate(feuille_saisie.get("lignes_exemple", [])[:8], row_recap):
                    bg = BLEU_CLAIR if r_idx % 2 == 0 else BLANC
                    for c_idx, val in enumerate(ligne[:5], 1):
                        if c_idx > len(cols_s):
                            break
                        cell = ws.cell(row=r_idx, column=c_idx, value=val)
                        type_col = cols_s[c_idx-1].get("type", "texte")
                        cell.font      = Font(name="Arial", size=10,
                                              color=("1F4E79" if type_col == "monnaie" else "000000"))
                        cell.fill      = PatternFill("solid", start_color=bg)
                        cell.alignment = Alignment(vertical="center",
                                                   horizontal="center" if type_col in ["monnaie","nombre","date"] else "left")
                        fmt_cell(cell, type_col)
                        brd(cell)

        # ── AUTRES TYPES DE FEUILLES (générique) ──────────────────
        else:
            if colonnes:
                for c_idx, col in enumerate(colonnes, 1):
                    cell = ws.cell(row=3, column=c_idx, value=col["entete"])
                    hdr(cell); brd(cell)
                    ws.column_dimensions[get_column_letter(c_idx)].width = col.get("largeur", 18)
                for r_idx, ligne in enumerate(lignes, 4):
                    for c_idx, val in enumerate(ligne, 1):
                        cell = ws.cell(row=r_idx, column=c_idx, value=val)
                        cell.font = Font(name="Arial", size=10)
                        cell.alignment = Alignment(vertical="center")
                        brd(cell)

    # ── MISE EN FORME FINALE : onglets colorés ─────────────────────
    couleurs_onglets = ["1F4E79", "1E8449", "D35400", "7D3C98"]
    for i, ws in enumerate(wb.worksheets):
        ws.sheet_properties.tabColor = couleurs_onglets[i % len(couleurs_onglets)]

    buf = BytesIO(); wb.save(buf); buf.seek(0)
    return buf


WHATSAPP_NUMBER = st.secrets.get("WHATSAPP_NUMBER", "2250171542505")
PREMIUM_MSG = "J'aimerais passer à la version Nova Premium pour bénéficier de la puissance 10^10 et de l'IA de pointe."
SUPPORT_MSG = "Bonjour, j'ai besoin d'assistance sur mon espace Nova AI."
whatsapp_premium_url = f"https://wa.me/{WHATSAPP_NUMBER}?text={PREMIUM_MSG.replace(' ', '%20')}"
whatsapp_support_url = f"https://wa.me/{WHATSAPP_NUMBER}?text={SUPPORT_MSG.replace(' ', '%20')}"

if "db" not in st.session_state:
    st.session_state["db"] = load_db()
if "current_user" not in st.session_state:
    st.session_state["current_user"] = None
if "view" not in st.session_state:
    st.session_state["view"] = "home"
if "is_glowing" not in st.session_state:
    st.session_state["is_glowing"] = False
if "show_premium_modal" not in st.session_state:
    st.session_state["show_premium_modal"] = False
if "show_service_warning" not in st.session_state:
    st.session_state["show_service_warning"] = False
if "auto_reply_gratuit" not in st.session_state:
    st.session_state["auto_reply_gratuit"] = get_auto_reply_setting()
if "contenu_fichier_source" not in st.session_state:
    st.session_state["contenu_fichier_source"] = ""
if "last_service_seen" not in st.session_state:
    st.session_state["last_service_seen"] = None
if "warning_triggered" not in st.session_state:
    st.session_state["warning_triggered"] = False
if "intro_played" not in st.session_state:
    st.session_state["intro_played"] = False
if "gemini_results" not in st.session_state:
    st.session_state["gemini_results"] = {}
if "premium_livrable" not in st.session_state:
    st.session_state["premium_livrable"] = None

if st.session_state["current_user"] is None:
    stored_user = st.query_params.get("user_id")
    if stored_user and stored_user in st.session_state["db"]["users"]:
        st.session_state["current_user"] = stored_user
    else:
        components.html("""
            <script>
            var uid = localStorage.getItem('nova_user_id');
            if (uid) {
                var url = new URL(window.location.href);
                url.searchParams.set('user_id', uid);
                window.location.href = url.toString();
            }
            </script>
        """, height=0)

if st.session_state["current_user"]:
    uid_connecte = st.session_state["current_user"]
    components.html(f"""
        <script>
        localStorage.setItem('nova_user_id', '{uid_connecte}');
        </script>
    """, height=0)

def inject_custom_css():
    # ── Détection Premium ─────────────────────────────────────────
    _user = st.session_state.get("current_user")
    _db   = st.session_state.get("db", {})
    _ud   = _db.get("users", {}).get(_user, {}) if _user else {}
    _premium = is_premium_actif(_ud)

    # ── CSS commun (base) ─────────────────────────────────────────
    st.markdown("""
        <style>
        @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@400;600;800&display=swap');
        * { font-family: 'Poppins', sans-serif; }
        .stApp {
            background: #0f0c29;
            background: -webkit-linear-gradient(to right, #24243e, #302b63, #0f0c29);
            background: linear-gradient(to right, #24243e, #302b63, #0f0c29);
            color: #ffffff;
            transition: filter 0.5s ease;
        }
        @keyframes glow-pulse {
            0% { filter: brightness(1) saturate(1); box-shadow: inset 0 0 0px transparent; }
            50% { filter: brightness(1.8) saturate(1.5); box-shadow: inset 0 0 100px rgba(0, 210, 255, 0.5); }
            100% { filter: brightness(1) saturate(1); box-shadow: inset 0 0 0px transparent; }
        }
        .main-title {
            background: linear-gradient(90deg, #00d2ff, #3a7bd5);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            font-weight: 800;
            font-size: 3.5rem !important;
            text-align: center;
            margin-bottom: 20px;
            text-shadow: 0px 0px 20px rgba(0, 210, 255, 0.3);
        }
        .stTabs [data-baseweb="tab-list"] {
            gap: 20px;
            background-color: rgba(255, 255, 255, 0.05);
            padding: 10px;
            border-radius: 15px;
            border: 1px solid rgba(255, 255, 255, 0.1);
        }
        .stTabs [data-baseweb="tab"] {
            height: 60px;
            white-space: pre-wrap;
            background-color: rgba(0, 210, 255, 0.1);
            border-radius: 10px;
            color: white !important;
            font-weight: 700 !important;
            font-size: 1.2rem !important;
            transition: all 0.3s ease;
            border: 1px solid transparent;
            padding: 0 25px;
        }
        .stTabs [data-baseweb="tab"]:nth-child(2) {
            border: 1px solid #2ecc71 !important;
            box-shadow: 0 0 15px rgba(46, 204, 113, 0.2);
            background-color: rgba(46, 204, 113, 0.1);
        }
        .stTabs [data-baseweb="tab"]:hover {
            background-color: rgba(0, 210, 255, 0.3);
            transform: translateY(-2px);
        }
        .stTabs [aria-selected="true"] {
            background-color: rgba(0, 210, 255, 0.6) !important;
            border: 1px solid #00d2ff !important;
            box-shadow: 0 0 20px rgba(0, 210, 255, 0.4);
        }
        @keyframes border-rainbow {
            0% { border-color: #00d2ff; box-shadow: 0 0 10px rgba(0, 210, 255, 0.3); }
            25% { border-color: #3a7bd5; box-shadow: 0 0 10px rgba(58, 123, 213, 0.3); }
            50% { border-color: #FFD700; box-shadow: 0 0 15px rgba(255, 215, 0, 0.3); }
            75% { border-color: #2ecc71; box-shadow: 0 0 10px rgba(46, 204, 113, 0.3); }
            100% { border-color: #00d2ff; box-shadow: 0 0 10px rgba(0, 210, 255, 0.3); }
        }
        .stTextInput label, .stSelectbox label, .stTextArea label {
            color: #00d2ff !important;
            font-weight: 600 !important;
            font-size: 1.1rem !important;
            margin-bottom: 5px;
        }
        div[data-baseweb="input"], div[data-baseweb="select"] > div {
            border: 1px solid rgba(0, 210, 255, 0.3) !important;
            background-color: rgba(0, 0, 0, 0.5) !important;
            color: white !important;
            border-radius: 10px !important;
        }
        .stTextArea textarea {
            background-color: rgba(0, 0, 0, 0.6) !important;
            color: white !important;
            border-radius: 10px !important;
            border: 2px solid #00d2ff !important;
            animation: border-rainbow 4s linear infinite;
            transition: transform 0.3s;
        }
        .stTextArea textarea:focus {
            transform: scale(1.01);
            animation: border-rainbow 1.5s linear infinite;
        }
        .logo-container {
            display: flex;
            justify-content: center;
            align-items: center;
            gap: 30px;
            margin-top: 20px;
            padding: 15px;
            background: rgba(255, 255, 255, 0.03);
            border-radius: 15px;
        }
        .logo-item {
            width: 45px;
            height: 45px;
            filter: grayscale(0.5) opacity(0.7);
            transition: all 0.3s ease;
        }
        .logo-item:hover {
            filter: grayscale(0) opacity(1);
            transform: translateY(-5px) scale(1.1);
        }
        .premium-card {
            background: rgba(20, 20, 30, 0.8);
            border: 2px solid #FFD700;
            border-radius: 20px;
            padding: 25px;
            text-align: center;
            margin-bottom: 30px;
            box-shadow: 0 0 30px rgba(255, 215, 0, 0.2);
            position: relative;
            overflow: hidden;
        }
        .premium-card::before {
            content: "";
            position: absolute;
            top: 0; left: 0; width: 100%; height: 5px;
            background: linear-gradient(90deg, #FFD700, #FF8C00, #FFD700);
        }
        .premium-title {
            color: #FFD700 !important;
            font-size: 1.5rem;
            font-weight: 800;
            text-transform: uppercase;
            margin-bottom: 10px;
            letter-spacing: 1px;
        }
        .premium-desc {
            color: #ffffff !important;
            font-size: 1rem;
            margin-bottom: 20px;
            line-height: 1.5;
        }
        .btn-gold {
            background: linear-gradient(45deg, #FFD700, #FF8C00);
            color: #000 !important;
            padding: 12px 30px;
            border-radius: 50px;
            text-decoration: none;
            font-weight: 800;
            font-size: 1.1rem;
            display: inline-block;
            box-shadow: 0 5px 15px rgba(255, 215, 0, 0.4);
            transition: transform 0.2s, box-shadow 0.2s;
            border: none;
            cursor: pointer;
        }
        .btn-gold:hover {
            transform: scale(1.05);
            box-shadow: 0 8px 25px rgba(255, 215, 0, 0.6);
        }
        .stButton>button {
            border-radius: 12px;
            padding: 0.8rem 2rem;
            background: linear-gradient(90deg, #00d2ff 0%, #3a7bd5 100%);
            border: none;
            color: white !important;
            font-weight: 700;
            font-size: 1.1rem;
            width: 100%;
            margin-top: 10px;
            box-shadow: 0 4px 10px rgba(0, 210, 255, 0.3);
            transition: 0.3s;
        }
        .stButton>button:hover {
            transform: translateY(-2px);
            box-shadow: 0 6px 20px rgba(0, 210, 255, 0.5);
        }
        .info-card {
            background: rgba(0, 0, 0, 0.4) !important;
            border-left: 4px solid #00d2ff;
            padding: 15px;
            border-radius: 0 10px 10px 0;
            margin-bottom: 15px;
            box-shadow: 2px 2px 10px rgba(0,0,0,0.3);
        }
        .info-title {
            color: #00d2ff !important;
            font-weight: bold;
            font-size: 1.1rem;
            display: block;
            margin-bottom: 8px;
            text-transform: uppercase;
        }
        .file-card {
            background: rgba(255, 255, 255, 0.08);
            border: 2px solid rgba(46, 204, 113, 0.5);
            border-radius: 15px;
            padding: 20px;
            margin-bottom: 15px;
            box-shadow: 0 4px 15px rgba(0, 0, 0, 0.3);
            animation: slideIn 0.5s ease;
        }
        @keyframes slideIn {
            from { opacity: 0; transform: translateY(10px); }
            to { opacity: 1; transform: translateY(0); }
        }
        .support-btn {
            display: block;
            text-decoration: none;
            background: transparent;
            border: 2px solid #25D366;
            color: #25D366 !important;
            padding: 10px;
            border-radius: 10px;
            font-weight: bold;
            text-align: center;
            margin-top: 10px;
            transition: 0.3s;
        }
        .support-btn:hover {
            background: #25D366;
            color: white !important;
        }
        .stProgress > div > div > div > div {
            background-image: linear-gradient(to right, #00d2ff , #3a7bd5);
        }
        .gemini-card {
            background: linear-gradient(135deg, rgba(0,210,255,0.08), rgba(58,123,213,0.12));
            border: 2px solid rgba(0,210,255,0.4);
            border-radius: 14px;
            padding: 16px 20px;
            margin: 12px 0;
        }
        .gemini-title {
            color: #00d2ff;
            font-weight: 800;
            font-size: 0.95rem;
            letter-spacing: 1px;
        }
        .gemini-sub {
            color: rgba(255,255,255,0.5);
            font-size: 0.75rem;
        }
        .badge-premium {
            display: inline-flex; align-items: center; gap: 6px;
            background: linear-gradient(135deg, #FFD700, #FF8C00);
            color: #000; font-weight: 800; font-size: 0.75rem;
            padding: 4px 12px; border-radius: 20px; text-transform: uppercase;
            box-shadow: 0 2px 10px rgba(255,215,0,0.4);
        }
        .badge-free {
            display: inline-flex; align-items: center; gap: 6px;
            background: rgba(255,255,255,0.1); color: rgba(255,255,255,0.5);
            font-size: 0.75rem; padding: 4px 12px; border-radius: 20px;
            border: 1px solid rgba(255,255,255,0.2);
        }
        .admin-premium-row {
            background: rgba(255,215,0,0.06); border: 1px solid rgba(255,215,0,0.25);
            border-radius: 14px; padding: 16px 20px; margin-bottom: 12px;
        }
        .admin-user-name { color: #fff; font-weight: 700; font-size: 1rem; }
        .admin-user-meta { color: rgba(255,255,255,0.45); font-size: 0.8rem; }
        .gemini-lock {
            background: linear-gradient(135deg, rgba(255,215,0,0.08), rgba(255,140,0,0.12));
            border: 2px solid rgba(255,215,0,0.4); border-radius: 14px; padding: 16px 20px; margin: 12px 0;
        }
        @keyframes nova-pulse {
            0%,100% { box-shadow: 0 0 20px rgba(255,215,0,0.4); }
            50%      { box-shadow: 0 0 60px rgba(255,215,0,0.9); }
        }
        .nova-processing {
            background: linear-gradient(135deg, rgba(255,215,0,0.1), rgba(255,140,0,0.08));
            border: 2px solid #FFD700; border-radius: 20px;
            padding: 30px; text-align: center; margin: 20px 0;
            animation: nova-pulse 2s ease-in-out infinite;
        }
        .nova-processing-title { color: #FFD700; font-size: 1.6rem; font-weight: 800; }
        .nova-processing-sub   { color: rgba(255,255,255,0.7); font-size: 1rem; margin-top: 8px; }
        .livrable-auto {
            background: linear-gradient(135deg, rgba(46,204,113,0.12), rgba(0,210,255,0.08));
            border: 2px solid #2ecc71; border-radius: 20px; padding: 28px; margin: 20px 0;
            box-shadow: 0 0 30px rgba(46,204,113,0.2);
        }
        .livrable-auto-title { color: #2ecc71; font-size: 1.4rem; font-weight: 800; }
        </style>
    """, unsafe_allow_html=True)

    # ── THÈME OR PREMIUM ──────────────────────────────────────────
    if _premium:
        st.markdown("""
        <style>
        @keyframes gold-shimmer {
            0%   { background-position: -300% center; }
            100% { background-position:  300% center; }
        }
        @keyframes gold-glow-pulse {
            0%,100% { box-shadow: inset 0 0 0px transparent; filter: brightness(1); }
            50%      { box-shadow: inset 0 0 120px rgba(255,215,0,0.18); filter: brightness(1.08); }
        }
        @keyframes gold-border-anim {
            0%   { border-color: #FFD700; box-shadow: 0 0 10px rgba(255,215,0,0.4); }
            50%  { border-color: #FF8C00; box-shadow: 0 0 20px rgba(255,140,0,0.5); }
            100% { border-color: #FFD700; box-shadow: 0 0 10px rgba(255,215,0,0.4); }
        }

        /* ── Fond général ── */
        .stApp {
            background: linear-gradient(135deg, #0a0800 0%, #1c1400 35%, #0d0a00 65%, #1a1000 100%) !important;
            animation: gold-glow-pulse 5s ease-in-out infinite !important;
        }

        /* ── Titre principal NOVA AI PLATFORM ── */
        .main-title {
            background: linear-gradient(90deg, #7a5500, #b8860b, #FFD700, #fff5c0, #FFD700, #b8860b, #7a5500) !important;
            background-size: 300% auto !important;
            -webkit-background-clip: text !important;
            -webkit-text-fill-color: transparent !important;
            animation: gold-shimmer 3s linear infinite !important;
            text-shadow: none !important;
            font-size: 3.8rem !important;
            letter-spacing: 2px !important;
        }

        /* ── Tabs ── */
        .stTabs [data-baseweb="tab-list"] {
            background-color: rgba(255,215,0,0.06) !important;
            border: 1px solid rgba(255,215,0,0.25) !important;
        }
        .stTabs [data-baseweb="tab"] {
            background-color: rgba(255,215,0,0.08) !important;
            border: 1px solid rgba(255,215,0,0.15) !important;
            color: #FFD700 !important;
        }
        .stTabs [data-baseweb="tab"]:nth-child(2) {
            border: 1px solid rgba(255,215,0,0.5) !important;
            box-shadow: 0 0 15px rgba(255,215,0,0.2) !important;
            background-color: rgba(255,215,0,0.1) !important;
        }
        .stTabs [data-baseweb="tab"]:hover {
            background-color: rgba(255,215,0,0.2) !important;
        }
        .stTabs [aria-selected="true"] {
            background: linear-gradient(135deg, rgba(255,215,0,0.35), rgba(255,140,0,0.25)) !important;
            border: 1px solid #FFD700 !important;
            box-shadow: 0 0 20px rgba(255,215,0,0.4) !important;
        }

        /* ── Labels formulaires ── */
        .stTextInput label, .stSelectbox label, .stTextArea label {
            color: #FFD700 !important;
        }

        /* ── Inputs / Selects ── */
        div[data-baseweb="input"], div[data-baseweb="select"] > div {
            border: 1px solid rgba(255,215,0,0.4) !important;
            background-color: rgba(20,12,0,0.7) !important;
        }
        .stTextArea textarea {
            background-color: rgba(20,12,0,0.8) !important;
            border: 2px solid #FFD700 !important;
            animation: gold-border-anim 3s ease-in-out infinite !important;
        }

        /* ── Boutons principaux ── */
        .stButton > button {
            background: linear-gradient(90deg, #7a5500, #b8860b, #FFD700, #b8860b, #7a5500) !important;
            background-size: 200% auto !important;
            color: #0a0800 !important;
            animation: gold-shimmer 3s linear infinite !important;
            box-shadow: 0 4px 18px rgba(255,215,0,0.35) !important;
            border: none !important;
        }
        .stButton > button:hover {
            box-shadow: 0 6px 28px rgba(255,215,0,0.6) !important;
            transform: translateY(-2px) !important;
        }

        /* ── Info cards sidebar ── */
        .info-card {
            border-left: 4px solid #FFD700 !important;
            background: rgba(20,12,0,0.6) !important;
        }
        .info-title { color: #FFD700 !important; }

        /* ── Support btn ── */
        .support-btn {
            border: 2px solid #FFD700 !important;
            color: #FFD700 !important;
        }
        .support-btn:hover {
            background: #FFD700 !important;
            color: #000 !important;
        }

        /* ── Barre de progression ── */
        .stProgress > div > div > div > div {
            background-image: linear-gradient(to right, #b8860b, #FFD700) !important;
        }

        /* ── Gemini card ── */
        .gemini-card {
            background: linear-gradient(135deg, rgba(255,215,0,0.08), rgba(255,140,0,0.06)) !important;
            border: 2px solid rgba(255,215,0,0.5) !important;
        }
        .gemini-title { color: #FFD700 !important; }

        /* ── File cards ── */
        .file-card {
            border: 2px solid rgba(255,215,0,0.5) !important;
            background: rgba(20,12,0,0.5) !important;
        }

        /* ── Livrable auto ── */
        .livrable-auto {
            background: linear-gradient(135deg, rgba(255,215,0,0.12), rgba(255,140,0,0.08)) !important;
            border: 2px solid #FFD700 !important;
            box-shadow: 0 0 35px rgba(255,215,0,0.25) !important;
        }
        .livrable-auto-title { color: #FFD700 !important; }

        /* ── Sidebar ── */
        [data-testid="stSidebar"] {
            background: linear-gradient(180deg, #0d0900, #1a1200) !important;
            border-right: 1px solid rgba(255,215,0,0.2) !important;
        }

        /* ── Divider ── */
        hr { border-color: rgba(255,215,0,0.2) !important; }

        /* ── Métriques admin ── */
        [data-testid="stMetric"] {
            background: rgba(255,215,0,0.06) !important;
            border: 1px solid rgba(255,215,0,0.2) !important;
            border-radius: 12px !important;
            padding: 10px !important;
        }

        /* ── Logo container ── */
        .logo-container {
            background: rgba(255,215,0,0.04) !important;
            border: 1px solid rgba(255,215,0,0.12) !important;
        }

        /* ── Glow global sur le body ── */
        body::before {
            content: '';
            position: fixed;
            top: 0; left: 0; right: 0; bottom: 0;
            background: radial-gradient(ellipse at 50% 0%, rgba(255,215,0,0.06) 0%, transparent 60%);
            pointer-events: none;
            z-index: 0;
        }
        </style>
        """, unsafe_allow_html=True)

    # ── THÈME OR EXCLUSIF PREMIUM ──────────────────────────────────
    user_now = st.session_state.get("current_user")
    db_now   = st.session_state.get("db", {})
    ud_now   = db_now.get("users", {}).get(user_now, {}) if user_now else {}
    if is_premium_actif(ud_now):
        st.markdown("""
        <style>
        /* ===== FOND OR PREMIUM ===== */
        .stApp {
            background: #2d1f00 !important;
            background: -webkit-linear-gradient(135deg, #3d2800 0%, #4a3200 20%, #3a2600 40%, #4d3500 60%, #3d2900 80%, #2d1f00 100%) !important;
            background: linear-gradient(135deg, #3d2800 0%, #4a3200 20%, #3a2600 40%, #4d3500 60%, #3d2900 80%, #2d1f00 100%) !important;
            color: #fff8e1 !important;
        }

        /* Halos lumineux dorés très visibles */
        .stApp::before {
            content: '';
            position: fixed;
            inset: 0;
            background:
                radial-gradient(ellipse at 10% 10%, rgba(255,215,0,0.30) 0%, transparent 40%),
                radial-gradient(ellipse at 90% 90%, rgba(255,160,0,0.25) 0%, transparent 40%),
                radial-gradient(ellipse at 50% 50%, rgba(255,200,0,0.12) 0%, transparent 60%),
                radial-gradient(ellipse at 85% 10%, rgba(255,215,0,0.20) 0%, transparent 35%),
                radial-gradient(ellipse at 15% 90%, rgba(255,140,0,0.18) 0%, transparent 35%);
            pointer-events: none;
            z-index: 0;
        }

        /* ===== TITRE PRINCIPAL OR ===== */
        .main-title {
            background: linear-gradient(90deg, #b8860b, #FFD700, #fff5c0, #FFD700, #b8860b) !important;
            background-size: 200% auto !important;
            -webkit-background-clip: text !important;
            -webkit-text-fill-color: transparent !important;
            animation: shimmer-gold 3s linear infinite !important;
            text-shadow: none !important;
            filter: drop-shadow(0 0 20px rgba(255,215,0,0.5));
        }
        @keyframes shimmer-gold {
            0%   { background-position: -200% center; }
            100% { background-position:  200% center; }
        }

        /* ===== TABS OR ===== */
        .stTabs [data-baseweb="tab-list"] {
            background-color: rgba(255,215,0,0.05) !important;
            border: 1px solid rgba(255,215,0,0.2) !important;
        }
        .stTabs [data-baseweb="tab"] {
            background-color: rgba(255,215,0,0.07) !important;
            border: 1px solid rgba(255,215,0,0.15) !important;
            color: #FFD700 !important;
        }
        .stTabs [data-baseweb="tab"]:nth-child(2) {
            border: 1px solid rgba(255,215,0,0.5) !important;
            box-shadow: 0 0 15px rgba(255,215,0,0.15) !important;
            background-color: rgba(255,215,0,0.1) !important;
        }
        .stTabs [data-baseweb="tab"]:hover {
            background-color: rgba(255,215,0,0.2) !important;
        }
        .stTabs [aria-selected="true"] {
            background-color: rgba(255,215,0,0.3) !important;
            border: 1px solid #FFD700 !important;
            box-shadow: 0 0 20px rgba(255,215,0,0.4) !important;
        }

        /* ===== INPUTS OR ===== */
        .stTextInput label, .stSelectbox label, .stTextArea label {
            color: #FFD700 !important;
        }
        div[data-baseweb="input"], div[data-baseweb="select"] > div {
            border: 1px solid rgba(255,215,0,0.6) !important;
            background-color: rgba(70,48,0,0.80) !important;
            color: #fff8e1 !important;
        }
        .stTextArea textarea {
            background-color: rgba(65,44,0,0.80) !important;
            color: #fff8e1 !important;
            border: 2px solid #FFD700 !important;
            animation: border-gold 4s linear infinite !important;
        }
        @keyframes border-gold {
            0%   { border-color: #FFD700; box-shadow: 0 0 14px rgba(255,215,0,0.55); }
            33%  { border-color: #FF8C00; box-shadow: 0 0 18px rgba(255,140,0,0.45); }
            66%  { border-color: #b8860b; box-shadow: 0 0 14px rgba(184,134,11,0.45); }
            100% { border-color: #FFD700; box-shadow: 0 0 14px rgba(255,215,0,0.55); }
        }

        /* ===== BOUTONS OR ===== */
        .stButton>button {
            background: linear-gradient(90deg, #8a6200, #c49a00, #FFD700, #c49a00, #8a6200) !important;
            background-size: 200% auto !important;
            color: #1a0f00 !important;
            box-shadow: 0 4px 22px rgba(255,215,0,0.55) !important;
            animation: shimmer-gold 3s linear infinite !important;
            font-weight: 800 !important;
        }
        .stButton>button:hover {
            box-shadow: 0 6px 32px rgba(255,215,0,0.75) !important;
            transform: translateY(-2px) !important;
        }

        /* ===== SIDEBAR OR ===== */
        section[data-testid="stSidebar"] {
            background: linear-gradient(180deg, #3a2800 0%, #4a3400 40%, #3a2800 100%) !important;
            border-right: 2px solid rgba(255,215,0,0.4) !important;
            box-shadow: 4px 0 25px rgba(255,215,0,0.12) !important;
        }

        /* ===== INFO-CARD OR ===== */
        .info-card {
            border-left: 4px solid #FFD700 !important;
            background: rgba(255,215,0,0.10) !important;
        }
        .info-title { color: #FFD700 !important; }

        /* ===== FILE-CARD OR ===== */
        .file-card {
            border: 2px solid rgba(255,215,0,0.5) !important;
            background: rgba(255,215,0,0.07) !important;
        }

        /* ===== PROGRESS BAR OR ===== */
        .stProgress > div > div > div > div {
            background-image: linear-gradient(to right, #b8860b, #FFD700, #FF8C00) !important;
        }

        /* ===== EXPANDER OR ===== */
        .streamlit-expanderHeader {
            color: #FFD700 !important;
            border: 1px solid rgba(255,215,0,0.3) !important;
            background: rgba(255,215,0,0.10) !important;
        }

        /* ===== DIVIDER OR ===== */
        hr { border-color: rgba(255,215,0,0.3) !important; }

        /* ===== METRIC OR ===== */
        [data-testid="stMetric"] {
            background: rgba(255,215,0,0.10) !important;
            border: 1px solid rgba(255,215,0,0.35) !important;
            border-radius: 12px !important;
            padding: 10px !important;
        }
        [data-testid="stMetricValue"] { color: #FFD700 !important; }

        /* ===== SUCCESS / INFO / WARNING OR ===== */
        .stSuccess {
            background: rgba(255,215,0,0.12) !important;
            border: 1px solid rgba(255,215,0,0.4) !important;
            color: #FFD700 !important;
        }
        .stInfo {
            background: rgba(255,215,0,0.08) !important;
            border: 1px solid rgba(255,215,0,0.3) !important;
        }

        /* ===== SUPPORT BTN OR ===== */
        .support-btn {
            border: 2px solid #FFD700 !important;
            color: #FFD700 !important;
        }
        .support-btn:hover {
            background: #FFD700 !important;
            color: #0a0800 !important;
        }

        /* ===== SCROLLBAR OR ===== */
        ::-webkit-scrollbar-thumb {
            background: linear-gradient(#FFD700, #b8860b) !important;
        }
        ::-webkit-scrollbar-track {
            background: #0a0800 !important;
        }
        </style>
        """, unsafe_allow_html=True)

    if st.session_state["is_glowing"]:
        st.markdown('<style>.stApp { animation: glow-pulse 1.5s ease-in-out infinite; }</style>', unsafe_allow_html=True)


def show_auth_page():
    st.markdown("""
    <style>
    @keyframes shimmer {
        0%   { background-position: -200% center; }
        100% { background-position:  200% center; }
    }
    @keyframes float-up {
        0%   { opacity: 0; transform: translateY(30px); }
        100% { opacity: 1; transform: translateY(0); }
    }
    @keyframes letter-pop {
        0%   { opacity: 0; transform: translateY(20px) scale(0.8); }
        60%  { transform: translateY(-4px) scale(1.05); }
        100% { opacity: 1; transform: translateY(0) scale(1); }
    }
    @keyframes glow-border {
        0%   { box-shadow: 0 0 8px rgba(255,215,0,0.3), inset 0 0 8px rgba(255,215,0,0.05); }
        50%  { box-shadow: 0 0 28px rgba(255,215,0,0.7), inset 0 0 20px rgba(255,215,0,0.08); }
        100% { box-shadow: 0 0 8px rgba(255,215,0,0.3), inset 0 0 8px rgba(255,215,0,0.05); }
    }
    @keyframes particle-drift {
        0%   { transform: translateY(0px) translateX(0px) rotate(0deg); opacity: 0.6; }
        33%  { transform: translateY(-18px) translateX(8px) rotate(120deg); opacity: 1; }
        66%  { transform: translateY(-8px) translateX(-6px) rotate(240deg); opacity: 0.7; }
        100% { transform: translateY(0px) translateX(0px) rotate(360deg); opacity: 0.6; }
    }
    @keyframes scanline {
        0%   { top: -10%; }
        100% { top: 110%; }
    }
    @keyframes pulse-dot {
        0%, 100% { opacity: 1; transform: scale(1); }
        50%       { opacity: 0.4; transform: scale(0.6); }
    }
    .auth-hero {
        text-align: center;
        padding: 40px 20px 10px 20px;
        animation: float-up 0.8s ease both;
    }
    .auth-logo-ring {
        width: 90px; height: 90px;
        border-radius: 50%;
        margin: 0 auto 18px auto;
        background: radial-gradient(circle at 35% 35%, #fff8e1, #FFD700 40%, #b8860b);
        box-shadow: 0 0 0 4px rgba(255,215,0,0.2), 0 0 40px rgba(255,215,0,0.5);
        display: flex; align-items: center; justify-content: center;
        font-size: 2.6rem;
        animation: glow-border 3s ease-in-out infinite;
        position: relative;
    }
    .auth-logo-ring::after {
        content: '';
        position: absolute;
        inset: -6px;
        border-radius: 50%;
        border: 2px dashed rgba(255,215,0,0.4);
        animation: particle-drift 6s linear infinite;
    }
    .auth-title-wrap { display: flex; justify-content: center; gap: 2px; flex-wrap: wrap; margin-bottom: 6px; }
    .auth-letter {
        font-size: 3rem;
        font-weight: 800;
        background: linear-gradient(90deg, #b8860b, #FFD700, #fff5c0, #FFD700, #b8860b);
        background-size: 200% auto;
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        animation: letter-pop 0.5s ease both, shimmer 3s linear infinite;
        display: inline-block;
        line-height: 1.1;
    }
    .auth-subtitle {
        color: rgba(255,215,0,0.65);
        font-size: 0.95rem;
        letter-spacing: 4px;
        text-transform: uppercase;
        animation: float-up 1s ease 0.5s both;
        margin-bottom: 6px;
    }
    .auth-tagline {
        color: rgba(255,255,255,0.4);
        font-size: 0.82rem;
        letter-spacing: 1.5px;
        animation: float-up 1s ease 0.8s both;
    }
    .auth-divider {
        display: flex; align-items: center; gap: 14px;
        margin: 28px auto 32px auto; max-width: 420px;
        animation: float-up 1s ease 1s both;
    }
    .auth-divider-line { flex: 1; height: 1px; background: linear-gradient(90deg, transparent, rgba(255,215,0,0.5), transparent); }
    .auth-divider-dot {
        width: 6px; height: 6px; border-radius: 50%; background: #FFD700;
        animation: pulse-dot 1.8s ease-in-out infinite;
    }
    .auth-card {
        background: linear-gradient(145deg, rgba(20,15,5,0.95), rgba(35,25,5,0.9));
        border: 1px solid rgba(255,215,0,0.35);
        border-radius: 22px;
        padding: 32px 28px 28px 28px;
        position: relative;
        overflow: hidden;
        animation: float-up 0.9s ease both;
        animation-delay: var(--card-delay, 0s);
    }
    .auth-card::before {
        content: '';
        position: absolute;
        top: 0; left: 0; right: 0; height: 3px;
        background: linear-gradient(90deg, #b8860b, #FFD700, #fff5c0, #FFD700, #b8860b);
        background-size: 200% auto;
        animation: shimmer 2.5s linear infinite;
        border-radius: 22px 22px 0 0;
    }
    .auth-card::after {
        content: '';
        position: absolute;
        width: 180px; height: 180px;
        border-radius: 50%;
        background: radial-gradient(circle, rgba(255,215,0,0.06), transparent 70%);
        bottom: -60px; right: -60px;
        pointer-events: none;
    }
    .auth-card .scanline {
        position: absolute;
        left: 0; right: 0; height: 2px;
        background: linear-gradient(90deg, transparent, rgba(255,215,0,0.12), transparent);
        animation: scanline 4s linear infinite;
        pointer-events: none;
    }
    .auth-card-header {
        display: flex; align-items: center; gap: 12px;
        margin-bottom: 22px;
    }
    .auth-card-icon {
        width: 44px; height: 44px; border-radius: 12px;
        display: flex; align-items: center; justify-content: center;
        font-size: 1.4rem;
        background: linear-gradient(135deg, rgba(255,215,0,0.15), rgba(255,215,0,0.05));
        border: 1px solid rgba(255,215,0,0.3);
        box-shadow: 0 0 12px rgba(255,215,0,0.15);
    }
    .auth-card-title {
        color: #FFD700 !important;
        font-size: 1.15rem;
        font-weight: 800;
        letter-spacing: 0.5px;
        margin: 0;
        text-transform: uppercase;
    }
    .auth-card-desc {
        color: rgba(255,255,255,0.35);
        font-size: 0.75rem;
        margin-top: 2px;
        letter-spacing: 0.5px;
    }
    .auth-page .stTextInput label {
        color: rgba(255,215,0,0.8) !important;
        font-size: 0.82rem !important;
        font-weight: 600 !important;
        letter-spacing: 1px;
        text-transform: uppercase;
    }
    .auth-page div[data-baseweb="input"] {
        background: rgba(0,0,0,0.5) !important;
        border: 1px solid rgba(255,215,0,0.25) !important;
        border-radius: 12px !important;
        transition: border-color 0.3s, box-shadow 0.3s !important;
    }
    .auth-page div[data-baseweb="input"]:focus-within {
        border-color: rgba(255,215,0,0.7) !important;
        box-shadow: 0 0 0 3px rgba(255,215,0,0.12) !important;
    }
    @keyframes btn-shimmer {
        0%   { background-position: -300% center; }
        100% { background-position:  300% center; }
    }
    @keyframes btn-float {
        0%, 100% { transform: translateY(0px);   box-shadow: 0 6px 25px rgba(255,215,0,0.45), 0 0 0 0 rgba(255,215,0,0.2); }
        50%       { transform: translateY(-4px);  box-shadow: 0 14px 35px rgba(255,215,0,0.65), 0 0 18px 4px rgba(255,215,0,0.15); }
    }
    @keyframes btn-glow-ring {
        0%, 100% { box-shadow: 0 6px 25px rgba(255,215,0,0.45), 0 0  0px rgba(255,215,0,0);   }
        50%       { box-shadow: 0 6px 25px rgba(255,215,0,0.45), 0 0 22px rgba(255,215,0,0.35); }
    }
    .auth-page .stButton > button {
        background: linear-gradient(90deg, #7a5500, #b8860b, #FFD700, #fff5c0, #FFD700, #b8860b, #7a5500) !important;
        background-size: 300% auto !important;
        color: #0a0800 !important;
        font-weight: 800 !important;
        font-size: 0.92rem !important;
        letter-spacing: 2.5px !important;
        text-transform: uppercase !important;
        border-radius: 50px !important;
        border: none !important;
        padding: 0.82rem 1.8rem !important;
        width: 100% !important;
        position: relative !important;
        overflow: hidden !important;
        animation: btn-shimmer 3s linear infinite, btn-float 3.5s ease-in-out infinite !important;
        cursor: pointer !important;
    }
    .auth-secure-badge {
        display: flex; align-items: center; justify-content: center;
        gap: 8px; margin-top: 28px;
        color: rgba(255,215,0,0.35);
        font-size: 0.72rem; letter-spacing: 1.5px; text-transform: uppercase;
        animation: float-up 1s ease 1.2s both;
    }
    .auth-secure-badge span { font-size: 0.9rem; }
    </style>
    """, unsafe_allow_html=True)

    letters = list("NOVA AI")
    letter_spans = "".join(
        f'<span class="auth-letter" style="animation-delay:{i*0.07:.2f}s">'
        f'{"&nbsp;" if c == " " else c}</span>'
        for i, c in enumerate(letters)
    )
    st.markdown(f"""
    <div class="auth-hero">
        <div class="auth-logo-ring">⚡</div>
        <div class="auth-title-wrap">{letter_spans}</div>
        <div class="auth-subtitle">Plateforme IA bureautique</div>
        <div class="auth-tagline">Intelligence · Excellence · Performance</div>
    </div>
    <div class="auth-divider">
        <div class="auth-divider-line"></div>
        <div class="auth-divider-dot"></div>
        <div class="auth-divider-line"></div>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns(2, gap="large")

    with col1:
        st.markdown("""
        <div class="auth-card" style="--card-delay:1.1s;">
            <div class="scanline"></div>
            <div class="auth-card-header">
                <div class="auth-card-icon">🔐</div>
                <div>
                    <div class="auth-card-title">Accès Membre</div>
                    <div class="auth-card-desc">Identifiez-vous pour accéder à votre espace</div>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        with st.form("login"):
            uid    = st.text_input("Identifiant Nova", placeholder="Votre identifiant...")
            wa_auth = st.text_input("Numéro WhatsApp", placeholder="Ex: 22501...")
            if st.form_submit_button("⚡ S'IDENTIFIER"):
                # Recharger depuis Supabase pour avoir les données fraîches
                fresh_db = load_db()
                st.session_state["db"] = fresh_db
                db = fresh_db
                if uid in db["users"] and db["users"][uid]["whatsapp"] == normalize_wa(wa_auth):
                    st.session_state["current_user"] = uid
                    st.session_state["view"] = "home"
                    st.query_params["user_id"] = uid
                    st.rerun()
                else:
                    st.error("❌ Identifiant ou numéro inconnu.")

    with col2:
        st.markdown("""
        <div class="auth-card" style="--card-delay:1.3s;">
            <div class="scanline"></div>
            <div class="auth-card-header">
                <div class="auth-card-icon">✨</div>
                <div>
                    <div class="auth-card-title">Nouveau Compte</div>
                    <div class="auth-card-desc">Rejoignez l'élite Nova AI dès maintenant</div>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        with st.form("signup"):
            new_uid = st.text_input("Identifiant au choix", placeholder="Choisissez un identifiant...")
            new_wa  = st.text_input("Votre WhatsApp (clé d'accès)", placeholder="Ex: 22507...")
            if st.form_submit_button("💎 REJOINDRE NOVA AI"):
                if new_uid and new_wa:
                    db = st.session_state["db"]
                    if new_uid not in db["users"]:
                        succes = save_user(new_uid, normalize_wa(new_wa))
                        if succes:
                            db["users"][new_uid] = {
                                "whatsapp": normalize_wa(new_wa),
                                "email": "Non renseigné",
                                "joined": str(datetime.now()),
                                "premium": False,
                                "premium_plan": None,
                                "premium_expiry": None,
                                "gen_used": 0,
                                "gen_date": None,
                            }
                            st.session_state["current_user"] = new_uid
                            st.session_state["view"] = "home"
                            st.session_state["db"] = load_db()
                            st.query_params["user_id"] = new_uid
                            st.rerun()
                        else:
                            st.error("❌ Impossible de créer le compte. Vérifie ta connexion ou contacte le support.")
                    else:
                        st.warning("⚠️ Identifiant déjà utilisé.")
                else:
                    st.error("Champs obligatoires.")


    st.markdown("""
    <div class="auth-secure-badge">
        <span>🔒</span> Connexion sécurisée &nbsp;·&nbsp; <span>⚡</span> Nova AI &nbsp;·&nbsp; <span>🛡️</span> Données protégées
    </div>
    """, unsafe_allow_html=True)

    audio_path_login = "login.mp3"
    if os.path.exists(audio_path_login):
        with open(audio_path_login, "rb") as f:
            audio_b64_login = __import__('base64').b64encode(f.read()).decode()
        components.html(f"""
            <script>
            (function() {{
                setTimeout(function() {{
                    var b64 = "{audio_b64_login}";
                    var binary = atob(b64);
                    var bytes = new Uint8Array(binary.length);
                    for (var i = 0; i < binary.length; i++) {{
                        bytes[i] = binary.charCodeAt(i);
                    }}
                    var blob = new Blob([bytes], {{type: "audio/mpeg"}});
                    var url = URL.createObjectURL(blob);
                    var audio = new Audio(url);
                    audio.volume = 1;
                    audio.play().catch(function(e) {{ console.log("Autoplay bloqué:", e); }});
                }}, 3000);
            }})();
            </script>
        """, height=0)


def main_dashboard():
    user = st.session_state["current_user"]
    db = st.session_state["db"]

    if user and user in db["users"]:
        ud = db["users"][user]
        if ud.get("premium") and not is_premium_actif(ud):
            desactiver_premium(user)
            st.session_state["db"] = load_db()
            db = st.session_state["db"]

    user_data     = db["users"].get(user, {}) if user else {}
    premium_actif = is_premium_actif(user_data)
    premium_info  = get_premium_info(user_data)

    with st.sidebar:
        st.markdown(f"### 👤 {user if user else 'Visiteur'}")
        if user:
            st.markdown(f"📱 **{db['users'][user]['whatsapp']}**")
            if premium_actif and premium_info:
                st.markdown(f"""
                <div style="margin:10px 0;">
                    <span class="badge-premium">⭐ PREMIUM — {premium_info['plan']}</span>
                    <div style="color:rgba(255,215,0,0.7);font-size:0.78rem;margin-top:6px;">
                        ⏳ Expire le {premium_info['expiry']} ({premium_info['jours_restants']}j restants)
                    </div>
                </div>""", unsafe_allow_html=True)
            else:
                st.markdown('<span class="badge-free">🔓 Compte Gratuit</span>', unsafe_allow_html=True)
            if st.button("Quitter la session"):
                st.session_state["current_user"] = None
                st.query_params.clear()
                components.html("<script>localStorage.removeItem('nova_user_id');</script>", height=0)
                st.rerun()
        else:
            if st.button("Connexion"):
                st.session_state["view"] = "auth"
                st.rerun()
        
        st.divider()
        st.markdown(f"""
            <div class="info-card">
                <span class="info-title">🚀 LIVRAISON NOVA</span>
                <span style="color:#eee; font-size:0.9rem;">
                    Vos résultats IA apparaissent dans l'onglet <b>"📂 MES LIVRABLES"</b>.
                    <br><br>
                    Suivi instantané 24h/24.
                </span>
            </div>
        """, unsafe_allow_html=True)
        st.markdown(f'<a href="{whatsapp_support_url}" target="_blank" class="support-btn">💬 Support Nova</a>', unsafe_allow_html=True)

    if premium_actif:
        st.markdown("""
        <div style="text-align:center; margin-bottom:20px;">
            <div style="font-size:2.2rem; margin-bottom:-10px; filter:drop-shadow(0 0 15px rgba(255,215,0,0.8));">👑</div>
            <h1 class='main-title' style="
                background: linear-gradient(90deg, #b8860b, #FFD700, #fff5c0, #FFD700, #b8860b);
                background-size: 200% auto;
                -webkit-background-clip: text;
                -webkit-text-fill-color: transparent;
                animation: shimmer-gold 3s linear infinite;
                filter: drop-shadow(0 0 25px rgba(255,215,0,0.6));
                font-size: 3.5rem !important;
                font-weight: 800 !important;
                margin-top: 0;
            ">NOVA AI PLATFORM</h1>
            <div style="
                color: rgba(255,215,0,0.6);
                font-size: 0.75rem;
                letter-spacing: 5px;
                text-transform: uppercase;
                margin-top: -10px;
            ">✦ Membre Premium Actif ✦</div>
        </div>
        """, unsafe_allow_html=True)
    else:
        st.markdown("<h1 class='main-title'>NOVA AI PLATFORM</h1>", unsafe_allow_html=True)

    if not st.session_state["intro_played"]:
        st.session_state["intro_played"] = True
        audio_path = "intro.mp3"
        if os.path.exists(audio_path):
            with open(audio_path, "rb") as f:
                audio_b64 = __import__('base64').b64encode(f.read()).decode()
            components.html(f"""
                <script>
                (function() {{
                    setTimeout(function() {{
                        var b64 = "{audio_b64}";
                        var binary = atob(b64);
                        var bytes = new Uint8Array(binary.length);
                        for (var i = 0; i < binary.length; i++) {{
                            bytes[i] = binary.charCodeAt(i);
                        }}
                        var blob = new Blob([bytes], {{type: "audio/mpeg"}});
                        var url = URL.createObjectURL(blob);
                        var audio = new Audio(url);
                        audio.volume = 1;
                        audio.play().catch(function(e) {{ console.log("Autoplay bloqué:", e); }});
                    }}, 3000);
                }})();
                </script>
            """, height=0)

    wa_jour = f"https://wa.me/{WHATSAPP_NUMBER}?text=Je%20souhaite%20l%27abonnement%20Nova%20Premium%20Journalier%20%C3%A0%20600%20FC."
    wa_10j  = f"https://wa.me/{WHATSAPP_NUMBER}?text=Je%20souhaite%20l%27abonnement%20Nova%20Premium%2010%20Jours%20%C3%A0%201000%20FC."
    wa_30j  = f"https://wa.me/{WHATSAPP_NUMBER}?text=Je%20souhaite%20l%27abonnement%20Nova%20Premium%2030%20Jours%20%C3%A0%202500%20FC."

    if premium_actif and premium_info:
        st.markdown(f"""
        <div style="background:linear-gradient(135deg,rgba(255,215,0,.12),rgba(255,140,0,.08));
             border:2px solid #FFD700;border-radius:20px;padding:20px;text-align:center;margin-bottom:20px;">
            <div style="font-size:1.4rem;font-weight:800;color:#FFD700;">
                ⭐ MEMBRE PREMIUM ACTIF — {premium_info['plan']}
            </div>
            <div style="color:rgba(255,255,255,.7);margin-top:6px;">
                🤖 Génération IA instantanée activée · Expire le <b>{premium_info['expiry']}</b>
                ({premium_info['jours_restants']} jour(s) restant(s))
            </div>
        </div>""", unsafe_allow_html=True)
    else:
        st.markdown("""
            <div class="premium-card">
                <div class="premium-title">⭐ ACCÉLÉRATEUR NOVA PREMIUM ⭐</div>
                <div class="premium-desc">
                    Passez au niveau supérieur : IA illimitée et puissance de calcul <b>10<sup>10</sup></b>.
                </div>
            </div>
        """, unsafe_allow_html=True)

        col_btn_center = st.columns([1, 2, 1])[1]
        with col_btn_center:
            if st.button("💎 ACTIVER NOVA PREMIUM", key="open_premium"):
                st.session_state["show_premium_modal"] = True
                st.rerun()

    if st.session_state["show_premium_modal"]:
        st.markdown("""
            <div style="
                background: linear-gradient(135deg, #1a1a2e, #16213e, #0f3460);
                border: 2px solid #FFD700;
                border-radius: 24px;
                padding: 35px 25px 30px 25px;
                margin: 10px 0 30px 0;
                box-shadow: 0 0 60px rgba(255,215,0,0.25);
            ">
                <h2 style="text-align:center; color:#FFD700; font-size:1.7rem; font-weight:800; margin-bottom:6px; letter-spacing:1px;">
                    ⭐ CHOISISSEZ VOTRE FORMULE NOVA PREMIUM
                </h2>
                <p style="text-align:center; color:rgba(255,255,255,0.55); margin-bottom:30px; font-size:0.95rem;">
                    Sélectionnez le plan qui correspond à vos besoins
                </p>
            </div>
        """, unsafe_allow_html=True)

        col1_p, col2_p, col3_p = st.columns(3)

        with col1_p:
            st.markdown("""
                <div style="background: rgba(255,255,255,0.05); border: 1px solid rgba(255,215,0,0.4); border-radius: 18px; padding: 28px 16px; text-align: center; min-height: 300px;">
                    <div style="font-size:2.5rem; margin-bottom:10px;">🌅</div>
                    <div style="color:#FFD700; font-weight:800; font-size:1.1rem; margin-bottom:6px; text-transform:uppercase;">Journalier</div>
                    <div style="color:white; font-size:2rem; font-weight:800; margin:10px 0;">600 FC</div>
                    <div style="color:rgba(255,255,255,0.45); font-size:0.8rem; margin-bottom:16px;">/ par jour</div>
                    <div style="background:rgba(255,215,0,0.1); border-radius:10px; padding:10px; margin-bottom:22px;">
                        <span style="color:#FFD700; font-size:0.9rem;">⚡ 2 générations IA / jour</span>
                    </div>
                </div>
            """, unsafe_allow_html=True)
            st.markdown(f'<a href="{wa_jour}" target="_blank" style="display:block; background:linear-gradient(45deg,#FFD700,#FF8C00); color:#000; font-weight:800; padding:12px; border-radius:50px; text-decoration:none; font-size:1rem; text-align:center; margin-top:10px;">Choisir cette formule</a>', unsafe_allow_html=True)

        with col2_p:
            st.markdown("""
                <div style="background: linear-gradient(135deg, rgba(0,210,255,0.15), rgba(58,123,213,0.15)); border: 2px solid #00d2ff; border-radius: 18px; padding: 28px 16px; text-align: center; min-height: 300px; position: relative;">
                    <div style="background:linear-gradient(90deg,#00d2ff,#3a7bd5); color:white; font-size:0.75rem; font-weight:800; padding:4px 16px; border-radius:20px; display:inline-block; margin-bottom:12px;">⭐ POPULAIRE</div>
                    <div style="font-size:2.5rem; margin-bottom:10px;">🔟</div>
                    <div style="color:#00d2ff; font-weight:800; font-size:1.1rem; margin-bottom:6px; text-transform:uppercase;">10 Jours</div>
                    <div style="color:white; font-size:2rem; font-weight:800; margin:10px 0;">1 000 FC</div>
                    <div style="color:rgba(255,255,255,0.45); font-size:0.8rem; margin-bottom:16px;">/ 10 jours</div>
                    <div style="background:rgba(0,210,255,0.1); border-radius:10px; padding:10px; margin-bottom:22px;">
                        <span style="color:#00d2ff; font-size:0.9rem;">⚡ 9 générations IA / jour</span>
                    </div>
                </div>
            """, unsafe_allow_html=True)
            st.markdown(f'<a href="{wa_10j}" target="_blank" style="display:block; background:linear-gradient(45deg,#00d2ff,#3a7bd5); color:#fff; font-weight:800; padding:12px; border-radius:50px; text-decoration:none; font-size:1rem; text-align:center; margin-top:10px;">Choisir cette formule</a>', unsafe_allow_html=True)

        with col3_p:
            st.markdown("""
                <div style="background: rgba(255,255,255,0.05); border: 1px solid rgba(46,204,113,0.4); border-radius: 18px; padding: 28px 16px; text-align: center; min-height: 300px;">
                    <div style="font-size:2.5rem; margin-bottom:10px;">👑</div>
                    <div style="color:#2ecc71; font-weight:800; font-size:1.1rem; margin-bottom:6px; text-transform:uppercase;">30 Jours</div>
                    <div style="color:white; font-size:2rem; font-weight:800; margin:10px 0;">2 500 FC</div>
                    <div style="color:rgba(255,255,255,0.45); font-size:0.8rem; margin-bottom:16px;">/ 30 jours</div>
                    <div style="background:rgba(46,204,113,0.1); border-radius:10px; padding:10px; margin-bottom:22px;">
                        <span style="color:#2ecc71; font-size:0.9rem;">♾️ Générations ILLIMITÉES</span>
                    </div>
                </div>
            """, unsafe_allow_html=True)
            st.markdown(f'<a href="{wa_30j}" target="_blank" style="display:block; background:linear-gradient(45deg,#2ecc71,#27ae60); color:#fff; font-weight:800; padding:12px; border-radius:50px; text-decoration:none; font-size:1rem; text-align:center; margin-top:10px;">Choisir cette formule</a>', unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)
        col_close = st.columns([1, 2, 1])[1]
        with col_close:
            if st.button("✕ Fermer", key="close_premium"):
                st.session_state["show_premium_modal"] = False
                st.rerun()

    # ── VÉRIFICATION AUTO-REPLY GRATUIT (tourne à chaque refresh) ──────
    if st.session_state.get("auto_reply_gratuit", False) and user:
        try:
            _fresh_demandes = supabase.table("demandes").select("*").execute().data
            _now = datetime.now()
            for _req in _fresh_demandes:
                # Services EXCLUS de l'auto-reply (premium obligatoire ou accord manuel)
                SERVICES_EXCLUS_AUTO = [
                    "📝 Exposé scolaire complet IA",
                ]
                if _req.get("service", "") in SERVICES_EXCLUS_AUTO:
                    continue  # Exposé = premium obligatoire ou accord manuel uniquement

                # Seulement les demandes des utilisateurs NON premium
                _req_user = _req.get("uid", "")
                _req_user_data = st.session_state["db"]["users"].get(_req_user, {})
                _is_premium_req = is_premium_actif(_req_user_data)
                if _is_premium_req:
                    continue  # Ignorer les premium
                # Vérifier si ça fait plus de 2 heures
                _ts_str = _req.get("timestamp", "")
                try:
                    _ts = datetime.fromisoformat(_ts_str)
                    _age_heures = (_now - _ts).total_seconds() / 3600
                except:
                    continue
                if _age_heures < 2:
                    continue  # Pas encore 2h
                # Vérifier que le service est supporté par Arsène AI
                _service_req = _req.get("service", "")
                if _service_req not in SERVICES_GEMINI:
                    continue
                # Vérifier pas déjà traité (status != "auto_done")
                if _req.get("status") == "auto_done":
                    continue
                # ── Générer automatiquement ──────────────────────────────
                _desc_req   = _req.get("description", "")
                _client_req = _req.get("uid", "")
                _wa_req     = _req.get("whatsapp", "")
                _req_id     = _req.get("id", "")
                _contenu_auto = generer_avec_gemini(_service_req, _desc_req, _client_req)
                if _contenu_auto.startswith("❌"):
                    continue  # Erreur API — on réessaiera au prochain refresh
                # Créer le fichier docx
                _buf_auto = creer_docx(_contenu_auto, _service_req, _client_req)
                _nom_auto = f"{_client_req}_{_service_req[:20].strip()}_auto.docx".replace(" ", "_").replace("/", "-")
                # Upload vers Supabase Storage
                _url_auto = upload_fichier_client(_client_req, _req_id, _buf_auto, _nom_auto)
                if _url_auto.startswith("ERREUR"):
                    continue
                # Sauvegarder le lien dans les livrables
                save_lien(_client_req, _service_req, _url_auto, _now.strftime("%d/%m/%Y"))
                # Marquer la demande comme traitée et la supprimer
                supabase.table("demandes").update({"status": "auto_done"}).eq("id", _req_id).execute()
                delete_demande(_req_id)
                # Email admin
                envoyer_email_auto_gratuit(_client_req, _wa_req, _service_req, _nom_auto, _desc_req)
        except Exception as _e_auto:
            pass  # Silencieux — ne pas perturber l'interface

    tab1, tab2 = st.tabs(["🚀 DÉPLOYER UNE TÂCHE", "📂 MES LIVRABLES (CLOUD)"])

    SERVICES_GEMINI = [
        "📝 Exposé scolaire complet IA",
        "📝 Création de Sujets & Examens",
        "📖 Fiche de Cours Professeur IA",
        "👔 CV & Lettre de Motivation",
        "⚙️ Pack Office (Word/Excel/PPT)",
        "📊 Data & Excel Analytics",
    ]

    with tab1:
        type_sujet_selectionne = None  # Initialisé ici, redéfini si service Sujets/Examens
        if st.session_state["premium_livrable"]:
            lv = st.session_state["premium_livrable"]
            st.markdown(f"""
            <div class="livrable-auto">
                <div class="livrable-auto-title">✅ Votre document est prêt !</div>
                <div style="color:rgba(255,255,255,.7);margin-top:6px;">
                    Généré en {lv['duree']}s · Service : <b>{lv['service']}</b>
                </div>
            </div>""", unsafe_allow_html=True)
            st.download_button(
                label="📥 TÉLÉCHARGER MON DOCUMENT",
                data=lv["buf"], file_name=lv["nom"], mime=lv["mime"],
                use_container_width=True
            )
            st.info("💡 Votre fichier est aussi disponible dans **📂 Mes Livrables** ci-dessus.")
            if st.button("🔄 Nouvelle mission", key="reset_livrable"):
                st.session_state["premium_livrable"] = None
                st.rerun()
            st.stop()

        SERVICE_PREREQUIS = {
            "📊 Data & Excel Analytics": {
                "icone": "📊",
                "titre": "Data & Excel Analytics",
                "intro": "Pour traiter vos données efficacement, merci de préciser dans votre cahier des charges :",
                "items": [
                    ("📁", "Le type de fichier à traiter (Excel, CSV, autre...)"),
                    ("🎯", "L'objectif de l'analyse (tableau de bord, graphiques, calculs...)"),
                    ("📋", "Une description des données ou colonnes présentes"),
                    ("🔢", "Le nombre approximatif de lignes ou d'entrées"),
                ],
                "note": "Vous pouvez également joindre un fichier exemple via WhatsApp après la soumission."
            },
            "⚙️ Pack Office (Word/Excel/PPT)": {
                "icone": "⚙️",
                "titre": "Pack Office — Word / Excel / PowerPoint",
                "intro": "Pour réaliser votre document Office au standard professionnel, précisez :",
                "items": [
                    ("📄", "Le type de document souhaité (Word, Excel ou PowerPoint)"),
                    ("🎯", "Le sujet ou contenu du document"),
                    ("📏", "Le nombre de pages ou diapositives souhaité"),
                    ("🎨", "Un style ou thème de couleurs si vous en avez un"),
                ],
                "note": "Précisez si vous souhaitez un logo ou une charte graphique particulière."
            },
            "🎨 Création Design IA": {
                "icone": "🎨",
                "titre": "Création Design IA",
                "intro": "Pour créer un design à la hauteur de vos attentes, indiquez-nous :",
                "items": [
                    ("🖼️", "Le type de visuel (affiche, bannière, logo, flyer...)"),
                    ("📐", "Le format ou dimensions souhaitées"),
                    ("🎨", "Les couleurs ou thème visuel préféré"),
                    ("✍️", "Les textes ou messages à intégrer"),
                    ("🏢", "Le nom de votre entreprise ou projet"),
                ],
                "note": "Une référence visuelle ou exemple que vous aimez accélérera le travail."
            },
            "📚 Affiches & Reçus": {
                "icone": "📚",
                "titre": "Affiches & Reçus",
                "intro": "Pour concevoir votre affiche ou reçu, nous aurons besoin de :",
                "items": [
                    ("🏢", "Le nom de votre entreprise ou organisation"),
                    ("📋", "Les informations à faire apparaître (prix, date, lieu, contacts...)"),
                    ("🎨", "La couleur principale ou identité visuelle"),
                    ("📐", "Le format désiré (A4, A5, reçu thermique...)"),
                ],
                "note": "Un logo ou image à intégrer peut être envoyé via WhatsApp."
            },
            "👔 CV & Lettre de Motivation": {
                "icone": "👔",
                "titre": "CV & Lettre de Motivation",
                "intro": "Pour rédiger votre CV ou lettre de motivation de façon percutante, fournissez :",
                "items": [
                    ("👤", "Votre nom complet et coordonnées"),
                    ("🎓", "Vos diplômes et formations"),
                    ("💼", "Vos expériences professionnelles"),
                    ("🎯", "Le poste ou secteur visé"),
                    ("✨", "Vos compétences clés et atouts"),
                ],
                "note": "Précisez si vous souhaitez uniquement le CV, la lettre, ou les deux."
            },
        }

        col_f, col_wa = st.columns(2)
        with col_f:
            st.markdown("#### 🛠️ Service Nova")
            service = st.selectbox(
                "Type d'intervention",
                [
                    "📊 Data & Excel Analytics",
                    "📖 Fiche de Cours Professeur IA",
                    "📎 Modifier mon Fichier (Word / Excel / PPT)",
                    "📝 Exposé scolaire complet IA",
                    "📝 Création de Sujets & Examens",
                    "⚙️ Pack Office (Word/Excel/PPT)",
                    "🎨 Création Design IA",
                    "📚 Affiches & Reçus",
                    "👔 CV & Lettre de Motivation",
                    "📄 Conversion & Fichier PDF",
                ]
            )
        with col_wa:
            st.markdown("#### 📞 Notification")
            default_wa = db["users"][user]["whatsapp"] if user else ""
            wa_display = st.text_input("WhatsApp de contact", value=default_wa, placeholder="225...")

        SERVICE_SAISIE = "📊 Data & Excel Analytics"

        if service != st.session_state["last_service_seen"]:
            st.session_state["last_service_seen"] = service
            st.session_state["warning_triggered"] = False
            st.session_state["show_service_warning"] = False
            if service != SERVICE_SAISIE and service in SERVICE_PREREQUIS:
                st.session_state["show_service_warning"] = True

        if st.session_state["show_service_warning"] and service in SERVICE_PREREQUIS and service != SERVICE_SAISIE:
            info = SERVICE_PREREQUIS[service]

            SERVICE_AUDIO = {
                "⚙️ Pack Office (Word/Excel/PPT)": "prerequis_office.mp3",
                "🎨 Création Design IA":           "prerequis_design.mp3",
                "📚 Affiches & Reçus":             "prerequis_affiches.mp3",
                "👔 CV & Lettre de Motivation":    "prerequis_cv.mp3",
            }

            st.info(f"""
**{info["icone"]} {info["titre"]} — Informations requises**

{info["intro"]}

{"".join(f"- {icone} {texte}\n" for icone, texte in info["items"])}
💡 *{info["note"]}*
""")
            audio_file = SERVICE_AUDIO.get(service)
            if audio_file and os.path.exists(audio_file):
                with open(audio_file, "rb") as f:
                    b64 = __import__('base64').b64encode(f.read()).decode()
                components.html(f"""
                    <script>
                    (function() {{
                        var binary = atob("{b64}");
                        var bytes = new Uint8Array(binary.length);
                        for (var i = 0; i < binary.length; i++) bytes[i] = binary.charCodeAt(i);
                        var blob = new Blob([bytes], {{type: "audio/mpeg"}});
                        var audio = new Audio(URL.createObjectURL(blob));
                        audio.volume = 1;
                        audio.play().catch(function(e) {{ console.log(e); }});
                    }})();
                    </script>
                """, height=0)
            col_mid = st.columns([1, 2, 1])[1]
            with col_mid:
                if st.button("✅ J'ai compris, je continue ma demande", key="close_service_warning"):
                    st.session_state["show_service_warning"] = False
                    components.html("<script>window.speechSynthesis.cancel();</script>", height=0)
                    st.rerun()

        # ── SÉLECTION DU TYPE DE SUJET (uniquement pour le service Sujets/Examens) ──
        type_sujet_selectionne = None
        if "Sujets" in service or "Examens" in service:
            st.markdown("#### 🎯 Type de sujet")
            TYPES_SUJETS = {
                "🔵 QCM — Questions à Choix Multiple": "QCM",
                "✅ Vrai ou Faux (avec justification)": "VRAI_FAUX",
                "🔤 Texte à Trous (lacunaire)": "TEXTE_TROU",
                "✍️ Questions Ouvertes (rédigées)": "QUESTIONS_OUVERTES",
                "🔀 Mixte (QCM + Vrai/Faux + Question ouverte)": "MIXTE",
                "📋 Cas Pratique / Étude de Cas": "CAS_PRATIQUE",
                "📐 Exercices de Calcul / Problèmes": "CALCUL",
                "🗺️ Étude de Document (texte, tableau, carte)": "ETUDE_DOCUMENT",
                "🔬 Schéma à Légender / Identification": "SCHEMA",
                "📝 Composition / Dissertation guidée": "DISSERTATION",
                "📄 Devoir Complet (comme sur les images)": "DEVOIR_COMPLET",
            }
            type_sujet_label = st.selectbox(
                "Choisissez le type d'exercice que vous voulez dans votre sujet",
                list(TYPES_SUJETS.keys()),
                help="Sélectionnez précisément le type de sujet souhaité. Arsène AI adaptera 100% du contenu à ce format."
            )
            type_sujet_selectionne = TYPES_SUJETS[type_sujet_label]

            TYPE_SUJET_DESCRIPTIONS = {
                "QCM": "**QCM sélectionné** — Arsène AI générera des questions à 4 choix (A/B/C/D) avec cases □ à cocher, distracteurs réalistes et corrigé si demandé.",
                "VRAI_FAUX": "**Vrai ou Faux sélectionné** — Arsène AI générera des affirmations à évaluer (V/F) avec lignes de justification pour les fausses réponses.",
                "TEXTE_TROU": "**Texte à trous sélectionné** — Arsène AI rédigera un texte cohérent avec des blancs à remplir et une liste de mots fournie.",
                "QUESTIONS_OUVERTES": "**Questions ouvertes sélectionnées** — Arsène AI formulera des questions de réflexion avec lignes de réponse proportionnelles au barème.",
                "MIXTE": "**Format Mixte sélectionné** — Arsène AI combinera QCM (Partie 1) + Vrai/Faux (Partie 2) + Question rédigée (Partie 3), barème équilibré.",
                "CAS_PRATIQUE": "**Cas Pratique sélectionné** — Arsène AI rédigera un texte/document contextualisé (Côte d'Ivoire) + questions d'analyse progressives.",
                "CALCUL": "**Exercices de Calcul sélectionnés** — Arsène AI rédigera des problèmes chiffrés contextualisés avec démarche guidée, formules rappelées et données réelles ivoiriennes.",
                "ETUDE_DOCUMENT": "**Étude de Document sélectionnée** — Arsène AI créera un document support (texte, tableau ou description de carte) + questions d'identification, analyse et interprétation.",
                "SCHEMA": "**Schéma à légender sélectionné** — Arsène AI décrira textuellement un schéma numéroté avec la liste des termes à placer et un corrigé de légendes.",
                "DISSERTATION": "**Dissertation guidée sélectionnée** — Arsène AI formulera un sujet de composition, fournira des consignes de méthode et proposera un plan détaillé guidé.",
                "DEVOIR_COMPLET": "**Devoir Complet sélectionné** — Arsène AI générera un vrai devoir ivoirien complet avec en-tête officiel + exercices variés progressifs (QCM → mise en situation → problème complexe) adaptés exactement au niveau et à la matière.",
            }
            st.info(TYPE_SUJET_DESCRIPTIONS.get(type_sujet_selectionne, ""))

        st.markdown("#### 📝 Spécifications de la mission")

        # Initialisations (re-évaluées à chaque run Streamlit)
        _niveau_val      = ""
        _matiere_val     = ""
        _fc_niveau_val   = ""
        _fc_matiere_val  = ""
        mf_fichier       = None
        mf_instructions  = ""
        conv_fichier     = None

        # ── FORMULAIRE STRUCTURÉ POUR SUJETS & EXAMENS ────────────────────────
        if "Sujets" in service or "Examens" in service:

            # ── TOGGLE FICHIER EN PREMIER ──────────────────────────────────
            st.markdown("""
            <style>
            @keyframes glowPulse2 {
                0%   { box-shadow: 0 0 6px 2px rgba(46,204,113,0.4), 0 0 12px 4px rgba(46,204,113,0.2); }
                50%  { box-shadow: 0 0 18px 6px rgba(46,204,113,0.9), 0 0 35px 12px rgba(46,204,113,0.4); }
                100% { box-shadow: 0 0 6px 2px rgba(46,204,113,0.4), 0 0 12px 4px rgba(46,204,113,0.2); }
            }
            @keyframes textPulse2 {
                0%   { opacity: 1; }
                50%  { opacity: 0.75; }
                100% { opacity: 1; }
            }
            .fichier-toggle-card2 {
                background: linear-gradient(135deg, rgba(46,204,113,0.12), rgba(39,174,96,0.08));
                border: 2px solid rgba(46,204,113,0.7);
                border-radius: 14px;
                padding: 16px 20px;
                margin: 0 0 6px 0;
                animation: glowPulse2 2s ease-in-out infinite;
            }
            .fichier-toggle-title2 {
                color: #2ecc71;
                font-weight: 800;
                font-size: 1.05rem;
                animation: textPulse2 2s ease-in-out infinite;
                display: block;
            }
            .fichier-toggle-sub2 {
                color: rgba(255,255,255,0.55);
                font-size: 0.8rem;
                margin-top: 4px;
                display: block;
            }
            </style>
            <div class="fichier-toggle-card2">
                <span class="fichier-toggle-title2">📂 ✨ Créer le sujet à partir d'un de mes fichiers</span>
                <span class="fichier-toggle-sub2">Importez votre cours ou leçon — Arsène AI génère un sujet basé uniquement sur votre document</span>
            </div>
            """, unsafe_allow_html=True)
            use_fichier_source = st.toggle(
                "✅ Activer cette option",
                key="use_fichier_source",
                help="Importez un fichier Word, PDF ou TXT — Arsène AI analysera son contenu."
            )

            if not use_fichier_source:
                st.markdown("""
                <div style="background:rgba(66,133,244,0.08);border:1px solid rgba(66,133,244,0.3);
                     border-radius:12px;padding:14px 18px;margin:14px 0;">
                    <span style="color:#4285f4;font-weight:700;">📋 Remplissez les champs ci-dessous — Nova s'appuie sur ces informations précises pour générer votre sujet</span>
                </div>
                """, unsafe_allow_html=True)
                col_a, col_b = st.columns(2)
            if not use_fichier_source:
                col_a, col_b = st.columns(2)
                with col_a:
                    exam_niveau = st.selectbox(
                        "🎓 Niveau scolaire *",
                        [
                            "── PRIMAIRE ──",
                            "CP1", "CP2", "CE1", "CE2", "CM1", "CM2 / CEPE",
                            "── COLLÈGE ──",
                            "6ème", "5ème", "4ème", "3ème / BEPC",
                            "── LYCÉE ──",
                            "2nde", "1ère - Série A1", "1ère - Série A2", "1ère - Série B",
                            "1ère - Série C", "1ère - Série D", "1ère - Série E",
                            "Terminale - Série A1", "Terminale - Série A2", "Terminale - Série B",
                            "Terminale - Série C", "Terminale - Série D", "Terminale - Série E",
                            "Terminale - Série F", "Terminale - Série G1", "Terminale - Série G2",
                            "Terminale - Série G3", "Terminale - Série H",
                            "── UNIVERSITÉ ──",
                            "Licence 1 (L1)", "Licence 2 (L2)", "Licence 3 (L3)",
                            "Master 1 (M1)", "Master 2 (M2)", "Doctorat",
                            "── CONCOURS ──",
                            "Concours ENS", "Concours CAFOP", "Concours INJS",
                            "Concours Fonction Publique", "Concours Douane / Police / Armée",
                            "Autre concours professionnel",
                        ],
                        index=0
                    )
                    exam_matiere = st.selectbox(
                        "📚 Matière / Discipline *",
                        [
                            "── TOUTES MATIÈRES ──",
                            "Français / Lettres", "Mathématiques",
                            "Sciences Physiques (PC)", "SVT / Biologie",
                            "Histoire-Géographie", "Économie / Gestion",
                            "Comptabilité", "Philosophie",
                            "EDHC / Éducation Civique",
                            "Anglais (LV1)", "Espagnol (LV2)", "Allemand (LV2)",
                            "Informatique / TIC",
                            "Technologie industrielle",
                            "EPS (Éducation Physique)",
                            "Arts Plastiques",
                            "Agronomie / Agriculture",
                            "Droit", "Économie politique",
                            "── PRIMAIRE ──",
                            "Lecture / Écriture (primaire)", "Calcul (primaire)",
                            "Sciences d'Éveil (primaire)",
                            "Histoire-Géo (primaire)", "ECM (primaire)",
                            "Autre matière (préciser dans les notes)",
                        ],
                        index=0
                    )
                with col_b:
                    exam_type_epreuve = st.selectbox(
                        "🎯 Type d'épreuve *",
                        [
                            "Devoir Surveillé (DS)", "Interrogation Écrite (IE)",
                            "Devoir de Maison (DM)", "Devoir du 1er Trimestre",
                            "Devoir du 2ème Trimestre", "Devoir du 3ème Trimestre",
                            "Examen Blanc / Blanc CEPE", "Examen Blanc / Brevet Blanc (BEPC)",
                            "Examen Blanc / BAC Blanc", "Épreuve de Rattrapage",
                            "Sujet de Concours", "Épreuve de Passage",
                            "Exercice de classe (rapide)",
                        ],
                        index=0
                    )
                    exam_duree = st.selectbox(
                        "⏱️ Durée prévue *",
                        [
                            "15 minutes", "30 minutes", "1 heure",
                            "1 heure 30", "2 heures", "2 heures 30",
                            "3 heures", "3 heures 30", "4 heures",
                            "Durée libre (DM / à domicile)",
                        ],
                        index=2
                    )
                col_c, col_d = st.columns(2)
                with col_c:
                    exam_nb_exercices = st.selectbox(
                        "📏 Nombre d'exercices / questions *",
                        [
                            "1 exercice", "2 exercices", "3 exercices",
                            "4 exercices", "5 exercices",
                            "10 questions", "15 questions", "20 questions",
                            "25 questions", "30 questions",
                            "Adapté automatiquement au niveau et à la durée",
                        ],
                        index=10
                    )
                with col_d:
                    exam_coefficient = st.selectbox(
                        "🔢 Coefficient",
                        ["1", "2", "3", "4", "5", "6", "7", "8"],
                        index=1
                    )
                col_e, col_f = st.columns(2)
                with col_e:
                    exam_etablissement = st.text_input(
                        "🏢 Établissement / Institution",
                        placeholder="Ex: Lycée Moderne de Cocody, CEG Treichville, UFHB..."
                    )
                with col_f:
                    exam_annee = st.text_input(
                        "📅 Année scolaire",
                        placeholder="Ex: 2024-2025",
                        value="2024-2025"
                    )
                exam_chapitre = st.text_input(
                    "📖 Chapitre / Notion spécifique (optionnel)",
                    placeholder="Ex: Les fractions, La cellule, La colonisation, La dérivation, La loi d'Ohm..."
                )
                exam_notes = st.text_area(
                    "💬 Informations complémentaires (optionnel)",
                    height=80,
                    placeholder="Ex: Avec corrigé, thème ivoirien, niveau difficile, chapitres 1 et 2, 4 QCM + 2 ouvertes..."
                )
            else:
                # Réinitialiser le contenu fichier si toggle désactivé
                st.session_state["contenu_fichier_source"] = ""
                contenu_fichier_source = ""
                # Valeurs par défaut quand fichier actif
                exam_niveau = "Non précisé"
                exam_matiere = "── TOUTES MATIÈRES ──"
                exam_type_epreuve = "Devoir Surveillé (DS)"
                exam_duree = "1 heure"
                exam_nb_exercices = "Adapté automatiquement au niveau et à la durée"
                exam_coefficient = "2"
                exam_etablissement = ""
                exam_annee = "2024-2025"
                exam_chapitre = ""
                exam_notes = ""

            contenu_fichier_source = ""
            if use_fichier_source:
                st.markdown('''
                <div style="background:rgba(46,204,113,0.08);border:1px solid rgba(46,204,113,0.3);
                     border-radius:10px;padding:12px 16px;margin:8px 0;">
                    <span style="color:#2ecc71;font-weight:700;">📄 Arsène AI va scanner votre document et créer le sujet uniquement à partir de son contenu</span>
                    <span style="color:rgba(255,255,255,.5);font-size:.82rem;display:block;margin-top:3px;">
                        Formats acceptés : PDF, Word (.docx), Texte (.txt)
                    </span>
                </div>
                ''', unsafe_allow_html=True)

                fichier_source = st.file_uploader(
                    "📂 Importer votre fichier *",
                    type=["pdf", "docx", "txt"],
                    key="fichier_source_exam"
                )

                if fichier_source:
                    with st.spinner("🔍 Lecture du fichier en cours..."):
                        try:
                            import io
                            fichier_bytes = fichier_source.read()

                            if fichier_source.name.endswith(".txt"):
                                contenu_fichier_source = fichier_bytes.decode("utf-8", errors="ignore")

                            elif fichier_source.name.endswith(".docx"):
                                from docx import Document as DocxDoc
                                doc_tmp = DocxDoc(io.BytesIO(fichier_bytes))
                                contenu_fichier_source = "\n".join([p.text for p in doc_tmp.paragraphs if p.text.strip()])

                            elif fichier_source.name.endswith(".pdf"):
                                try:
                                    import fitz  # PyMuPDF
                                    pdf_doc = fitz.open(stream=fichier_bytes, filetype="pdf")
                                    contenu_fichier_source = "\n".join([page.get_text() for page in pdf_doc])
                                except ImportError:
                                    import subprocess
                                    subprocess.run(["pip", "install", "PyMuPDF", "--break-system-packages", "-q"])
                                    import fitz
                                    pdf_doc = fitz.open(stream=fichier_bytes, filetype="pdf")
                                    contenu_fichier_source = "\n".join([page.get_text() for page in pdf_doc])

                            # Tronquer si trop long (max ~6000 caractères pour le prompt)
                            if len(contenu_fichier_source) > 6000:
                                contenu_fichier_source = contenu_fichier_source[:6000] + "\n...[document tronqué]"

                            if contenu_fichier_source.strip():
                                st.session_state["contenu_fichier_source"] = contenu_fichier_source
                                st.success(f"✅ **{fichier_source.name}** lu avec succès — {len(contenu_fichier_source)} caractères extraits")
                            else:
                                st.warning("⚠️ Le fichier semble vide ou illisible")
                                contenu_fichier_source = ""
                                st.session_state["contenu_fichier_source"] = ""

                        except Exception as e_fic:
                            st.error(f"❌ Erreur lecture fichier : {e_fic}")
                            contenu_fichier_source = ""
                            st.session_state["contenu_fichier_source"] = ""
                else:
                    # Récupérer depuis session_state si déjà chargé
                    if st.session_state["contenu_fichier_source"]:
                        contenu_fichier_source = st.session_state["contenu_fichier_source"]
                        st.success(f"✅ Fichier en mémoire — {len(contenu_fichier_source)} caractères")
                    else:
                        st.info("⬆️ Importez votre fichier pour que Nova génère le sujet à partir de son contenu")

                # ── CAHIER DES CHARGES SPÉCIFIQUE FICHIER ─────────────────────
                st.markdown("""
                <div style="background:rgba(255,165,0,0.08);border:1px solid rgba(255,165,0,0.4);
                     border-radius:10px;padding:12px 16px;margin:10px 0 4px 0;">
                    <span style="color:#FFA500;font-weight:700;">📋 Instructions pour le sujet</span>
                    <span style="color:rgba(255,255,255,.5);font-size:.8rem;display:block;margin-top:2px;">
                        Arsène AI se base <b>uniquement sur votre fichier</b> — précisez ici ce que vous voulez comme sujet
                    </span>
                </div>
                """, unsafe_allow_html=True)
                fichier_instructions = st.text_area(
                    "✍️ Décrivez le sujet que vous voulez *",
                    height=110,
                    placeholder="Ex: Fais un QCM de 10 questions sur les définitions...\nEx: Crée un devoir avec 2 exercices de calcul basés sur les formules du document...\nEx: Génère une interrogation sur la partie 2 du cours uniquement...",
                    key="fichier_instructions"
                )

            # ── CONSTRUCTION AUTOMATIQUE DU PROMPT STRUCTURÉ ──────────────────
            _niveau_val = exam_niveau if not exam_niveau.startswith("──") else ""
            _matiere_val = exam_matiere if not exam_matiere.startswith("──") else ""

            _etab_val = exam_etablissement.strip() if exam_etablissement.strip() else "Établissement non précisé"
            _annee_val = exam_annee.strip() if exam_annee.strip() else "2024-2025"
            _fichier_instr = fichier_instructions.strip() if use_fichier_source and "fichier_instructions" in dir() else ""

            prompt = f"""FICHE DE COMMANDE NOVA EXAM — INFORMATIONS STRUCTURÉES :

🎓 NIVEAU SCOLAIRE     : {_niveau_val if _niveau_val else "Non précisé"}
📚 MATIÈRE             : {_matiere_val if _matiere_val else "Non précisée"}
🎯 TYPE D'ÉPREUVE      : {exam_type_epreuve}
📏 EXERCICES/QUESTIONS : {exam_nb_exercices}
⏱️ DURÉE               : {exam_duree}
🔢 COEFFICIENT         : {exam_coefficient}
🏢 ÉTABLISSEMENT       : {_etab_val}
📅 ANNÉE SCOLAIRE      : {_annee_val}
📖 CHAPITRE/NOTION     : {"BASÉ SUR LE FICHIER FOURNI — voir document source ci-dessous" if use_fichier_source else (exam_chapitre if exam_chapitre.strip() else "Choisir un chapitre cohérent avec le programme officiel du niveau")}
💬 INSTRUCTIONS CLIENT : {"" + _fichier_instr if use_fichier_source else (exam_notes.strip() if exam_notes.strip() else "Aucune")}

INSTRUCTIONS NOVA EXAM :
- Respecte EXACTEMENT le niveau "{_niveau_val}" — applique le programme officiel MENET-FP de cette classe
- Génère UNIQUEMENT des notions au programme de ce niveau — rien hors-programme
- Adapte le vocabulaire, la complexité et la longueur à l'âge de l'élève de ce niveau
- Si un chapitre/notion est précisé, le sujet porte EXCLUSIVEMENT sur ce chapitre
- Si "avec corrigé" dans les notes, inclure le corrigé complet après ---SAUT_DE_PAGE---

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SECTION CRITIQUE — FORMAT AUTHENTIQUE DES DEVOIRS IVOIRIENS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

① EN-TÊTE OBLIGATOIRE — Reproduis ce modèle EXACTEMENT :

{_etab_val}* {_etab_val}*
Année Scolaire : {_annee_val}                    Coefficient : {exam_coefficient}
Classe : {_niveau_val if _niveau_val else "___"}                              Durée : {exam_duree}
[ligne vide]
┌─────────────────────────────────────────────┐
│    DEVOIR DE {(_matiere_val.upper() if _matiere_val else "___")}    │
└─────────────────────────────────────────────┘


RÈGLES EN-TÊTE :
- Nom établissement répété 2 fois sur la même ligne séparé par *
- Année + Coefficient sur la MÊME ligne (gauche / droite)
- Classe + Durée sur la MÊME ligne (gauche / droite)
- Titre DEVOIR encadré ou en gras centré
- NE JAMAIS ajouter REPUBLIQUE, MINISTERE, DIRECTION non demandés

② STRUCTURE DES EXERCICES — ADAPTE-TOI INTELLIGEMMENT AU NIVEAU DÉTECTÉ :

Le niveau demandé est : {_niveau_val}

Tu es un expert du système éducatif ivoirien (MENET-FP). Tu connais parfaitement comment les devoirs sont structurés à CHAQUE niveau en Côte d'Ivoire. Adapte la structure, la complexité, le vocabulaire et les types d'exercices exactement comme le ferait un vrai professeur ivoirien de ce niveau.

▶ PRIMAIRE (CP1 → CM2) :
- Exercices courts, simples, directs
- Pas de mise en situation complexe
- Dictée, calcul mental, lecture, problèmes simples du quotidien ivoirien
- Consignes en langage simple accessible à l'enfant
- 3 à 4 exercices maximum, jamais de sous-questions complexes

▶ COLLÈGE — 6ème / 5ème :
- Exercice 1 : QCM ou Vrai/Faux ou tableau à compléter (facile)
- Exercice 2 : Définitions, propriétés, applications directes du cours
- Exercice 3 : Problème avec mise en situation simple, questions progressives 1- 2- 3-
- Exercice 4 (si demandé) : Problème de la vie courante ivoirienne, niveau accessible
- Langage simple, données concrètes, pas de démonstrations formelles

▶ COLLÈGE — 4ème / 3ème (BEPC) :
- Exercice 1 : QCM / Vrai-Faux / affirmations V ou F (mise en route)
- Exercice 2 : Questions courtes directes (définitions, calculs simples, compléter)
- Exercice 3 : TEXTE DE MISE EN SITUATION + 4-5 questions progressives avec sous-questions (1.1 / 1.2 / 1.3)
- Exercice 4 : PROBLÈME COMPLEXE ancré dans la réalité ivoirienne (PME, plantation, lycée CI) — démonstrations, calculs multi-étapes
- Barème : ex1-2 = 30%, ex3-4 = 70% des points

▶ LYCÉE — 2nde / 1ère :
- Structure plus libre, moins de QCM, plus de rédaction et démonstration
- Exercice 1 : Questions de cours ou vérification des connaissances (définitions, théorèmes)
- Exercice 2 : Application directe, calculs guidés
- Exercice 3 : Problème ouvert avec situation complexe, plusieurs sous-parties (A, B, C)
- Exercice 4 : Problème de synthèse liant plusieurs notions du programme
- Attendu : justifications rigoureuses, raisonnement logique développé

▶ LYCÉE — TERMINALE / BAC :
- Format proche des épreuves officielles du BAC ivoirien
- 3 à 4 exercices longs avec parties A / B / C
- Chaque partie autonome mais liée au thème général
- Calculs complexes, démonstrations formelles, analyse critique
- Problèmes souvent en contexte scientifique ou socio-économique ivoirien
- Barème précis sur /20, questions numérotées I- II- III- ou 1) 2) 3)

▶ POST-BAC (BTS, Licence, Master) :
- Format académique supérieur : énoncé dense, problème unique divisé en parties
- Partie I / Partie II / Partie III avec sous-questions a) b) c)
- Niveau de rigueur élevé, démonstrations formelles exigées
- Mise en contexte professionnel ou scientifique sérieux
- Pas de QCM — uniquement questions ouvertes et problèmes

RÈGLES GÉNÉRALES POUR TOUS LES NIVEAUX :
- Chaque exercice : ## EXERCICE N : (X points)
- Données toujours précises avec vrais chiffres
- Contextes ivoiriens authentiques : noms ivoiriens, villes CI (Abidjan, Bouaké, Yamoussoukro...), produits locaux (cacao, café, anacarde...), monnaie FCFA
- NE JAMAIS inventer une structure générique — colle au vrai format du niveau
"""
            # Injecter le contenu du fichier source si fourni
            if contenu_fichier_source.strip():
                prompt += """
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🎯 DIRECTIVE PRINCIPALE — CE QUE LE CLIENT VEUT EXACTEMENT :
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
""" + (_fichier_instr if _fichier_instr else "Crée un sujet adapté au niveau et à la durée indiqués, basé sur le document ci-dessous.") + """

⚠️ RÈGLE ABSOLUE : Exécute EXACTEMENT la demande ci-dessus. 
Le document ci-dessous est ta SEULE source de contenu — tu puises les notions, définitions et formules UNIQUEMENT dans ce texte. N'invente rien hors du document.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📄 DOCUMENT SOURCE (contenu à utiliser) :
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
""" + contenu_fichier_source + """
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""
            # Afficher un résumé de la commande
            if use_fichier_source:
                # Mode fichier : validation basée sur le fichier + instructions
                if contenu_fichier_source.strip() and _fichier_instr.strip():
                    st.success(f"✅ Commande prête — sujet généré à partir de votre fichier")
                elif not contenu_fichier_source.strip():
                    st.warning("⚠️ Importez votre fichier pour continuer")
                else:
                    st.warning("⚠️ Décrivez le sujet que vous voulez dans le cahier des charges")
            else:
                if _niveau_val and _matiere_val and not _niveau_val.startswith("──") and not _matiere_val.startswith("──"):
                    st.success(f"✅ Commande prête : **{_matiere_val}** · **{_niveau_val}** · **{exam_type_epreuve}** · **{exam_duree}**")
                else:
                    if not _niveau_val or _niveau_val.startswith("──"):
                        st.warning("⚠️ Sélectionnez un niveau scolaire précis (pas le titre de catégorie)")
                    if not _matiere_val or _matiere_val.startswith("──"):
                        st.warning("⚠️ Sélectionnez une matière précise (pas le titre de catégorie)")

        elif "Exposé" in service:
            st.markdown('''
            <div style="background:rgba(255,165,0,0.08);border:1px solid rgba(255,165,0,0.35);
                 border-radius:12px;padding:14px 18px;margin-bottom:14px;">
                <span style="font-weight:700;color:#FFA500;">📝 Remplissez les champs — Nova génère votre exposé complet, structuré et prêt à présenter</span>
            </div>
            ''', unsafe_allow_html=True)

            col_a, col_b = st.columns(2)
            with col_a:
                exp_niveau = st.selectbox("🎓 Niveau scolaire *", [
                    "── PRIMAIRE ──","CP1","CP2","CE1","CE2","CM1","CM2",
                    "── COLLÈGE ──","6ème","5ème","4ème","3ème",
                    "── LYCÉE ──","2nde",
                    "1ère - Série A1","1ère - Série A2","1ère - Série B",
                    "1ère - Série C","1ère - Série D","1ère - Série E",
                    "Terminale - Série A1","Terminale - Série A2","Terminale - Série B",
                    "Terminale - Série C","Terminale - Série D","Terminale - Série E",
                    "── UNIVERSITÉ / BTS ──",
                    "Licence 1 (L1)","Licence 2 (L2)","Licence 3 (L3)",
                    "Master 1 (M1)","Master 2 (M2)",
                    "BTS 1ère année","BTS 2ème année",
                ], index=0, key="exp_niveau")
                exp_matiere = st.selectbox("📚 Matière / Discipline *", [
                    "── TOUTES MATIÈRES ──",
                    "Français / Lettres","Mathématiques","Sciences Physiques (PC)","SVT / Biologie",
                    "Histoire-Géographie","Économie / Gestion","Comptabilité","Philosophie",
                    "EDHC / Éducation Civique","Anglais (LV1)","Espagnol (LV2)","Allemand (LV2)",
                    "Informatique / TIC","Technologie industrielle","EPS","Arts Plastiques",
                    "Agronomie / Agriculture","Droit","Économie politique",
                    "Sciences de la Vie et de la Terre","Chimie","Physique",
                    "Autre (préciser dans les notes)",
                ], index=0, key="exp_matiere")
            with col_b:
                exp_pages = st.selectbox("📏 Nombre de pages souhaité *", [
                    "Adapté automatiquement au niveau",
                    "2 - 3 pages (court)",
                    "4 - 5 pages (standard collège)",
                    "6 - 8 pages (standard lycée)",
                    "8 - 12 pages (université)",
                    "12 - 20 pages (mémoire / rapport)",
                ], index=0, key="exp_pages")
                exp_langue = st.selectbox("🌍 Langue de rédaction *", [
                    "Français","Anglais","Français + résumé en anglais",
                ], index=0, key="exp_langue")

            exp_sujet = st.text_input("🎯 Thème / Sujet de l'exposé *",
                placeholder="Ex: Le changement climatique en Afrique, La Révolution française, L'ADN et la génétique...",
                key="exp_sujet")

            col_c, col_d = st.columns(2)
            with col_c:
                exp_type = st.selectbox("📄 Type d'exposé *", [
                    "Exposé scolaire classique (Introduction + Développement + Conclusion)",
                    "Rapport de recherche structuré (Parties I, II, III)",
                    "Exposé oral (plan + fiches de présentation)",
                    "Compte-rendu de TP / Expérience scientifique",
                    "Commentaire de texte / document",
                    "Dissertation guidée",
                ], index=0, key="exp_type")
            with col_d:
                exp_etablissement = st.text_input("🏢 Établissement (optionnel)",
                    placeholder="Ex: Lycée Moderne de Cocody, Université FHB...",
                    key="exp_etablissement")

            exp_notes = st.text_area("💬 Instructions complémentaires (optionnel)", height=70,
                placeholder="Ex: Insister sur le rôle de la Côte d'Ivoire, Niveau très basique, Inclure des statistiques récentes...",
                key="exp_notes")

            # Construire le prompt
            _exp_niveau_val  = exp_niveau  if not exp_niveau.startswith("──")  else ""
            _exp_matiere_val = exp_matiere if not exp_matiere.startswith("──") else ""
            prompt = f"""FICHE DE COMMANDE NOVA EXPOSÉ :
🎯 SUJET            : {exp_sujet.strip() or "Non précisé"}
🎓 NIVEAU           : {_exp_niveau_val or "Non précisé"}
📚 MATIÈRE          : {_exp_matiere_val or "Non précisée"}
📄 TYPE             : {exp_type}
📏 PAGES            : {exp_pages}
🌍 LANGUE           : {exp_langue}
🏢 ÉTABLISSEMENT    : {exp_etablissement.strip() or "Non précisé"}
💬 INSTRUCTIONS     : {exp_notes.strip() or "Aucune"}
"""
            if _exp_niveau_val and _exp_matiere_val and exp_sujet.strip():
                st.success(f"✅ Commande prête : **{exp_sujet.strip()[:45]}** · **{_exp_matiere_val}** · **{_exp_niveau_val}**")
            elif exp_sujet.strip():
                st.info("💡 Précisez le niveau et la matière pour un meilleur résultat")
            else:
                if not exp_sujet.strip(): st.warning("⚠️ Entrez le thème / sujet de l'exposé")
                if not _exp_niveau_val:   st.warning("⚠️ Sélectionnez un niveau scolaire précis")
                if not _exp_matiere_val:  st.warning("⚠️ Sélectionnez une matière précise")

        elif "Fiche de Cours" in service:
            st.markdown('''
            <div style="background:rgba(155,89,182,0.08);border:1px solid rgba(155,89,182,0.35);
                 border-radius:12px;padding:14px 18px;margin-bottom:14px;">
                <span style="color:#9b59b6;font-weight:700;">📖 Remplissez les champs — Nova génère une fiche de cours complète utilisable directement en classe</span>
            </div>
            ''', unsafe_allow_html=True)
            col_a, col_b = st.columns(2)
            with col_a:
                fc_niveau = st.selectbox("🎓 Niveau de la classe *", [
                    "── PRIMAIRE ──","CP1","CP2","CE1","CE2","CM1","CM2",
                    "── COLLÈGE ──","6ème","5ème","4ème","3ème",
                    "── LYCÉE ──","2nde",
                    "1ère - Série A1","1ère - Série A2","1ère - Série B",
                    "1ère - Série C","1ère - Série D","1ère - Série E",
                    "Terminale - Série A1","Terminale - Série A2","Terminale - Série B",
                    "Terminale - Série C","Terminale - Série D","Terminale - Série E",
                    "── UNIVERSITÉ / BTS ──",
                    "Licence 1 (L1)","Licence 2 (L2)","Licence 3 (L3)",
                    "Master 1 (M1)","Master 2 (M2)",
                    "BTS 1ère année","BTS 2ème année",
                ], index=0, key="fc_niveau")
                fc_matiere = st.selectbox("📚 Matière *", [
                    "── TOUTES MATIÈRES ──",
                    "Français / Lettres","Mathématiques","Sciences Physiques (PC)","SVT / Biologie",
                    "Histoire-Géographie","Économie / Gestion","Comptabilité","Philosophie",
                    "EDHC / Éducation Civique","Anglais (LV1)","Espagnol (LV2)","Allemand (LV2)",
                    "Informatique / TIC","Technologie industrielle","EPS","Arts Plastiques",
                    "Agronomie / Agriculture","Droit","Économie politique",
                    "── PRIMAIRE ──",
                    "Lecture / Écriture (primaire)","Calcul (primaire)",
                    "Sciences d'Éveil (primaire)","Histoire-Géo (primaire)","ECM (primaire)",
                    "Autre matière (préciser dans les notes)",
                ], index=0, key="fc_matiere")
            with col_b:
                fc_duree = st.selectbox("⏱️ Durée de la séance *", [
                    "30 minutes","45 minutes","1 heure",
                    "1 heure 30","2 heures","3 heures",
                    "Cours complet (plusieurs séances)",
                ], index=2, key="fc_duree")
                fc_type = st.selectbox("📄 Type de document *", [
                    "Fiche de cours complète (objectifs + développement + exercices)",
                    "Cours magistral (contenu seul, très détaillé)",
                    "Plan de leçon (structure + timing + activités)",
                    "Fiche de révision élève (synthèse dense)",
                    "Cours + Exercices corrigés",
                    "Fiche méthode (procédure étape par étape)",
                ], index=0, key="fc_type")
            col_c, col_d = st.columns(2)
            with col_c:
                fc_chapitre = st.text_input("📖 Titre du chapitre / Notion *",
                    placeholder="Ex: Les fractions, La photosynthèse, La colonisation...",
                    key="fc_chapitre")
            with col_d:
                fc_etablissement = st.text_input("🏢 Établissement (optionnel)",
                    placeholder="Ex: Lycée Moderne de Cocody...", key="fc_etablissement")
            fc_inclure = st.multiselect("✨ Éléments à inclure", [
                "Objectifs pédagogiques détaillés","Prérequis des élèves",
                "Exercices d'application","Corrigé complet des exercices",
                "Évaluation formative (fin de séance)","Résumé / Synthèse à retenir",
                "Exemples contextualisés (Côte d'Ivoire)","Définitions des termes clés",
                "Référence au manuel officiel MENET-FP","Tableau récapitulatif",
            ], default=[
                "Objectifs pédagogiques détaillés","Exercices d'application",
                "Corrigé complet des exercices","Exemples contextualisés (Côte d'Ivoire)",
                "Définitions des termes clés","Résumé / Synthèse à retenir",
            ], key="fc_inclure")
            fc_notes = st.text_area("💬 Instructions supplémentaires (optionnel)", height=70,
                placeholder="Ex: Niveau très faible — simplifier, Inclure exemples en FCFA...",
                key="fc_notes")
            _fc_niveau_val  = fc_niveau  if not fc_niveau.startswith("──")  else ""
            _fc_matiere_val = fc_matiere if not fc_matiere.startswith("──") else ""
            _fc_inclure_str = ", ".join(fc_inclure) if fc_inclure else "Éléments standards"
            prompt = f"""FICHE DE COMMANDE NOVA COURS :
🎓 NIVEAU          : {_fc_niveau_val or "Non précisé"}
📚 MATIÈRE         : {_fc_matiere_val or "Non précisée"}
📖 CHAPITRE        : {fc_chapitre.strip() or "Choisir un chapitre cohérent"}
📄 TYPE DOCUMENT   : {fc_type}
⏱️ DURÉE           : {fc_duree}
🏢 ÉTABLISSEMENT   : {fc_etablissement.strip() or "Non précisé"}
✨ ÉLÉMENTS        : {_fc_inclure_str}
💬 INSTRUCTIONS    : {fc_notes.strip() or "Aucune"}
"""
            if _fc_niveau_val and _fc_matiere_val and fc_chapitre.strip():
                st.success(f"✅ Fiche prête : **{fc_chapitre.strip()[:40]}** · **{_fc_matiere_val}** · **{_fc_niveau_val}**")
            elif _fc_niveau_val and _fc_matiere_val:
                st.info("💡 Précisez le titre du chapitre pour un meilleur résultat")
            else:
                if not _fc_niveau_val: st.warning("⚠️ Sélectionnez un niveau scolaire précis")
                if not _fc_matiere_val: st.warning("⚠️ Sélectionnez une matière précise")

        elif "Modifier" in service and "Fichier" in service:
            st.markdown('''
            <div style="background:rgba(0,210,255,0.07);border:1px solid rgba(0,210,255,0.3);
                 border-radius:12px;padding:14px 18px;margin-bottom:14px;">
                <span style="color:#00d2ff;font-weight:700;">📎 Importez votre fichier et décrivez vos modifications — notre équipe s'en charge</span>
            </div>
            ''', unsafe_allow_html=True)
            mf_fichier = st.file_uploader("📂 Importer votre fichier *",
                type=["docx","xlsx","xls","pptx","ppt","doc","pdf","csv","txt"],
                help="Formats : Word, Excel, PowerPoint, PDF, CSV, TXT",
                key="mf_fichier")
            if mf_fichier:
                st.success(f"✅ **{mf_fichier.name}** · {round(mf_fichier.size/1024,1)} Ko")
            mf_type_modif = st.selectbox("🛠️ Type de modification *", [
                "Correction et amélioration du contenu",
                "Mise en forme / Design professionnel",
                "Ajout de données / contenu supplémentaire",
                "Restructuration complète du document",
                "Traduction (Français ↔ Anglais)",
                "Fusion de plusieurs documents",
                "Extraction et réorganisation des données",
                "Création de graphiques / tableaux depuis les données",
                "Correction orthographique et grammaticale",
                "Adapter à un nouveau modèle / template",
                "Autre modification (préciser dans les instructions)",
            ], key="mf_type_modif")
            mf_instructions = st.text_area("✍️ Décrivez précisément ce que vous voulez *",
                height=130,
                placeholder="Ex: Mettre le tableau en page 3 en format auto, ajouter colonne Total avec formule, corriger les fautes...",
                key="mf_instructions")
            mf_urgence = st.selectbox("⏱️ Délai souhaité", [
                "Dès que possible (standard)",
                "Urgent — dans les 2 heures",
                "Très urgent — dans 1 heure",
            ], key="mf_urgence")
            _mf_nom = mf_fichier.name if mf_fichier else "Aucun fichier"
            prompt = f"""FICHE MODIFICATION FICHIER :
📎 FICHIER         : {_mf_nom}
🛠️ MODIFICATION    : {mf_type_modif}
⏱️ DÉLAI           : {mf_urgence}
✍️ INSTRUCTIONS    :
{mf_instructions.strip() or "(aucune instruction)"}
NOTE : fichier original joint via lien ci-dessous.
"""
            if mf_fichier and mf_instructions.strip():
                st.success("✅ Demande complète — notre équipe s'en charge")
            else:
                if not mf_fichier: st.warning("⚠️ Importez votre fichier")
                if not mf_instructions.strip(): st.warning("⚠️ Décrivez les modifications souhaitées")

        elif "Conversion" in service:
            # ════════════════════════════════════════════════════════════════
            # SERVICE CONVERSION — 100% Python, pas de Gemini, pas de WhatsApp
            # ════════════════════════════════════════════════════════════════
            st.markdown('''
            <div style="background:rgba(46,204,113,0.08);border:1px solid rgba(46,204,113,0.35);
                 border-radius:12px;padding:14px 18px;margin-bottom:14px;">
                <span style="font-weight:700;color:#2ecc71;">🔄 Conversion instantanée — 100% automatique, aucune attente</span>
                <span style="color:rgba(255,255,255,.5);font-size:.82rem;display:block;margin-top:4px;">
                    Importez votre fichier ci-dessous, choisissez le format de sortie — votre fichier est prêt en quelques secondes.
                </span>
            </div>
            ''', unsafe_allow_html=True)

            conv_type = st.selectbox("🔄 Type de conversion *", [
                "📝 Word (.docx) → PDF",
                "📊 Excel (.xlsx) → PDF",
                "📊 CSV → Excel (.xlsx)",
                "📊 Excel (.xlsx) → CSV",
                "🎞️ PowerPoint (.pptx) → PDF",
                "📄 PDF → Word (.docx)",
                "📄 PDF → PowerPoint (.pptx)",
            ], key="conv_type")

            # Définir les types acceptés selon la conversion choisie
            if "Word" in conv_type and "→ PDF" in conv_type:
                types_acceptes = ["docx", "doc"]
                label_upload   = "📂 Importer votre fichier Word (.docx)"
                suffix_in      = ".docx"
            elif "Excel" in conv_type and "→ PDF" in conv_type:
                types_acceptes = ["xlsx", "xls"]
                label_upload   = "📂 Importer votre fichier Excel (.xlsx)"
                suffix_in      = ".xlsx"
            elif "CSV → Excel" in conv_type:
                types_acceptes = ["csv"]
                label_upload   = "📂 Importer votre fichier CSV"
                suffix_in      = ".csv"
            elif "Excel" in conv_type and "→ CSV" in conv_type:
                types_acceptes = ["xlsx", "xls"]
                label_upload   = "📂 Importer votre fichier Excel (.xlsx)"
                suffix_in      = ".xlsx"
            elif "PowerPoint" in conv_type and "→ PDF" in conv_type:
                types_acceptes = ["pptx", "ppt"]
                label_upload   = "📂 Importer votre fichier PowerPoint (.pptx)"
                suffix_in      = ".pptx"
            elif "PDF → Word" in conv_type:
                types_acceptes = ["pdf"]
                label_upload   = "📂 Importer votre fichier PDF"
                suffix_in      = ".pdf"
            else:  # PDF → PPT
                types_acceptes = ["pdf"]
                label_upload   = "📂 Importer votre fichier PDF"
                suffix_in      = ".pdf"

            conv_fichier = st.file_uploader(label_upload, type=types_acceptes, key="conv_fichier")

            if conv_fichier:
                st.success(f"✅ **{conv_fichier.name}** · {round(conv_fichier.size/1024, 1)} Ko importé")

                if st.button("⚡ CONVERTIR MAINTENANT", use_container_width=True, key="btn_convertir"):
                    with st.spinner("🔄 Conversion en cours..."):
                        try:
                            import tempfile, os as _os, subprocess as _sub
                            fichier_bytes = conv_fichier.read()

                            # ── Word / Excel / PPT → PDF via LibreOffice ─────
                            if "→ PDF" in conv_type:
                                with tempfile.NamedTemporaryFile(suffix=suffix_in, delete=False) as tmp_in:
                                    tmp_in.write(fichier_bytes)
                                    tmp_in_path = tmp_in.name
                                tmp_out_dir = tempfile.mkdtemp()
                                res = _sub.run(
                                    ["libreoffice", "--headless", "--convert-to", "pdf",
                                     "--outdir", tmp_out_dir, tmp_in_path],
                                    capture_output=True, text=True, timeout=90
                                )
                                nom_base = _os.path.splitext(_os.path.basename(tmp_in_path))[0]
                                pdf_path = _os.path.join(tmp_out_dir, nom_base + ".pdf")
                                if _os.path.exists(pdf_path):
                                    with open(pdf_path, "rb") as f:
                                        buf_out = f.read()
                                    nom_sortie = conv_fichier.name.rsplit(".", 1)[0] + ".pdf"
                                    st.success("✅ Conversion réussie !")
                                    st.download_button("📥 TÉLÉCHARGER LE PDF", data=buf_out,
                                        file_name=nom_sortie, mime="application/pdf",
                                        use_container_width=True, key="dl_pdf_out")
                                else:
                                    st.error(f"❌ Erreur LibreOffice : {res.stderr or res.stdout}")
                                try: _os.unlink(tmp_in_path)
                                except: pass

                            # ── CSV → Excel ──────────────────────────────────
                            elif "CSV → Excel" in conv_type:
                                import pandas as pd
                                from io import BytesIO as _BytesIO
                                df = pd.read_csv(_BytesIO(fichier_bytes))
                                buf_out = _BytesIO()
                                df.to_excel(buf_out, index=False, engine="openpyxl")
                                buf_out.seek(0)
                                nom_sortie = conv_fichier.name.replace(".csv", "") + ".xlsx"
                                st.success("✅ Conversion réussie !")
                                st.download_button("📥 TÉLÉCHARGER L'EXCEL", data=buf_out.read(),
                                    file_name=nom_sortie,
                                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                    use_container_width=True, key="dl_xlsx_out")

                            # ── Excel → CSV ──────────────────────────────────
                            elif "Excel" in conv_type and "→ CSV" in conv_type:
                                import pandas as pd
                                from io import BytesIO as _BytesIO
                                df = pd.read_excel(_BytesIO(fichier_bytes))
                                csv_str = df.to_csv(index=False, encoding="utf-8-sig")
                                nom_sortie = conv_fichier.name.rsplit(".", 1)[0] + ".csv"
                                st.success("✅ Conversion réussie !")
                                st.download_button("📥 TÉLÉCHARGER LE CSV", data=csv_str.encode("utf-8-sig"),
                                    file_name=nom_sortie, mime="text/csv",
                                    use_container_width=True, key="dl_csv_out")

                            # ── PDF → Word via pdf2docx ───────────────────────
                            elif "PDF → Word" in conv_type:
                                try:
                                    from pdf2docx import Converter as PdfConverter
                                except ImportError:
                                    _sub.run(["pip", "install", "pdf2docx",
                                              "--break-system-packages", "-q"], check=True)
                                    from pdf2docx import Converter as PdfConverter
                                with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_in:
                                    tmp_in.write(fichier_bytes)
                                    tmp_in_path = tmp_in.name
                                tmp_out_path = tmp_in_path.replace(".pdf", ".docx")
                                cv = PdfConverter(tmp_in_path)
                                cv.convert(tmp_out_path, start=0, end=None)
                                cv.close()
                                with open(tmp_out_path, "rb") as f:
                                    buf_out = f.read()
                                nom_sortie = conv_fichier.name.replace(".pdf", "") + ".docx"
                                st.success("✅ Conversion réussie !")
                                st.download_button("📥 TÉLÉCHARGER LE WORD", data=buf_out,
                                    file_name=nom_sortie,
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True, key="dl_word_out")
                                try: _os.unlink(tmp_in_path); _os.unlink(tmp_out_path)
                                except: pass

                            # ── PDF → PPT via LibreOffice Impress ────────────
                            else:
                                with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_in:
                                    tmp_in.write(fichier_bytes)
                                    tmp_in_path = tmp_in.name
                                tmp_out_dir = tempfile.mkdtemp()
                                res = _sub.run(
                                    ["libreoffice", "--headless", "--convert-to", "pptx",
                                     "--outdir", tmp_out_dir, tmp_in_path],
                                    capture_output=True, text=True, timeout=90
                                )
                                nom_base  = _os.path.splitext(_os.path.basename(tmp_in_path))[0]
                                pptx_path = _os.path.join(tmp_out_dir, nom_base + ".pptx")
                                if _os.path.exists(pptx_path):
                                    with open(pptx_path, "rb") as f:
                                        buf_out = f.read()
                                    nom_sortie = conv_fichier.name.replace(".pdf", "") + ".pptx"
                                    st.success("✅ Conversion réussie !")
                                    st.download_button("📥 TÉLÉCHARGER LE POWERPOINT", data=buf_out,
                                        file_name=nom_sortie,
                                        mime="application/vnd.openxmlformats-officedocument.presentationml.presentation",
                                        use_container_width=True, key="dl_pptx_out")
                                else:
                                    st.error(f"❌ Erreur conversion PDF→PPT : {res.stderr or res.stdout}")
                                try: _os.unlink(tmp_in_path)
                                except: pass

                        except Exception as e:
                            st.error(f"❌ Erreur lors de la conversion : {e}")
            else:
                st.info("⬆️ Importez votre fichier pour démarrer la conversion")

            prompt = "CONVERSION_AUTO"  # Pas de prompt Gemini ni WhatsApp

        else:
            # ── CHAMP TEXTE LIBRE POUR LES AUTRES SERVICES ────────────────────
            prompt = st.text_area("Cahier des charges Nova", height=150, placeholder="Détaillez votre projet pour une exécution parfaite...")

        if service == SERVICE_SAISIE and service in SERVICE_PREREQUIS:
            if prompt and not st.session_state["warning_triggered"]:
                st.session_state["warning_triggered"] = True
                st.session_state["show_service_warning"] = True
                st.rerun()

            if st.session_state["show_service_warning"]:
                info = SERVICE_PREREQUIS[service]
                st.info(f"""
**{info["icone"]} {info["titre"]} — Informations requises**

{info["intro"]}

{"".join(f"- {icone} {texte}\n" for icone, texte in info["items"])}
💡 *{info["note"]}*
""")
                if os.path.exists("prerequis_excel.mp3"):
                    with open("prerequis_excel.mp3", "rb") as f:
                        b64_excel = __import__('base64').b64encode(f.read()).decode()
                    components.html(f"""
                        <script>
                        (function() {{
                            var binary = atob("{b64_excel}");
                            var bytes = new Uint8Array(binary.length);
                            for (var i = 0; i < binary.length; i++) bytes[i] = binary.charCodeAt(i);
                            var blob = new Blob([bytes], {{type: "audio/mpeg"}});
                            var audio = new Audio(URL.createObjectURL(blob));
                            audio.volume = 1;
                            audio.play().catch(function(e) {{ console.log(e); }});
                        }})();
                        </script>
                    """, height=0)
                col_mid = st.columns([1, 2, 1])[1]
                with col_mid:
                    if st.button("✅ J'ai compris, je continue ma demande", key="close_service_warning"):
                        st.session_state["show_service_warning"] = False
                        components.html("<script>window.speechSynthesis.cancel();</script>", height=0)
                        st.rerun()
        
        st.markdown("""
        <div class="logo-container">
            <svg class="logo-item" viewBox="0 0 24 24" fill="#217346"><path d="M16.2 21H2.8c-.4 0-.8-.4-.8-.8V3.8c0-.4.4-.8.8-.8h13.4c.4 0 .8.4.8.8v16.4c0 .4-.4.8-.8.8z"/><path d="M14.7 15.3l-2.2-3.3 2.2-3.3h-1.6l-1.4 2.2-1.4-2.2H8.7l2.2 3.3-2.2 3.3h1.6l1.4-2.2 1.4 2.2z" fill="white"/></svg>
            <svg class="logo-item" viewBox="0 0 24 24" fill="#2b579a"><path d="M16.2 21H2.8c-.4 0-.8-.4-.8-.8V3.8c0-.4.4-.8.8-.8h13.4c.4 0 .8.4.8.8v16.4c0 .4-.4.8-.8.8z"/><path d="M11.5 15.3V8.7h1.4c.8 0 1.4.3 1.8.8.4.5.6 1.1.6 1.8s-.2 1.3-.6 1.8c-.4.5-1 .8-1.8.8h-1.4z" fill="white"/></svg>
            <svg class="logo-item" viewBox="0 0 24 24" fill="#3776ab"><path d="M12 2C6.47 2 2 6.47 2 12s4.47 10 10 10 10-4.47 10-10S17.53 2 12 2zm-1 14.5h-1v-5h1v5zm0-6.5h-1V9h1v1z"/></svg>
            <svg class="logo-item" viewBox="0 0 24 24" fill="#d24726"><path d="M16.2 21H2.8c-.4 0-.8-.4-.8-.8V3.8c0-.4.4-.8.8-.8h13.4c.4 0 .8.4.8.8v16.4c0 .4-.4.8-.8.8z"/><path d="M8.7 8.7h1.5v5.1h2.5v1.5H8.7V8.7z" fill="white"/></svg>
            <svg class="logo-item" viewBox="0 0 24 24" fill="#ff9900"><path d="M12 2L2 7l10 5 10-5-10-5zM2 17l10 5 10-5M2 12l10 5 10-5"/></svg>
        </div>
        <p style="text-align:center; color:rgba(255,255,255,0.4); font-size:0.8rem; margin-top:5px;">Data • Dev • Design • Expertise • Rapidité</p>
        """, unsafe_allow_html=True)

        champs_manquants = []
        if not wa_display:
            champs_manquants.append("WhatsApp de contact")
        # Pour Sujets & Examens : le prompt est auto-construit via les selectbox
        # On ne vérifie PAS "Cahier des charges" mais les champs structurés
        if "Sujets" in service or "Examens" in service:
            if use_fichier_source:
                # Mode fichier : vérifier fichier + instructions
                if not contenu_fichier_source.strip():
                    champs_manquants.append("Fichier source")
                if not _fichier_instr.strip():
                    champs_manquants.append("Description du sujet souhaité")
            else:
                if not _niveau_val or _niveau_val.startswith("──"):
                    champs_manquants.append("Niveau scolaire")
                if not _matiere_val or _matiere_val.startswith("──"):
                    champs_manquants.append("Matière")
        elif "Fiche de Cours" in service:
            if not _fc_niveau_val or fc_niveau.startswith("──"):
                champs_manquants.append("Niveau scolaire")
            if not _fc_matiere_val or fc_matiere.startswith("──"):
                champs_manquants.append("Matière")
        elif "Modifier" in service and "Fichier" in service:
            if not mf_fichier:
                champs_manquants.append("Fichier à modifier")
            if not mf_instructions.strip():
                champs_manquants.append("Instructions de modification")
        else:
            if not prompt:
                champs_manquants.append("Cahier des charges")
        if champs_manquants:
            st.markdown(f"""
            <div style="
                background: rgba(241,196,15,0.08);
                border: 1px dashed rgba(241,196,15,0.4);
                border-radius: 10px;
                padding: 10px 16px;
                margin-top: 8px;
                color: rgba(241,196,15,0.85);
                font-size: 0.85rem;
            ">
                ⚠️ Veuillez s'il vous plaît détailler vos besoins comme il se doit, afin d'éviter que votre demande soit refusée.
            </div>
            """, unsafe_allow_html=True)

        if premium_actif and service in SERVICES_GEMINI:
            _udata_q = st.session_state["db"]["users"].get(user, {})
            _restant  = quota_restant(_udata_q)
            _plan_q   = _udata_q.get("premium_plan", "")
            _quota_q  = PLANS_PREMIUM.get(_plan_q, {}).get("generations", 0)
            _used_q, _ = get_gen_quota(_udata_q)
            _is_illimite = _quota_q >= 999
            _couleur_quota = "#FFD700" if _is_illimite else ("#2ecc71" if _restant > 1 else ("#FFD700" if _restant == 1 else "#e74c3c"))
            if _is_illimite:
                _quota_txt = "♾️ Illimité"
            elif _restant > 0:
                _quota_txt = f"{_used_q}/{_quota_q} utilisées — ✅ {_restant} restante(s)"
            else:
                _quota_txt = f"{_used_q}/{_quota_q} utilisées — 🚫 Quota atteint"
            st.markdown(f"""
            <div style="background:linear-gradient(135deg,rgba(255,215,0,.1),rgba(255,140,0,.06));
                 border:1px solid rgba(255,215,0,.5);border-radius:12px;padding:12px 18px;margin:10px 0;">
                <span style="color:#FFD700;font-weight:800;">⚡ PREMIUM — Génération IA automatique activée</span>
                <span style="color:rgba(255,255,255,.5);font-size:.8rem;display:block;margin-top:3px;">
                    Votre document sera généré et livré en moins d'1 minute.
                </span>
                <span style="color:{_couleur_quota};font-size:.85rem;font-weight:700;display:block;margin-top:6px;">
                    📊 Générations aujourd'hui : {_quota_txt}
                </span>
            </div>""", unsafe_allow_html=True)

        label_btn = "⚡ GÉNÉRER MAINTENANT AVEC L'IA NOVA" if (premium_actif and service in SERVICES_GEMINI) else "ACTIVER L'ALGORITHME NOVA"
        if "Conversion" not in service and st.button(label_btn):
            if not user:
                st.session_state["view"] = "auth"
                st.rerun()

            elif premium_actif and service in SERVICES_GEMINI and not champs_manquants:
                import threading

                # ── VÉRIFICATION DU QUOTA DE GÉNÉRATIONS ──────────────────
                user_data_frais = st.session_state["db"]["users"].get(user, {})
                restant = quota_restant(user_data_frais)
                plan_actuel = user_data_frais.get("premium_plan", "")
                quota_max = PLANS_PREMIUM.get(plan_actuel, {}).get("generations", 0)
                used_auj, _ = get_gen_quota(user_data_frais)

                if restant <= 0 and quota_max < 999:
                    st.error(f"🚫 Limite de générations atteinte pour aujourd'hui ({used_auj}/{quota_max} utilisées).")
                    st.info("💡 Votre quota se renouvelle demain, ou contactez Nova pour upgrader votre plan.")
                    # Basculer en mode demande manuelle
                    st.session_state["is_glowing"] = True
                    st.rerun()
                else:
                    processing_box = st.empty()
                    processing_box.markdown(f"""
                    <div class="nova-processing">
                        <div class="nova-processing-title">⚡ GÉNÉRATION EN COURS</div>
                        <div class="nova-processing-sub">Génération automatique · {'♾️ Illimité' if quota_max >= 999 else f'Quota restant après cette génération : {restant - 1}/{quota_max}'}</div>
                    </div>""", unsafe_allow_html=True)

                    barre = st.progress(0)
                    label_prog = st.empty()
                    t_start = time.time()
                    result_holder = {}

                    def generer():
                        try:
                            # Enrichir le prompt avec le type de sujet sélectionné (pour Sujets/Examens)
                            prompt_enrichi = prompt
                            if type_sujet_selectionne and ("Sujets" in service or "Examens" in service):
                                TYPE_SUJET_LABELS_FR = {
                                    "QCM": "QCM (Questions à Choix Multiple — 4 options A/B/C/D, cases □, UN SEUL TYPE)",
                                    "VRAI_FAUX": "VRAI ou FAUX UNIQUEMENT (V/F + justification si faux, UN SEUL TYPE)",
                                    "TEXTE_TROU": "TEXTE À TROUS UNIQUEMENT (lacunaire + liste de mots, UN SEUL TYPE)",
                                    "QUESTIONS_OUVERTES": "QUESTIONS OUVERTES UNIQUEMENT (rédigées + lignes de réponse, UN SEUL TYPE)",
                                    "MIXTE": "FORMAT MIXTE (Partie 1 QCM + Partie 2 Vrai/Faux + Partie 3 Question ouverte)",
                                    "CAS_PRATIQUE": "CAS PRATIQUE / ÉTUDE DE CAS (texte CI contextualisé + questions d'analyse)",
                                    "CALCUL": "EXERCICES DE CALCUL / PROBLÈMES (chiffrés, contextualisés CI, formules rappelées)",
                                    "ETUDE_DOCUMENT": "ÉTUDE DE DOCUMENT (document support + questions d'exploitation)",
                                    "SCHEMA": "SCHÉMA À LÉGENDER (description numérotée + termes à placer + corrigé)",
                                    "DISSERTATION": "COMPOSITION / DISSERTATION GUIDÉE (sujet + méthode + plan guidé)",
                                    "DEVOIR_COMPLET": "DEVOIR COMPLET AUTHENTIQUE IVOIRIEN",
                                }
                                label_fr = TYPE_SUJET_LABELS_FR.get(type_sujet_selectionne, type_sujet_selectionne)
                                prompt_enrichi = f"""{prompt}

⚠️ TYPE DE SUJET IMPOSÉ PAR L'UTILISATEUR — RESPECTER ABSOLUMENT :
TYPE UNIQUE SÉLECTIONNÉ : {label_fr}

RÈGLE ABSOLUE : Génère UNIQUEMENT ce type d'exercice. Ne pas mélanger avec d'autres types sauf si MIXTE ou DEVOIR_COMPLET est sélectionné.
Si QCM → QCM seulement. Si VRAI_FAUX → Vrai/Faux seulement. Si TEXTE_TROU → Texte à trous seulement.
Si QUESTIONS_OUVERTES → Questions ouvertes seulement. Si CALCUL → Calculs seulement.
Si ETUDE_DOCUMENT → Étude de document seulement. Si SCHEMA → Schéma à légender seulement.
Si DISSERTATION → Composition guidée seulement. Si CAS_PRATIQUE → Cas pratique seulement.
Si DEVOIR_COMPLET → Vrai devoir ivoirien COMPLET : applique EXACTEMENT la Section Critique (en-tête officiel + structure progressive adaptée au niveau détecté). Mélange intelligent de tous les types adaptés au niveau."""
                            contenu = generer_avec_gemini(service, prompt_enrichi, user)
                            if contenu.startswith("❌"):
                                result_holder["erreur"] = contenu
                                return
                            if service == "📊 Data & Excel Analytics":
                                buf  = creer_xlsx(prompt_enrichi, user)
                                nom  = f"{user}_{service[:20].strip()}.xlsx".replace(" ", "_").replace("/", "-")
                                mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                            else:
                                buf  = creer_docx(contenu, service, user)
                                type_suffix = f"_{type_sujet_selectionne}" if type_sujet_selectionne else ""
                                nom  = f"{user}_{service[:20].strip()}{type_suffix}.docx".replace(" ", "_").replace("/", "-")
                                mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            result_holder["buf"]  = buf
                            result_holder["nom"]  = nom
                            result_holder["mime"] = mime
                        except Exception as e:
                            result_holder["erreur"] = f"❌ Erreur : {e}"

                    thread = threading.Thread(target=generer)
                    thread.start()

                    pct = 0
                    while thread.is_alive():
                        elapsed = time.time() - t_start
                        pct = min(int(elapsed / 60 * 90), 90)
                        barre.progress(pct)
                        label_prog.markdown(f"<p style='text-align:center;color:#FFD700;font-weight:bold;'>⚡ Génération en cours... {pct}%</p>", unsafe_allow_html=True)
                        time.sleep(0.5)
                    thread.join()

                    barre.progress(100)
                    label_prog.markdown("<p style='text-align:center;color:#2ecc71;font-weight:bold;'>✅ Document généré !</p>", unsafe_allow_html=True)
                    time.sleep(0.8)
                    barre.empty(); label_prog.empty(); processing_box.empty()

                    duree = int(time.time() - t_start)

                    if "erreur" in result_holder:
                        st.error(result_holder["erreur"])
                        st.info("💡 Votre demande a été transmise à l'équipe Nova pour traitement manuel.")
                        new_req = {
                            "id": hashlib.md5(str(datetime.now()).encode()).hexdigest()[:8],
                            "user": user, "service": service,
                            "desc": prompt, "whatsapp": normalize_wa(wa_display),
                            "status": "Traitement Nova en cours...", "incomplet": False,
                            "champs_manquants": [], "timestamp": str(datetime.now()),
                        }
                        st.session_state["db"]["demandes"].append(new_req)
                        save_demande(new_req)
                    else:
                        # Incrémenter le compteur de générations
                        incrementer_gen(user)
                        save_lien(user, service, f"__local__{result_holder['nom']}", datetime.now().strftime("%d/%m/%Y"))
                        # Email admin — Gemini a déjà répondu
                        wa_display_local = st.session_state["db"]["users"].get(user, {}).get("whatsapp", "—")
                        envoyer_notification_gemini_ok(user, wa_display_local, service, result_holder["nom"], demande_complete=prompt)
                        st.session_state["premium_livrable"] = {
                            "buf":     result_holder["buf"],
                            "nom":     result_holder["nom"],
                            "mime":    result_holder["mime"],
                            "service": service,
                            "duree":   duree,
                        }
                        st.session_state["db"] = load_db()
                        st.rerun()

            else:
                st.session_state["is_glowing"] = True
                st.rerun()

        if st.session_state["is_glowing"]:
            progress_placeholder = st.empty()
            status_text = st.empty()
            bar = progress_placeholder.progress(0)
            for percent_complete in range(100):
                time.sleep(0.02)
                bar.progress(percent_complete + 1)
                status_text.markdown(f"<p style='text-align:center; color:#00d2ff; font-size:1.2rem; font-weight:bold;'>NOVA PROCESSING : {percent_complete + 1}%</p>", unsafe_allow_html=True)

            statut = "En attente de vérification (informations incomplètes)" if champs_manquants else "Traitement Nova en cours..."

            req_id_new = hashlib.md5(str(datetime.now()).encode()).hexdigest()[:8]
            fichier_info = ""
            if "Modifier" in service and "Fichier" in service and mf_fichier:
                with st.spinner("📤 Upload du fichier..."):
                    url_f = upload_fichier_client(
                        user if user else "guest", req_id_new,
                        mf_fichier.getvalue(), mf_fichier.name)
                if url_f.startswith("ERREUR"):
                    st.warning(f"⚠️ Fichier non uploadé ({url_f})")
                else:
                    fichier_info = f"\n📎 FICHIER CLIENT : {url_f}"
            desc_finale = (prompt if prompt else "(aucune description fournie)") + fichier_info
            new_req = {
                "id": req_id_new,
                "user": user if user else "guest",
                "service": service,
                "desc": desc_finale,
                "whatsapp": normalize_wa(wa_display) if wa_display else "(non renseigné)",
                "status": statut,
                "incomplet": bool(champs_manquants),
                "champs_manquants": champs_manquants,
                "timestamp": str(datetime.now())
            }
            st.session_state["db"]["demandes"].append(new_req)
            save_demande(new_req)
            envoyer_notification(
                client_nom  = user if user else "Visiteur",
                client_wa   = normalize_wa(wa_display) if wa_display else "(non renseigné)",
                service     = service,
                description = desc_finale
            )
            st.session_state["db"] = load_db()
            st.session_state["is_glowing"] = False
            progress_placeholder.empty()
            status_text.empty()
            if user:
                st.success("✅ Mission enregistrée ! L'équipe Nova examinera votre demande.")

                audio_path_confirm = "confirmation.mp3"
                if os.path.exists(audio_path_confirm):
                    with open(audio_path_confirm, "rb") as f:
                        audio_b64_confirm = __import__('base64').b64encode(f.read()).decode()
                    components.html(f"""
                        <script>
                        (function() {{
                            var b64 = "{audio_b64_confirm}";
                            var binary = atob(b64);
                            var bytes = new Uint8Array(binary.length);
                            for (var i = 0; i < binary.length; i++) {{
                                bytes[i] = binary.charCodeAt(i);
                            }}
                            var blob = new Blob([bytes], {{type: "audio/mpeg"}});
                            var url = URL.createObjectURL(blob);
                            var audio = new Audio(url);
                            audio.volume = 1;
                            audio.play().catch(function(e) {{ console.log("Autoplay bloqué:", e); }});
                        }})();
                        </script>
                    """, height=0)

                st.rerun()
            else:
                st.session_state["view"] = "auth"
                st.rerun()

    with tab2:
        if not user:
            st.warning("🔒 Authentification requise pour accéder au Cloud Nova.")
        else:
            fresh_db = load_db()
            user_links = fresh_db["liens"].get(user, [])
            user_reqs = [r for r in fresh_db["demandes"] if r["user"] == user]
            
            st.markdown("""
                <div style="background: rgba(46, 204, 113, 0.1); padding: 15px; border-radius: 10px; border: 1px dashed #2ecc71; margin-bottom: 20px; text-align: center;">
                    <h2 style="color: #2ecc71; margin: 0;">📥 HUB DE TÉLÉCHARGEMENT NOVA</h2>
                    <p style="color: white; font-size: 0.9rem;">Accédez à vos actifs numériques terminés.</p>
                </div>
            """, unsafe_allow_html=True)

            if st.session_state["premium_livrable"]:
                lv = st.session_state["premium_livrable"]
                st.markdown(f"""
                <div class="livrable-auto">
                    <div class="livrable-auto-title">⚡ Livrable IA Premium</div>
                    <div style="color:rgba(255,255,255,.7);margin-top:4px;">Généré en {lv['duree']}s · {lv['service']}</div>
                </div>""", unsafe_allow_html=True)
                st.download_button(
                    "📥 TÉLÉCHARGER MON DOCUMENT",
                    data=lv["buf"], file_name=lv["nom"], mime=lv["mime"],
                    use_container_width=True
                )
                st.divider()

            if user_links:
                col_titre, col_vider = st.columns([3, 1])
                with col_titre:
                    st.markdown("#### 📁 Mes livrables")
                with col_vider:
                    if st.button("🗑️ Vider l'historique", key="vider_historique", use_container_width=True):
                        delete_all_liens(user)
                        st.success("Historique vidé !")
                        st.rerun()

                for link in user_links:
                    if link["url"].startswith("__refus__"):
                        msg_refus = link["url"].replace("__refus__", "")
                        st.markdown(f"""
                        <div class="file-card" style="border-color:rgba(231,76,60,0.5); background:rgba(231,76,60,0.05);">
                            <div style="display:flex;justify-content:space-between;align-items:center;">
                                <div>
                                    <h3 style="color:#e74c3c;margin:0;">❌ Mission refusée — {link['name']}</h3>
                                    <p style="color:#aaa;font-size:.85rem;margin:5px 0;">Le {link.get('date','')}</p>
                                    <p style="color:#eee;font-size:.9rem;margin-top:8px;">{msg_refus}</p>
                                </div>
                            </div>
                        </div>""", unsafe_allow_html=True)
                    elif link["url"].startswith("__local__"):
                        st.markdown(f"""
                        <div class="file-card" style="border-color:rgba(255,215,0,.5);">
                            <div style="display:flex;justify-content:space-between;align-items:center;">
                                <div>
                                    <h3 style="color:#FFD700;margin:0;">⭐ {link['name']}</h3>
                                    <p style="color:#aaa;font-size:.85rem;margin:5px 0;">Généré le {link.get('date',"Aujourd'hui")} · Téléchargez depuis l'onglet Déployer</p>
                                </div>
                                <span class="badge-premium">IA AUTO</span>
                            </div>
                        </div>""", unsafe_allow_html=True)
                    else:
                        st.markdown(f"""
                        <div class="file-card">
                            <div style="display: flex; justify-content: space-between; align-items: center;">
                                <div>
                                    <h3 style="color:#00d2ff; margin:0;">💎 {link['name']}</h3>
                                    <p style="color:#aaa; font-size:0.85rem; margin: 5px 0;">Finalisé le {link.get('date', "Aujourd'hui")}</p>
                                </div>
                                <a href="{link['url']}" target="_blank" style="text-decoration:none;">
                                    <button style="padding:10px 25px; background:#2ecc71; color:white; border:none; border-radius:30px; font-weight:bold; cursor:pointer; box-shadow: 0 4px 10px rgba(46,204,113,0.3);">
                                        📥 TÉLÉCHARGER
                                    </button>
                                </a>
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
            
            if user_reqs:
                st.markdown("#### ⏳ Missions Nova en préparation")
                for r in user_reqs:
                    st.markdown(f"""
                        <div class="file-card" style="border-left: 5px solid #f1c40f; border-color: rgba(241, 196, 15, 0.3);">
                            <div style="display: flex; justify-content: space-between; align-items: center;">
                                <div>
                                    <strong style="color: #f1c40f;">{r['service']}</strong><br>
                                    <span style="color:#eee; font-size: 0.9rem;">Status: {r['status']}</span>
                                </div>
                                <div class="spinner" style="width: 20px; height: 20px; border: 3px solid rgba(255,255,255,0.1); border-top: 3px solid #f1c40f; border-radius: 50%; animation: spin 1s linear infinite;"></div>
                            </div>
                        </div>
                        <style>@keyframes spin {{ 0% {{ transform: rotate(0deg); }} 100% {{ transform: rotate(360deg); }} }}</style>
                    """, unsafe_allow_html=True)
            
            if not user_links and not user_reqs:
                st.info("Votre espace Nova est vide. Déployez votre première tâche !")
            
            st.write("---")
            st.markdown("### 🆘 Support Nova Direct")
            col_rel, col_sup = st.columns(2)
            with col_rel:
                relance_msg = f"Bonjour, je souhaite un status sur ma mission Nova (ID: {user})."
                wa_relance = f"https://wa.me/{WHATSAPP_NUMBER}?text={relance_msg.replace(' ', '%20')}"
                st.markdown(f'<a href="{wa_relance}" target="_blank" class="support-btn" style="border-color:#f1c40f; color:#f1c40f !important;">🔔 Relancer Nova</a>', unsafe_allow_html=True)
            with col_sup:
                st.markdown(f'<a href="{whatsapp_support_url}" target="_blank" class="support-btn">🙋 Agent Nova</a>', unsafe_allow_html=True)

    with st.expander("🛠 Console Admin Nova"):
        if st.text_input("Master Key", type="password") == ADMIN_CODE:

            current_db = st.session_state["db"]
            admin_tab1, admin_tab2 = st.tabs(["📋 MISSIONS", "👑 GESTION PREMIUM"])

            with admin_tab1:
                st.markdown("### 🛡️ Panneau de contrôle Nova")

                # ── TOGGLE RÉPONSE AUTOMATIQUE PLAN GRATUIT ───────────────
                col_toggle_l, col_toggle_r = st.columns([3, 1])
                with col_toggle_l:
                    _auto_status = "🟢 ACTIVÉ" if st.session_state["auto_reply_gratuit"] else "🔴 DÉSACTIVÉ"
                    st.markdown(f"""
                    <div style="background:rgba(255,255,255,0.05);border:1px solid rgba(255,255,255,0.15);
                         border-radius:10px;padding:12px 16px;margin-bottom:12px;">
                        <span style="font-weight:700;color:#FFD700;">🤖 Réponse automatique — Plan Gratuit</span>
                        <span style="color:rgba(255,255,255,0.5);font-size:0.82rem;display:block;margin-top:3px;">
                            Si activé : Arsène AI répond automatiquement aux demandes gratuites après <b>2 heures</b> d'attente, sans validation manuelle.
                        </span>
                        <span style="font-weight:800;font-size:0.95rem;margin-top:6px;display:block;">
                            Statut actuel : {_auto_status}
                        </span>
                    </div>
                    """, unsafe_allow_html=True)
                with col_toggle_r:
                    st.markdown("<br>", unsafe_allow_html=True)
                    if st.session_state["auto_reply_gratuit"]:
                        if st.button("🔴 Désactiver", key="btn_toggle_auto", use_container_width=True):
                            st.session_state["auto_reply_gratuit"] = False
                            set_auto_reply_setting(False)
                            st.success("✅ Réponse automatique désactivée")
                            st.rerun()
                    else:
                        if st.button("🟢 Activer", key="btn_toggle_auto", use_container_width=True):
                            st.session_state["auto_reply_gratuit"] = True
                            set_auto_reply_setting(True)
                            st.success("✅ Réponse automatique activée — Arsène AI répondra après 2h")
                            st.rerun()

                st.divider()

                if not current_db["demandes"]:
                    st.info("✅ Aucune mission en attente.")

                def wa_url(numero, texte):
                    encoded = texte.replace(" ", "%20").replace("'", "%27").replace("\n", "%0A")
                    return f"https://wa.me/{numero}?text={encoded}"

                for i, req in enumerate(current_db["demandes"]):
                    client_wa_raw    = req.get("whatsapp", "(non renseigné)")
                    client_wa        = normalize_wa(client_wa_raw)
                    client_nom       = req.get("user", "Inconnu")
                    service          = req.get("service", "—")
                    description      = req.get("desc", "(aucune description)")
                    req_id           = req.get("id", f"{i+1}")
                    timestamp        = req.get("timestamp", "")[:16] if req.get("timestamp") else "—"
                    est_incomplet    = req.get("incomplet", False)
                    champs_manquants = req.get("champs_manquants", [])
                    client_premium   = is_premium_actif(current_db["users"].get(client_nom, {}))

                    if i > 0:
                        st.divider()

                    col_info, col_badge = st.columns([4, 1])
                    with col_info:
                        st.markdown(f"**Mission `#{req_id}`** · {timestamp}" + (" — ⚠️ *Incomplet : " + ", ".join(champs_manquants) + "*" if est_incomplet else ""))
                        st.markdown(f"👤 **Client :** {client_nom}")
                        st.markdown(f"📱 **WhatsApp :** {client_wa}")
                        st.markdown(f"🛠️ **Service demandé :** {service}")
                        st.markdown(f"📝 **Détails de la demande :** {description}")
                        if "Modifier" in service and "Fichier" in service:
                            _url_dl = None
                            _nom_dl = "fichier_client"
                            for _l in description.split("\n"):
                                if "📎 FICHIER CLIENT :" in _l:
                                    _url_dl = _l.replace("📎 FICHIER CLIENT :", "").strip()
                                if "FICHIER" in _l and ":" in _l and "MODIF" not in _l and "NOTE" not in _l and "📎 FICHIER CLIENT" not in _l:
                                    _nom_dl = _l.split(":")[-1].strip()
                            if _url_dl and _url_dl.startswith("http"):
                                st.markdown(f'''<a href="{_url_dl}" target="_blank"
                                   style="display:inline-flex;align-items:center;gap:8px;
                                   background:rgba(0,210,255,0.12);border:1px solid rgba(0,210,255,0.5);
                                   border-radius:10px;padding:10px 20px;color:#00d2ff;
                                   font-weight:700;text-decoration:none;margin-top:6px;">
                                   📥 TÉLÉCHARGER LE FICHIER — {_nom_dl}</a>''',
                                   unsafe_allow_html=True)
                            else:
                                st.info("📎 Fichier non disponible")
                    with col_badge:
                        if client_premium:
                            st.markdown('<span class="badge-premium">⭐ PREMIUM</span>', unsafe_allow_html=True)
                        else:
                            st.markdown('<span class="badge-free">🔓 Gratuit</span>', unsafe_allow_html=True)

                    if est_incomplet and champs_manquants:
                        champs_str = ", ".join(champs_manquants)
                        msg_rejet = (f"Bonjour {client_nom}, nous avons reçu votre demande Nova AI "
                                     f"concernant : {service}. Cependant, nous ne pouvons pas la traiter "
                                     f"car les informations suivantes sont manquantes : {champs_str}. "
                                     f"Merci de soumettre à nouveau votre demande en complétant tous les champs. "
                                     f"— Équipe Nova AI ⚡")
                    else:
                        msg_rejet = (f"Bonjour {client_nom}, nous avons bien reçu votre demande Nova AI "
                                     f"concernant : {service}. Malheureusement, nous ne sommes pas en mesure "
                                     f"de traiter cette mission pour le moment. Merci de nous recontacter. "
                                     f"— Équipe Nova AI ⚡")

                    msg_succes = (f"✅ Bonjour {client_nom} ! Votre mission Nova AI ({service}) est terminée ! "
                                  f"Rendez-vous dans votre espace Nova pour récupérer votre livrable. "
                                  f"Merci de votre confiance. — Équipe Nova AI ⚡")

                    msg_recu = (f"📬 Bonjour {client_nom}, nous confirmons la réception de votre demande "
                                f"Nova AI : {service}. Votre mission est en cours de traitement. "
                                f"Vous serez notifié dès qu'elle sera finalisée. — Équipe Nova AI ⚡")

                    col_rejet, col_recu, col_succes = st.columns(3)
                    with col_rejet:
                        if st.button("❌ Rejeter", key=f"rej_{i}", use_container_width=True):
                            save_refus(client_nom, service, msg_rejet)
                            delete_demande(req["id"])
                            st.session_state["db"] = load_db()
                            st.success(f"Mission refusée — notification envoyée dans les livrables de {client_nom}.")
                            st.rerun()
                    with col_recu:
                        st.markdown(f'<a href="{wa_url(client_wa, msg_recu)}" target="_blank" style="display:block; text-align:center; padding:10px; border-radius:10px; background:rgba(255,215,0,0.1); border:1px solid rgba(255,215,0,0.4); color:#FFD700; font-weight:700; text-decoration:none;">📬 Reçu</a>', unsafe_allow_html=True)
                    with col_succes:
                        st.markdown(f'<a href="{wa_url(client_wa, msg_succes)}" target="_blank" style="display:block; text-align:center; padding:10px; border-radius:10px; background:rgba(46,204,113,0.15); border:1px solid rgba(46,204,113,0.5); color:#2ecc71; font-weight:700; text-decoration:none;">✅ Succès</a>', unsafe_allow_html=True)

                    if service in SERVICES_GEMINI:
                        st.markdown("<br>", unsafe_allow_html=True)
                        st.markdown(f"""
                        <div class="gemini-card">
                            <div class="gemini-title">🤖 Arsène AI — GÉNÉRATION AUTOMATIQUE DISPONIBLE</div>
                            <div class="gemini-sub">Génère le document complet en .docx en 30-60 secondes</div>
                        </div>
                        """, unsafe_allow_html=True)

                        if st.button(f"🔍 Voir modèles disponibles", key=f"diag_{req_id}"):
                            with st.spinner("Interrogation de l'API Arsène AI..."):
                                modeles_dispo = get_modeles_disponibles(st.secrets["GEMINI_API_KEY"])
                            if modeles_dispo:
                                st.success(f"✅ {len(modeles_dispo)} modèles trouvés :")
                                for m in modeles_dispo:
                                    st.code(m)
                            else:
                                st.error("❌ Aucun modèle disponible — vérifiez votre clé API.")

                        if st.button(f"⚡ APPROUVER & GÉNÉRER AVEC ARSÈNE AI", key=f"gemini_{req_id}", use_container_width=True):
                            with st.spinner("🔍 Détection automatique du meilleur modèle disponible..."):
                                modeles_dispo = get_modeles_disponibles(st.secrets["GEMINI_API_KEY"])
                                if modeles_dispo:
                                    st.info(f"✅ Modèle sélectionné : **{modeles_dispo[0]}**")
                                else:
                                    st.error("❌ Aucun modèle Gemini disponible pour cette clé API.")
                            with st.spinner("🤖 Arsène AI génère le document... (30-60 secondes)"):
                                contenu = generer_avec_gemini(service, description, client_nom)

                            if contenu.startswith("❌"):
                                st.error(contenu)
                            else:
                                st.session_state["gemini_results"][req_id] = {
                                    "contenu": contenu,
                                    "service": service,
                                    "client": client_nom
                                }
                                st.success("✅ Document généré avec succès !")
                                st.rerun()

                        if req_id in st.session_state["gemini_results"]:
                            result = st.session_state["gemini_results"][req_id]

                            with st.expander("👁️ Aperçu du contenu généré", expanded=False):
                                st.markdown(result["contenu"])

                            if st.button("🚀 LIVRER DIRECTEMENT AU CLIENT", key=f"livrer_auto_{req_id}", use_container_width=True):
                                try:
                                    SERVICE_EXCEL = "📊 Data & Excel Analytics"
                                    if result["service"] == SERVICE_EXCEL:
                                        buf = creer_xlsx(result.get("desc", ""), result["client"])
                                        nom_fichier = f"{client_nom}_Suivi_Depenses.xlsx".replace(" ", "_")
                                        mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                                    else:
                                        buf = creer_docx(result["contenu"], result["service"], result["client"])
                                        nom_fichier = f"{client_nom}_{result['service'][:20].strip()}.docx".replace(" ", "_").replace("/", "-")
                                        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

                                    with st.spinner("⬆️ Upload du fichier vers Supabase..."):
                                        url_fichier = upload_fichier_client(client_nom, req_id, buf, nom_fichier)

                                    if url_fichier.startswith("ERREUR"):
                                        st.error(f"❌ Upload échoué : {url_fichier}")
                                    else:
                                        save_lien(client_nom, service, url_fichier, datetime.now().strftime("%d/%m/%Y"))
                                        delete_demande(req["id"])
                                        if req_id in st.session_state["gemini_results"]:
                                            del st.session_state["gemini_results"][req_id]
                                        st.session_state["db"] = load_db()
                                        st.success(f"✅ Fichier livré directement dans les livrables de {client_nom} !")
                                        st.rerun()
                                except Exception as e:
                                    st.error(f"Erreur livraison : {e}")

                    st.markdown("<br>", unsafe_allow_html=True)
                    url_dl = st.text_input("🔗 Lien de livraison manuel (optionnel)", key=f"url_{i}", placeholder="https://drive.google.com/...")
                    if st.button("📦 LIVRER AVEC LIEN MANUEL", key=f"btn_{i}", use_container_width=True):
                        if url_dl:
                            save_lien(req['user'], req['service'], url_dl, datetime.now().strftime("%d/%m/%Y"))
                            delete_demande(req['id'])
                            if req_id in st.session_state["gemini_results"]:
                                del st.session_state["gemini_results"][req_id]
                            st.session_state["db"] = load_db()
                            st.success(f"✅ Mission livrée à {client_nom} !")
                            st.rerun()

            with admin_tab2:
                st.markdown("### 👑 Gestion des membres Premium")
                total  = len(current_db["users"])
                prems  = [u for u, d in current_db["users"].items() if is_premium_actif(d)]
                c1, c2, c3 = st.columns(3)
                c1.metric("👥 Total membres", total)
                c2.metric("⭐ Premium actifs", len(prems))
                c3.metric("🔓 Gratuits", total - len(prems))
                st.divider()

                st.markdown("#### ➕ Activer / Gérer un Premium")
                co1, co2, co3 = st.columns([2, 2, 1])
                with co1:
                    uid_target = st.selectbox("Membre", options=list(current_db["users"].keys()),
                        format_func=lambda u: f"{u} {'⭐' if is_premium_actif(current_db['users'][u]) else '🔓'}")
                with co2:
                    plan_ch = st.selectbox("Plan", list(PLANS_PREMIUM.keys()),
                        format_func=lambda p: f"{PLANS_PREMIUM[p]['emoji']} {p} — {PLANS_PREMIUM[p]['prix']}")
                with co3:
                    st.markdown("<br>", unsafe_allow_html=True)
                    if st.button("⚡ ACTIVER", key="btn_act_global"):
                        activer_premium(uid_target, plan_ch)
                        st.session_state["db"] = load_db()
                        st.success(f"✅ Premium **{plan_ch}** activé pour **{uid_target}** !")
                        st.rerun()

                st.divider()
                filtre = st.radio("Afficher", ["Tous", "Premium uniquement", "Gratuits uniquement"], horizontal=True)

                for uid_m, udata in current_db["users"].items():
                    p_actif = is_premium_actif(udata)
                    p_info  = get_premium_info(udata)
                    if filtre == "Premium uniquement" and not p_actif: continue
                    if filtre == "Gratuits uniquement" and p_actif:    continue

                    col_m, col_a = st.columns([3, 2])
                    with col_m:
                        badge = f'<span class="badge-premium">⭐ {udata.get("premium_plan","—")}</span>' if p_actif else '<span class="badge-free">🔓 Gratuit</span>'
                        exp_txt = f"<br><small style='color:rgba(255,215,0,.6);'>Expire : {p_info['expiry']} ({p_info['jours_restants']}j)</small>" if p_actif and p_info else ""
                        st.markdown(f"""<div class="admin-premium-row">
                            <div>
                                <div class="admin-user-name">👤 {uid_m}</div>
                                <div class="admin-user-meta">📱 {udata.get('whatsapp','—')} · {str(udata.get('joined',''))[:10]}</div>
                                {exp_txt}
                            </div>
                            <div>{badge}</div>
                        </div>""", unsafe_allow_html=True)
                    with col_a:
                        if p_actif:
                            cp1, cp2 = st.columns(2)
                            with cp1:
                                ext_p = st.selectbox("", list(PLANS_PREMIUM.keys()), key=f"ext_{uid_m}",
                                    format_func=lambda p: f"{PLANS_PREMIUM[p]['emoji']} {p}")
                                if st.button("➕ Prolonger", key=f"pro_{uid_m}"):
                                    curr_exp = datetime.fromisoformat(udata.get("premium_expiry", datetime.now().isoformat()))
                                    new_exp  = max(curr_exp, datetime.now()) + timedelta(days=PLANS_PREMIUM[ext_p]["jours"])
                                    update_premium_status(uid_m, True, ext_p, new_exp.isoformat())
                                    st.session_state["db"] = load_db()
                                    st.success(f"✅ Prolongé jusqu'au {new_exp.strftime('%d/%m/%Y')} !")
                                    st.rerun()
                            with cp2:
                                st.markdown("<br>", unsafe_allow_html=True)
                                if st.button("🗑️ Révoquer", key=f"rev_{uid_m}"):
                                    desactiver_premium(uid_m)
                                    st.session_state["db"] = load_db()
                                    st.warning(f"Premium révoqué pour {uid_m}.")
                                    st.rerun()
                        else:
                            ap = st.selectbox("", list(PLANS_PREMIUM.keys()), key=f"act_{uid_m}",
                                format_func=lambda p: f"{PLANS_PREMIUM[p]['emoji']} {p}")
                            if st.button("⚡ Activer", key=f"actbtn_{uid_m}"):
                                activer_premium(uid_m, ap)
                                st.session_state["db"] = load_db()
                                st.success(f"✅ Premium activé pour {uid_m} !")
                                st.rerun()


inject_custom_css()

components.html("""
    <script>
    const user = localStorage.getItem('nova_user');
    const urlParams = new URLSearchParams(window.parent.location.search);
    const currentUser = urlParams.get('user_id');
    
    if (user && !currentUser && !window.parent.location.href.includes('logout')) {
        window.parent.location.href = window.parent.location.origin + window.parent.location.pathname + '?user_id=' + user;
    }
    if (!currentUser && user && window.parent.location.href.includes('logout')) {
        localStorage.removeItem('nova_user');
    }
    if (currentUser && user !== currentUser) {
        localStorage.setItem('nova_user', currentUser);
    }
    </script>
""", height=0)

# Masquer l'iframe vide créée par components.html
st.markdown("""
    <style>
    iframe[title="components.v1.html"] { display: none !important; height: 0 !important; }
    </style>
""", unsafe_allow_html=True)

if st.session_state["view"] == "auth" and st.session_state["current_user"] is None:
    show_auth_page()
else:
    main_dashboard()
